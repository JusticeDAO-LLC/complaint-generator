from __future__ import annotations

from datetime import datetime, timezone
from html import escape
import json
from pathlib import Path
import re
from typing import Any, Dict, Iterable, List, Optional
from urllib.parse import urlencode

from complaint_phases import ComplaintPhase
from document_optimization import AgenticDocumentOptimizer
from intake_status import build_intake_case_review_summary, build_intake_status_summary
from workflow_phase_guidance import (
    build_drafting_document_generation_phase_guidance,
    build_graph_analysis_phase_guidance,
    build_workflow_phase_plan,
    build_workflow_phase_warning_entries,
)


DEFAULT_OUTPUT_DIR = Path(__file__).resolve().parent / "tmp" / "generated_documents"
DEFAULT_RELIEF = [
    "Compensatory damages in an amount to be proven at trial.",
    "Pre- and post-judgment interest as allowed by law.",
    "Reasonable attorney's fees and costs where authorized.",
    "Injunctive and declaratory relief sufficient to stop the unlawful conduct.",
    "Such other and further relief as the Court deems just and proper.",
]

STATE_DEFAULT_RELIEF = [
    "General and special damages according to proof.",
    "Costs of suit incurred herein.",
    "Such other and further relief as the Court deems just and proper.",
]
ACTOR_CRITIC_PHASE_FOCUS_ORDER = ("graph_analysis", "document_generation", "intake_questioning")
DEFAULT_ACTOR_CRITIC_PRIORITY_METRICS = {
    "empathy": 0.22,
    "question_quality": 0.58,
    "information_extraction": 0.40,
    "coverage": 0.40,
    "efficiency": 0.62,
}
_CONFIRMATION_PLACEHOLDER_PATTERN = re.compile(
    r"\b(?:needs?\s+confirmation|to\s+be\s+confirmed|confirm(?:ed|ation)?\s+pending|tbd|unknown|not\s+sure|unclear|pending)\b",
    flags=re.IGNORECASE,
)


def _utcnow() -> datetime:
    return datetime.now(timezone.utc)


def _slugify(value: str) -> str:
    text = re.sub(r"[^A-Za-z0-9]+", "-", str(value or "document").strip().lower())
    return text.strip("-") or "document"


def _unique_preserving_order(values: Iterable[str]) -> List[str]:
    seen = set()
    unique: List[str] = []
    for value in values:
        text = str(value or "").strip()
        if not text:
            continue
        key = text.lower()
        if key in seen:
            continue
        seen.add(key)
        unique.append(text)
    return unique


def _coerce_list(value: Any) -> List[Any]:
    if isinstance(value, list):
        return value
    if value is None:
        return []
    return [value]


def _dedupe_text_values(values: Iterable[Any]) -> List[str]:
    seen = set()
    normalized_values: List[str] = []
    for value in values:
        text = str(value or "").strip()
        if not text:
            continue
        if text in seen:
            continue
        seen.add(text)
        normalized_values.append(text)
    return normalized_values


def _format_timeline_date(value: Any) -> str:
    text = str(value or "").strip()
    if not text:
        return ""
    try:
        return datetime.strptime(text, "%Y-%m-%d").strftime("%B %-d, %Y")
    except ValueError:
        return text


def _chronology_fact_label(fact: Dict[str, Any]) -> str:
    event_label = str((fact if isinstance(fact, dict) else {}).get("event_label") or "").strip()
    if event_label:
        return event_label
    predicate_family = str((fact if isinstance(fact, dict) else {}).get("predicate_family") or "").strip().replace("_", " ")
    if predicate_family:
        return predicate_family.title()
    return "Event"


def _join_chronology_segments(segments: List[str]) -> str:
    if not segments:
        return ""
    if len(segments) == 1:
        return segments[0]
    if len(segments) == 2:
        return f"{segments[0]} and {segments[1]}"
    return f"{', '.join(segments[:-1])}, and {segments[-1]}"


def _build_anchored_chronology_summary_from_case_file(intake_case_file: Dict[str, Any], *, limit: int = 3) -> List[str]:
    case_file = intake_case_file if isinstance(intake_case_file, dict) else {}
    facts = [dict(item) for item in list(case_file.get("canonical_facts") or []) if isinstance(item, dict)]
    relations = [dict(item) for item in list(case_file.get("timeline_relations") or []) if isinstance(item, dict)]
    if not facts or not relations:
        return []

    fact_by_id = {
        str(fact.get("fact_id") or "").strip(): fact
        for fact in facts
        if str(fact.get("fact_id") or "").strip()
    }
    relation_records = []
    for relation in relations:
        if str(relation.get("relation_type") or "").strip().lower() != "before":
            continue
        source_id = str(relation.get("source_fact_id") or "").strip()
        target_id = str(relation.get("target_fact_id") or "").strip()
        source_fact = fact_by_id.get(source_id)
        target_fact = fact_by_id.get(target_id)
        if not source_fact or not target_fact:
            continue
        source_date = _format_timeline_date((source_fact.get("temporal_context") or {}).get("start_date") or relation.get("source_start_date"))
        target_date = _format_timeline_date((target_fact.get("temporal_context") or {}).get("start_date") or relation.get("target_start_date"))
        if not source_date or not target_date:
            continue
        relation_records.append(
            {
                "key": (source_id, target_id),
                "source_id": source_id,
                "target_id": target_id,
                "source_fact": source_fact,
                "target_fact": target_fact,
                "source_date": source_date,
                "target_date": target_date,
            }
        )
    if not relation_records:
        return []

    outgoing: Dict[str, List[Dict[str, Any]]] = {}
    incoming_count: Dict[str, int] = {}
    for record in relation_records:
        outgoing.setdefault(record["source_id"], []).append(record)
        incoming_count[record["target_id"]] = incoming_count.get(record["target_id"], 0) + 1
        incoming_count.setdefault(record["source_id"], incoming_count.get(record["source_id"], 0))

    lines: List[str] = []
    seen = set()
    used_keys = set()
    for record in relation_records:
        if len(lines) >= limit:
            break
        if record["key"] in used_keys:
            continue
        if incoming_count.get(record["source_id"], 0) != 0 or len(outgoing.get(record["source_id"], [])) != 1:
            continue
        chain = [record]
        next_id = record["target_id"]
        temp_used = {record["key"]}
        while len(outgoing.get(next_id, [])) == 1 and incoming_count.get(next_id, 0) == 1:
            next_record = outgoing[next_id][0]
            if next_record["key"] in temp_used:
                break
            chain.append(next_record)
            temp_used.add(next_record["key"])
            next_id = next_record["target_id"]
        if len(chain) < 2:
            continue
        segments = [
            f"{_chronology_fact_label(chain[0]['source_fact'])} on {chain[0]['source_date']}"
        ]
        segments.extend(
            f"{_chronology_fact_label(item['target_fact'])} on {item['target_date']}"
            for item in chain
        )
        line = f"{_join_chronology_segments(segments)} occurred in sequence."
        last_target = chain[-1]["target_fact"]
        target_context = last_target.get("temporal_context") if isinstance(last_target.get("temporal_context"), dict) else {}
        if target_context.get("derived_from_relative_anchor"):
            relative_markers = [str(item) for item in list(target_context.get("relative_markers") or []) if str(item)]
            if relative_markers:
                line = line.rstrip(".") + f". The later date is currently derived from reported timing ({relative_markers[0]})."
        key = line.lower()
        if key in seen:
            continue
        seen.add(key)
        used_keys.update(temp_used)
        lines.append(line)

    for record in relation_records:
        if len(lines) >= limit:
            break
        if record["key"] in used_keys:
            continue
        source_label = _chronology_fact_label(record["source_fact"])
        target_label = _chronology_fact_label(record["target_fact"])
        line = f"{source_label} on {record['source_date']} preceded {target_label.lower()} on {record['target_date']}."
        target_context = record["target_fact"].get("temporal_context") if isinstance(record["target_fact"].get("temporal_context"), dict) else {}
        if target_context.get("derived_from_relative_anchor"):
            relative_markers = [str(item) for item in list(target_context.get("relative_markers") or []) if str(item)]
            if relative_markers:
                line = line.rstrip(".") + f". The later date is currently derived from reported timing ({relative_markers[0]})."
        key = line.lower()
        if key in seen:
            continue
        seen.add(key)
        lines.append(line)
    return lines


def _safe_float(value: Any, default: float = 0.0) -> float:
    try:
        return float(value)
    except Exception:
        return float(default)


def _coerce_bool(value: Any, default: bool = False) -> bool:
    if isinstance(value, bool):
        return value
    if isinstance(value, (int, float)):
        return bool(value)
    text = str(value or "").strip().lower()
    if not text:
        return default
    if text in {"1", "true", "yes", "y", "on", "available", "enabled"}:
        return True
    if text in {"0", "false", "no", "n", "off", "unavailable", "disabled"}:
        return False
    return default


def _extract_text_candidates(value: Any) -> List[str]:
    if isinstance(value, str):
        return [value]
    if isinstance(value, list):
        results: List[str] = []
        for item in value:
            results.extend(_extract_text_candidates(item))
        return results
    if isinstance(value, dict):
        keys = (
            "fact",
            "fact_text",
            "text",
            "summary",
            "description",
            "name",
            "parsed_text_preview",
            "claim_element",
            "claim_element_text",
            "answer",
            "question",
            "title",
            "relevance",
        )
        results = []
        for key in keys:
            if key in value and value.get(key):
                results.extend(_extract_text_candidates(value.get(key)))
        return results
    return []


def _contains_date_anchor(value: Any) -> bool:
    text = str(value or "")
    return bool(
        re.search(
            r"\b(?:\d{1,2}/\d{1,2}/\d{2,4}|\d{4}-\d{2}-\d{2}|(?:jan|feb|mar|apr|may|jun|jul|aug|sep|sept|oct|nov|dec)[a-z]*\.?\s+\d{1,2}(?:,\s+\d{2,4})?)\b",
            text,
            flags=re.IGNORECASE,
        )
    )


def _contains_actor_marker(value: Any) -> bool:
    lowered = str(value or "").lower()
    return any(
        marker in lowered
        for marker in (
            "who at hacc",
            "caseworker",
            "housing specialist",
            "program manager",
            "hearing officer",
            "staff",
            "supervisor",
            "director",
            "coordinator",
            "name",
            "title",
        )
    )


def _contains_causation_marker(value: Any) -> bool:
    lowered = str(value or "").lower()
    if not lowered:
        return False
    return (
        any(marker in lowered for marker in ("because", "as a result", "after", "following", "in retaliation", "retaliat", "days after", "weeks after", "shortly after"))
        and any(marker in lowered for marker in ("complained", "reported", "grievance", "appeal", "protected activity", "requested accommodation"))
        and any(marker in lowered for marker in ("adverse action", "termination", "denial", "loss of assistance", "retaliat"))
    )


def _contains_hearing_timing_marker(value: Any) -> bool:
    lowered = str(value or "").lower()
    return any(
        marker in lowered
        for marker in (
            "hearing request",
            "requested a hearing",
            "requested review",
            "review request",
            "informal hearing request",
            "grievance request",
        )
    ) and any(marker in lowered for marker in ("date", "on ", "after", "before", "within", "days", "weeks"))


def _contains_response_date_marker(value: Any) -> bool:
    lowered = str(value or "").lower()
    return any(
        marker in lowered
        for marker in (
            "response date",
            "responded on",
            "response was",
            "review decision",
            "hearing outcome",
            "notice date",
            "decision date",
        )
    ) and _contains_date_anchor(value)


def _contains_staff_identity_marker(value: Any) -> bool:
    lowered = str(value or "").lower()
    return (
        "hacc" in lowered
        and any(marker in lowered for marker in ("name", "title", "staff", "caseworker", "manager", "officer", "specialist", "director"))
    )


def _contains_confirmation_placeholder(value: Any) -> bool:
    return bool(_CONFIRMATION_PLACEHOLDER_PATTERN.search(str(value or "")))


def _extract_latest_adversarial_priority_findings(value: Any) -> List[str]:
    findings: List[str] = []
    if isinstance(value, str):
        text = value.strip()
        if text:
            findings.append(text)
        return findings
    if isinstance(value, list):
        for item in value:
            findings.extend(_extract_latest_adversarial_priority_findings(item))
        return findings
    if isinstance(value, dict):
        candidate_keys = (
            "summary",
            "finding",
            "message",
            "reason",
            "title",
            "description",
            "priority",
            "text",
        )
        for key in candidate_keys:
            if key in value and value.get(key):
                findings.extend(_extract_latest_adversarial_priority_findings(value.get(key)))
        for nested_key in (
            "priorities",
            "priority_findings",
            "findings",
            "issues",
            "gaps",
            "latest_priorities",
            "latest_adversarial_priorities",
            "critical_findings",
            "latest_batch_findings",
            "latest_batch_priorities",
            "latest_adversarial_batch",
            "latest_adversarial_batch_summary",
        ):
            if nested_key in value and value.get(nested_key):
                findings.extend(_extract_latest_adversarial_priority_findings(value.get(nested_key)))
    return findings


def _has_chronology_gap_priority(findings: Iterable[str]) -> bool:
    lowered = " ".join(str(item or "").strip().lower() for item in findings if str(item or "").strip())
    if not lowered:
        return False
    chronology_markers = (
        "chronology",
        "exact date",
        "date gap",
        "response timing",
        "response date",
        "sequence",
        "timeline",
        "follow up",
        "follow-up",
        "critical chronology",
    )
    return any(marker in lowered for marker in chronology_markers)


def _has_decision_or_document_precision_priority(findings: Iterable[str]) -> bool:
    lowered = " ".join(str(item or "").strip().lower() for item in findings if str(item or "").strip())
    if not lowered:
        return False
    precision_markers = (
        "decision-maker",
        "decision maker",
        "adverse action",
        "documentary artifact",
        "document artifact",
        "documentary evidence",
        "specific decision",
        "who made",
        "who decided",
        "named actor",
        "artifact precision",
    )
    return any(marker in lowered for marker in precision_markers)


def _roman(index: int) -> str:
    numerals = [
        (1000, "M"),
        (900, "CM"),
        (500, "D"),
        (400, "CD"),
        (100, "C"),
        (90, "XC"),
        (50, "L"),
        (40, "XL"),
        (10, "X"),
        (9, "IX"),
        (5, "V"),
        (4, "IV"),
        (1, "I"),
    ]
    value = max(1, int(index))
    result = []
    for number, symbol in numerals:
        while value >= number:
            result.append(symbol)
            value -= number
    return "".join(result)


def _safe_call(target: Any, method_name: str, *args: Any, **kwargs: Any) -> Any:
    method = getattr(target, method_name, None)
    if not callable(method):
        return None
    try:
        return method(*args, **kwargs)
    except Exception:
        return None


def _merge_status(current: str, candidate: str) -> str:
    order = {
        "ready": 0,
        "warning": 1,
        "blocked": 2,
        "critical": 3,
    }
    current_status = str(current or "ready")
    candidate_status = str(candidate or "ready")
    return candidate_status if order.get(candidate_status, 0) > order.get(current_status, 0) else current_status


def _build_runtime_workflow_optimization_guidance(
    *,
    mediator: Any,
    drafting_readiness: Dict[str, Any],
    workflow_phase_plan: Dict[str, Any],
    document_optimization: Optional[Dict[str, Any]] = None,
) -> Dict[str, Any]:
    optimization_report = document_optimization if isinstance(document_optimization, dict) else {}
    existing_guidance = (
        optimization_report.get("workflow_optimization_guidance")
        if isinstance(optimization_report.get("workflow_optimization_guidance"), dict)
        else {}
    )
    if existing_guidance:
        guidance = dict(existing_guidance)
        if workflow_phase_plan and "workflow_phase_plan" not in guidance:
            guidance["workflow_phase_plan"] = dict(workflow_phase_plan)
        return guidance

    intake_case_summary = build_intake_case_review_summary(mediator)
    intake_sections = (
        intake_case_summary.get("intake_sections")
        if isinstance(intake_case_summary.get("intake_sections"), dict)
        else {}
    )
    question_summary = (
        intake_case_summary.get("question_candidate_summary")
        if isinstance(intake_case_summary.get("question_candidate_summary"), dict)
        else {}
    )
    claim_packet_summary = (
        intake_case_summary.get("claim_support_packet_summary")
        if isinstance(intake_case_summary.get("claim_support_packet_summary"), dict)
        else {}
    )
    candidate_claims = _coerce_list(intake_case_summary.get("candidate_claims"))
    claim_types = _unique_preserving_order(
        str((claim or {}).get("claim_type") or "").strip()
        for claim in candidate_claims
        if isinstance(claim, dict)
    )
    intake_focus_areas = _unique_preserving_order(
        str(section_name)
        for section_name, payload in intake_sections.items()
        if isinstance(payload, dict) and str(payload.get("status") or "").strip().lower() != "complete"
    )
    graph_focus_areas = _unique_preserving_order(
        [
            *claim_types,
            *[
                "claim_support_packets"
                if int(claim_packet_summary.get("unsupported_element_count") or 0) > 0
                else ""
            ],
        ]
    )
    document_focus_areas = _unique_preserving_order(
        [
            str(section_name)
            for section_name, payload in dict(drafting_readiness.get("sections") or {}).items()
            if isinstance(payload, dict) and str(payload.get("status") or "").strip().lower() != "ready"
        ]
    )
    cross_phase_findings = []
    if intake_focus_areas and graph_focus_areas:
        cross_phase_findings.append(
            "Intake follow-up gaps remain linked to graph support gaps, so unresolved intake sections should be closed before final drafting."
        )
    if graph_focus_areas and document_focus_areas:
        cross_phase_findings.append(
            "Graph support issues are still affecting drafting readiness, especially in claims-for-relief and chronology-dependent allegations."
        )
    return {
        "workflow_phase_plan": dict(workflow_phase_plan or {}),
        "phase_scorecards": {
            "intake_questioning": {
                "status": "warning" if intake_focus_areas else "ready",
                "focus_areas": intake_focus_areas,
                "question_candidate_count": int(question_summary.get("count") or 0),
            },
            "graph_analysis": {
                "status": "warning" if graph_focus_areas else "ready",
                "focus_areas": graph_focus_areas,
                "unsupported_element_count": int(claim_packet_summary.get("unsupported_element_count") or 0),
            },
            "document_generation": {
                "status": str(drafting_readiness.get("status") or "ready").strip().lower() or "ready",
                "focus_areas": document_focus_areas,
                "warning_count": int(drafting_readiness.get("warning_count") or 0),
            },
        },
        "cross_phase_findings": cross_phase_findings,
        "complaint_type_generalization_summary": {
            "complaint_types": claim_types,
            "complaint_type_count": len(claim_types),
        },
        "document_handoff_summary": {
            "ready_for_document_optimization": str(drafting_readiness.get("status") or "").strip().lower() == "ready",
            "drafting_status": str(drafting_readiness.get("status") or "ready").strip().lower() or "ready",
            "blocking_warning_count": int(drafting_readiness.get("warning_count") or 0),
        },
    }


class FormalComplaintDocumentBuilder:
    def __init__(self, mediator: Any):
        self.mediator = mediator

    def build_package(
        self,
        *,
        user_id: Optional[str] = None,
        court_name: str = "United States District Court",
        district: str = "",
        county: Optional[str] = None,
        division: Optional[str] = None,
        court_header_override: Optional[str] = None,
        case_number: Optional[str] = None,
        lead_case_number: Optional[str] = None,
        related_case_number: Optional[str] = None,
        assigned_judge: Optional[str] = None,
        courtroom: Optional[str] = None,
        title_override: Optional[str] = None,
        plaintiff_names: Optional[List[str]] = None,
        defendant_names: Optional[List[str]] = None,
        requested_relief: Optional[List[str]] = None,
        jury_demand: Optional[bool] = None,
        jury_demand_text: Optional[str] = None,
        signer_name: Optional[str] = None,
        signer_title: Optional[str] = None,
        signer_firm: Optional[str] = None,
        signer_bar_number: Optional[str] = None,
        signer_contact: Optional[str] = None,
        additional_signers: Optional[List[Dict[str, str]]] = None,
        declarant_name: Optional[str] = None,
        service_method: Optional[str] = None,
        service_recipients: Optional[List[str]] = None,
        service_recipient_details: Optional[List[Dict[str, str]]] = None,
        signature_date: Optional[str] = None,
        verification_date: Optional[str] = None,
        service_date: Optional[str] = None,
        affidavit_title: Optional[str] = None,
        affidavit_intro: Optional[str] = None,
        affidavit_facts: Optional[List[str]] = None,
        affidavit_supporting_exhibits: Optional[List[Dict[str, str]]] = None,
        affidavit_include_complaint_exhibits: Optional[bool] = None,
        affidavit_venue_lines: Optional[List[str]] = None,
        affidavit_jurat: Optional[str] = None,
        affidavit_notary_block: Optional[List[str]] = None,
        enable_agentic_optimization: bool = False,
        optimization_max_iterations: int = 2,
        optimization_target_score: float = 0.9,
        optimization_provider: Optional[str] = None,
        optimization_model_name: Optional[str] = None,
        optimization_llm_config: Optional[Dict[str, Any]] = None,
        optimization_persist_artifacts: bool = False,
        output_dir: Optional[str] = None,
        output_formats: Optional[List[str]] = None,
    ) -> Dict[str, Any]:
        resolved_user_id = self._resolve_user_id(user_id)
        formats = self._normalize_formats(output_formats)
        draft = self.build_draft(
            user_id=resolved_user_id,
            court_name=court_name,
            district=district,
            county=county,
            division=division,
            court_header_override=court_header_override,
            case_number=case_number,
            lead_case_number=lead_case_number,
            related_case_number=related_case_number,
            assigned_judge=assigned_judge,
            courtroom=courtroom,
            title_override=title_override,
            plaintiff_names=plaintiff_names,
            defendant_names=defendant_names,
            requested_relief=requested_relief,
            jury_demand=jury_demand,
            jury_demand_text=jury_demand_text,
            signer_name=signer_name,
            signer_title=signer_title,
            signer_firm=signer_firm,
            signer_bar_number=signer_bar_number,
            signer_contact=signer_contact,
            additional_signers=additional_signers,
            declarant_name=declarant_name,
            service_method=service_method,
            service_recipients=service_recipients,
            service_recipient_details=service_recipient_details,
            signature_date=signature_date,
            verification_date=verification_date,
            service_date=service_date,
            affidavit_title=affidavit_title,
            affidavit_intro=affidavit_intro,
            affidavit_facts=affidavit_facts,
            affidavit_supporting_exhibits=affidavit_supporting_exhibits,
            affidavit_include_complaint_exhibits=affidavit_include_complaint_exhibits,
            affidavit_venue_lines=affidavit_venue_lines,
            affidavit_jurat=affidavit_jurat,
            affidavit_notary_block=affidavit_notary_block,
        )
        document_optimization = None
        if enable_agentic_optimization:
            draft, document_optimization = self._optimize_draft(
                draft,
                user_id=resolved_user_id,
                max_iterations=optimization_max_iterations,
                target_score=optimization_target_score,
                provider=optimization_provider,
                model_name=optimization_model_name,
                llm_config=optimization_llm_config,
                persist_artifacts=optimization_persist_artifacts,
            )
        drafting_readiness = self._build_drafting_readiness(
            user_id=resolved_user_id,
            draft=draft,
        )
        intake_summary_handoff = self._build_intake_summary_handoff(document_optimization)
        if intake_summary_handoff:
            drafting_readiness["intake_summary_handoff"] = dict(intake_summary_handoff)
            handoff_payload = (
                dict(drafting_readiness.get("drafting_handoff") or {})
                if isinstance(drafting_readiness.get("drafting_handoff"), dict)
                else {}
            )
            handoff_payload["intake_summary_handoff_available"] = True
            handoff_payload["intake_summary_handoff_keys"] = sorted(
                str(key)
                for key in intake_summary_handoff.keys()
                if str(key).strip()
            )
            confirmation = (
                intake_summary_handoff.get("complainant_summary_confirmation")
                if isinstance(intake_summary_handoff.get("complainant_summary_confirmation"), dict)
                else {}
            )
            confirmation_snapshot = (
                confirmation.get("confirmed_summary_snapshot")
                if isinstance(confirmation.get("confirmed_summary_snapshot"), dict)
                else {}
            )
            priority_summary = (
                confirmation_snapshot.get("adversarial_intake_priority_summary")
                if isinstance(confirmation_snapshot.get("adversarial_intake_priority_summary"), dict)
                else {}
            )
            handoff_payload["uncovered_intake_objectives"] = _dedupe_text_values(
                priority_summary.get("uncovered_objectives") or []
            )
            drafting_readiness["drafting_handoff"] = handoff_payload
        workflow_phase_plan = self._build_runtime_workflow_phase_plan(
            drafting_readiness=drafting_readiness,
            document_optimization=document_optimization,
        )
        if not workflow_phase_plan:
            workflow_phase_plan = (
                dict(drafting_readiness.get("workflow_phase_plan") or {})
                if isinstance(drafting_readiness.get("workflow_phase_plan"), dict)
                else {}
            )
        if workflow_phase_plan:
            drafting_readiness["workflow_phase_plan"] = workflow_phase_plan
            self._refresh_drafting_readiness_workflow_warnings(
                drafting_readiness=drafting_readiness,
                workflow_phase_plan=workflow_phase_plan,
            )
        workflow_optimization_guidance = _build_runtime_workflow_optimization_guidance(
            mediator=self.mediator,
            drafting_readiness=drafting_readiness,
            workflow_phase_plan=workflow_phase_plan,
            document_optimization=document_optimization,
        )
        filing_checklist = self._build_filing_checklist(drafting_readiness)
        self._annotate_filing_checklist_review_links(
            filing_checklist=filing_checklist,
            drafting_readiness=drafting_readiness,
            user_id=resolved_user_id,
        )
        draft["drafting_readiness"] = drafting_readiness
        if workflow_phase_plan:
            draft["workflow_phase_plan"] = workflow_phase_plan
        if workflow_optimization_guidance:
            draft["workflow_optimization_guidance"] = workflow_optimization_guidance
        draft["filing_checklist"] = filing_checklist
        draft["affidavit"] = self._build_affidavit(draft)
        claim_support_temporal_handoff = self._build_claim_support_temporal_handoff(document_optimization)
        formalization_gate = self._build_formalization_gate_payload(drafting_readiness)
        source_context = draft.get("source_context") if isinstance(draft.get("source_context"), dict) else {}
        enriched_source_context = dict(source_context)
        if claim_support_temporal_handoff:
            enriched_source_context["claim_support_temporal_handoff"] = claim_support_temporal_handoff
        if formalization_gate:
            enriched_source_context["formalization_gate"] = formalization_gate
        if enriched_source_context:
            draft["source_context"] = enriched_source_context
        if formalization_gate:
            draft["formalization_gate"] = formalization_gate
        drafting_handoff = (
            dict(drafting_readiness.get("drafting_handoff") or {})
            if isinstance(drafting_readiness.get("drafting_handoff"), dict)
            else {}
        )
        if drafting_handoff:
            draft["drafting_handoff"] = drafting_handoff
        artifacts = self.render_artifacts(
            draft,
            output_dir=output_dir,
            output_formats=formats,
        )
        package_payload = {
            "draft": draft,
            "drafting_readiness": drafting_readiness,
            "filing_checklist": filing_checklist,
            "artifacts": artifacts,
            "document_optimization": document_optimization,
            "workflow_optimization_guidance": workflow_optimization_guidance,
            "intake_summary_handoff": intake_summary_handoff,
            "output_formats": formats,
            "generated_at": _utcnow().isoformat(),
        }
        if formalization_gate:
            package_payload["formalization_gate"] = formalization_gate
        if workflow_phase_plan:
            package_payload["workflow_phase_plan"] = workflow_phase_plan
        if drafting_handoff:
            package_payload["drafting_handoff"] = drafting_handoff
        if claim_support_temporal_handoff:
            package_payload["claim_support_temporal_handoff"] = claim_support_temporal_handoff
        return package_payload

    def build_draft(
        self,
        *,
        user_id: str,
        court_name: str,
        district: str,
        county: Optional[str],
        division: Optional[str],
        court_header_override: Optional[str],
        case_number: Optional[str],
        lead_case_number: Optional[str],
        related_case_number: Optional[str],
        assigned_judge: Optional[str],
        courtroom: Optional[str],
        title_override: Optional[str],
        plaintiff_names: Optional[List[str]],
        defendant_names: Optional[List[str]],
        requested_relief: Optional[List[str]],
        jury_demand: Optional[bool],
        jury_demand_text: Optional[str],
        signer_name: Optional[str],
        signer_title: Optional[str],
        signer_firm: Optional[str],
        signer_bar_number: Optional[str],
        signer_contact: Optional[str],
        additional_signers: Optional[List[Dict[str, str]]],
        declarant_name: Optional[str],
        service_method: Optional[str],
        service_recipients: Optional[List[str]],
        service_recipient_details: Optional[List[Dict[str, str]]],
        signature_date: Optional[str],
        verification_date: Optional[str],
        service_date: Optional[str],
        affidavit_title: Optional[str],
        affidavit_intro: Optional[str],
        affidavit_facts: Optional[List[str]],
        affidavit_supporting_exhibits: Optional[List[Dict[str, str]]],
        affidavit_include_complaint_exhibits: Optional[bool],
        affidavit_venue_lines: Optional[List[str]],
        affidavit_jurat: Optional[str],
        affidavit_notary_block: Optional[List[str]],
    ) -> Dict[str, Any]:
        affidavit_overrides = self._build_affidavit_overrides(
            affidavit_title=affidavit_title,
            affidavit_intro=affidavit_intro,
            affidavit_facts=affidavit_facts,
            affidavit_supporting_exhibits=affidavit_supporting_exhibits,
            affidavit_include_complaint_exhibits=affidavit_include_complaint_exhibits,
            affidavit_venue_lines=affidavit_venue_lines,
            affidavit_jurat=affidavit_jurat,
            affidavit_notary_block=affidavit_notary_block,
        )
        canonical_generate = getattr(self.mediator, "generate_formal_complaint", None)
        if callable(canonical_generate):
            try:
                result = canonical_generate(
                    user_id=user_id,
                    court_name=court_name,
                    district=district,
                    county=county,
                    division=division,
                    court_header_override=court_header_override,
                    case_number=case_number,
                    lead_case_number=lead_case_number,
                    related_case_number=related_case_number,
                    assigned_judge=assigned_judge,
                    courtroom=courtroom,
                    title_override=title_override,
                    plaintiff_names=plaintiff_names,
                    defendant_names=defendant_names,
                    requested_relief=requested_relief,
                    jury_demand=jury_demand,
                    jury_demand_text=jury_demand_text,
                    signer_name=signer_name,
                    signer_title=signer_title,
                    signer_firm=signer_firm,
                    signer_bar_number=signer_bar_number,
                    signer_contact=signer_contact,
                    additional_signers=additional_signers,
                    declarant_name=declarant_name,
                    service_method=service_method,
                    service_recipients=service_recipients,
                    service_recipient_details=service_recipient_details,
                    signature_date=signature_date,
                    verification_date=verification_date,
                    service_date=service_date,
                    affidavit_title=affidavit_title,
                    affidavit_intro=affidavit_intro,
                    affidavit_facts=affidavit_facts,
                    affidavit_supporting_exhibits=affidavit_supporting_exhibits,
                    affidavit_include_complaint_exhibits=affidavit_include_complaint_exhibits,
                    affidavit_venue_lines=affidavit_venue_lines,
                    affidavit_jurat=affidavit_jurat,
                    affidavit_notary_block=affidavit_notary_block,
                )
            except TypeError:
                result = None
            if isinstance(result, dict) and isinstance(result.get("formal_complaint"), dict):
                draft = self._adapt_formal_complaint_to_package_draft(result["formal_complaint"])
                draft["affidavit_overrides"] = affidavit_overrides
                draft["affidavit"] = self._build_affidavit(draft)
                draft["draft_text"] = self._render_draft_text(draft)
                return draft

        return self._build_legacy_draft(
            user_id=user_id,
            court_name=court_name,
            district=district,
            county=county,
            division=division,
            court_header_override=court_header_override,
            case_number=case_number,
            lead_case_number=lead_case_number,
            related_case_number=related_case_number,
            assigned_judge=assigned_judge,
            courtroom=courtroom,
            title_override=title_override,
            plaintiff_names=plaintiff_names,
            defendant_names=defendant_names,
            requested_relief=requested_relief,
            jury_demand=jury_demand,
            jury_demand_text=jury_demand_text,
            signer_name=signer_name,
            signer_title=signer_title,
            signer_firm=signer_firm,
            signer_bar_number=signer_bar_number,
            signer_contact=signer_contact,
            additional_signers=additional_signers,
            declarant_name=declarant_name,
            service_method=service_method,
            service_recipients=service_recipients,
            service_recipient_details=service_recipient_details,
            signature_date=signature_date,
            verification_date=verification_date,
            service_date=service_date,
            affidavit_title=affidavit_title,
            affidavit_intro=affidavit_intro,
            affidavit_facts=affidavit_facts,
            affidavit_supporting_exhibits=affidavit_supporting_exhibits,
            affidavit_include_complaint_exhibits=affidavit_include_complaint_exhibits,
            affidavit_venue_lines=affidavit_venue_lines,
            affidavit_jurat=affidavit_jurat,
            affidavit_notary_block=affidavit_notary_block,
        )

    def _build_legacy_draft(
        self,
        *,
        user_id: str,
        court_name: str,
        district: str,
        county: Optional[str],
        division: Optional[str],
        court_header_override: Optional[str],
        case_number: Optional[str],
        lead_case_number: Optional[str],
        related_case_number: Optional[str],
        assigned_judge: Optional[str],
        courtroom: Optional[str],
        title_override: Optional[str],
        plaintiff_names: Optional[List[str]],
        defendant_names: Optional[List[str]],
        requested_relief: Optional[List[str]],
        jury_demand: Optional[bool],
        jury_demand_text: Optional[str],
        signer_name: Optional[str],
        signer_title: Optional[str],
        signer_firm: Optional[str],
        signer_bar_number: Optional[str],
        signer_contact: Optional[str],
        additional_signers: Optional[List[Dict[str, str]]],
        declarant_name: Optional[str],
        service_method: Optional[str],
        service_recipients: Optional[List[str]],
        service_recipient_details: Optional[List[Dict[str, str]]],
        signature_date: Optional[str],
        verification_date: Optional[str],
        service_date: Optional[str],
        affidavit_title: Optional[str],
        affidavit_intro: Optional[str],
        affidavit_facts: Optional[List[str]],
        affidavit_supporting_exhibits: Optional[List[Dict[str, str]]],
        affidavit_include_complaint_exhibits: Optional[bool],
        affidavit_venue_lines: Optional[List[str]],
        affidavit_jurat: Optional[str],
        affidavit_notary_block: Optional[List[str]],
    ) -> Dict[str, Any]:
        state = getattr(self.mediator, "state", None)
        phase_manager = getattr(self.mediator, "phase_manager", None)
        intake_case_file = phase_manager.get_phase_data(ComplaintPhase.INTAKE, "intake_case_file") if phase_manager else None
        generated_complaint = self._get_existing_formal_complaint()
        classification = getattr(state, "legal_classification", {}) or {}
        statutes = _coerce_list(getattr(state, "applicable_statutes", []) or [])
        requirements = getattr(state, "summary_judgment_requirements", {}) or {}
        support_summary = _safe_call(self.mediator, "summarize_claim_support", user_id=user_id) or {}
        support_claims = support_summary.get("claims", {}) if isinstance(support_summary, dict) else {}
        claim_types = self._derive_claim_types(generated_complaint, classification, support_claims, requirements)
        plaintiffs, defendants = self._derive_parties(
            generated_complaint,
            plaintiff_names=plaintiff_names,
            defendant_names=defendant_names,
        )
        title = title_override or generated_complaint.get("title") or self._derive_title(plaintiffs, defendants)
        exhibits = self._collect_exhibits(user_id=user_id, claim_types=claim_types, support_claims=support_claims)
        facts = self._collect_general_facts(generated_complaint, classification, state)
        facts = self._annotate_lines_with_exhibits(facts, exhibits)
        claims_for_relief = self._build_claims_for_relief(
            user_id=user_id,
            claim_types=claim_types,
            requirements=requirements,
            statutes=statutes,
            support_claims=support_claims,
            exhibits=exhibits,
        )
        factual_allegations = self._build_factual_allegations(
            summary_of_facts=facts,
            claims_for_relief=claims_for_relief,
        )
        relief_items = _unique_preserving_order(
            list(requested_relief or [])
            + list(generated_complaint.get("prayer_for_relief", []) or [])
            + self._extract_requested_relief_from_facts(facts)
            + (STATE_DEFAULT_RELIEF if str(classification.get("jurisdiction") or "").strip().lower() == "state" else DEFAULT_RELIEF)
        )
        jury_demand_block = self._build_jury_demand(jury_demand=jury_demand, jury_demand_text=jury_demand_text)
        court_header = self._build_court_header(
            court_name=court_name,
            district=district,
            county=county,
            division=division,
            override=court_header_override,
        )
        jurisdiction_statement = self._build_jurisdiction_statement(
            classification=classification,
            statutes=statutes,
            court_name=court_name,
        )
        venue_statement = self._build_venue_statement(
            district=district,
            county=county,
            division=division,
            classification=classification,
            court_name=court_name,
        )
        nature_of_action = self._build_nature_of_action(
            claim_types=claim_types,
            classification=classification,
            statutes=statutes,
            court_name=court_name,
        )
        legal_standards = self._build_legal_standards_summary(statutes=statutes, requirements=requirements)
        signature_block = self._build_signature_block(
            plaintiffs,
            signer_name=signer_name,
            signer_title=signer_title,
            signer_firm=signer_firm,
            signer_bar_number=signer_bar_number,
            signer_contact=signer_contact,
            additional_signers=additional_signers,
            signature_date=signature_date,
        )
        verification = self._build_verification(
            plaintiffs,
            declarant_name=declarant_name,
            signer_name=signer_name,
            verification_date=verification_date,
            jurisdiction=classification.get("jurisdiction"),
        )
        certificate_of_service = self._build_certificate_of_service(
            plaintiffs,
            defendants,
            signer_name=signer_name,
            service_method=service_method,
            service_recipients=service_recipients,
            service_recipient_details=service_recipient_details,
            service_date=service_date,
            jurisdiction=classification.get("jurisdiction"),
        )

        draft = {
            "court_header": court_header,
            "case_caption": {
                "plaintiffs": plaintiffs,
                "defendants": defendants,
                "case_number": case_number or "________________",
                "county": county.strip().upper() if isinstance(county, str) and county.strip() else None,
                "lead_case_number": lead_case_number.strip() if isinstance(lead_case_number, str) and lead_case_number.strip() else None,
                "related_case_number": related_case_number.strip() if isinstance(related_case_number, str) and related_case_number.strip() else None,
                "assigned_judge": assigned_judge.strip() if isinstance(assigned_judge, str) and assigned_judge.strip() else None,
                "courtroom": courtroom.strip() if isinstance(courtroom, str) and courtroom.strip() else None,
                "jury_demand_notice": "JURY TRIAL DEMANDED" if jury_demand_block else None,
                "document_title": "COMPLAINT",
            },
            "title": title,
            "nature_of_action": nature_of_action,
            "parties": {
                "plaintiffs": plaintiffs,
                "defendants": defendants,
            },
            "jurisdiction_statement": jurisdiction_statement,
            "venue_statement": venue_statement,
            "factual_allegations": factual_allegations,
            "summary_of_facts": facts,
            "anchored_chronology_summary": _build_anchored_chronology_summary_from_case_file(
                intake_case_file if isinstance(intake_case_file, dict) else {}
            ),
            "claims_for_relief": claims_for_relief,
            "legal_standards": legal_standards,
            "requested_relief": relief_items,
            "jury_demand": jury_demand_block,
            "exhibits": exhibits,
            "signature_block": signature_block,
            "verification": verification,
            "certificate_of_service": certificate_of_service,
            "source_context": {
                "user_id": user_id,
                "claim_types": claim_types,
                "district": district,
                "jurisdiction": classification.get("jurisdiction", "unknown"),
                "generated_at": _utcnow().isoformat(),
            },
            "affidavit_overrides": self._build_affidavit_overrides(
                affidavit_title=affidavit_title,
                affidavit_intro=affidavit_intro,
                affidavit_facts=affidavit_facts,
                affidavit_supporting_exhibits=affidavit_supporting_exhibits,
                affidavit_include_complaint_exhibits=affidavit_include_complaint_exhibits,
                affidavit_venue_lines=affidavit_venue_lines,
                affidavit_jurat=affidavit_jurat,
                affidavit_notary_block=affidavit_notary_block,
            ),
        }
        self._annotate_claim_temporal_gap_hints(draft)
        self._attach_allegation_references(draft)
        self._annotate_case_caption_display(draft)
        draft["affidavit"] = self._build_affidavit(draft)
        draft["draft_text"] = self._render_draft_text(draft)
        return draft

    def _build_anchored_chronology_summary(self) -> List[str]:
        phase_manager = getattr(self.mediator, "phase_manager", None)
        intake_case_file = phase_manager.get_phase_data(ComplaintPhase.INTAKE, "intake_case_file") if phase_manager else None
        return _build_anchored_chronology_summary_from_case_file(
            intake_case_file if isinstance(intake_case_file, dict) else {}
        )

    def _optimize_draft(
        self,
        draft: Dict[str, Any],
        *,
        user_id: Optional[str],
        max_iterations: int,
        target_score: float,
        provider: Optional[str],
        model_name: Optional[str],
        llm_config: Optional[Dict[str, Any]],
        persist_artifacts: bool,
    ) -> tuple[Dict[str, Any], Dict[str, Any]]:
        optimizer = AgenticDocumentOptimizer(
            self.mediator,
            builder=self,
            provider=provider,
            model_name=model_name,
            max_iterations=max_iterations,
            target_score=target_score,
            persist_artifacts=persist_artifacts,
        )
        report = optimizer.optimize_draft(
            draft=draft,
            user_id=user_id,
            drafting_readiness={},
            config={
                "provider": provider,
                "model_name": model_name,
                "max_iterations": max_iterations,
                "target_score": target_score,
                "persist_artifacts": persist_artifacts,
                "llm_config": dict(llm_config or {}),
            },
        )
        optimized_draft = report.get("draft") or dict(draft)
        optimized_draft["summary_of_facts"] = self._normalize_text_lines(optimized_draft.get("summary_of_facts", []))
        optimized_draft["factual_allegations"] = self._expand_allegation_sources(
            optimized_draft.get("factual_allegations", []),
            limit=24,
        ) or self._expand_allegation_sources(draft.get("factual_allegations", []), limit=24)
        for claim in _coerce_list(optimized_draft.get("claims_for_relief")):
            if not isinstance(claim, dict):
                continue
            claim["supporting_facts"] = self._expand_allegation_sources(
                claim.get("supporting_facts", []),
                limit=10,
            ) or self._normalize_text_lines(claim.get("supporting_facts", []))
        self._annotate_claim_temporal_gap_hints(optimized_draft)
        self._attach_allegation_references(optimized_draft)
        self._annotate_case_caption_display(optimized_draft)
        optimized_draft["affidavit"] = self._build_affidavit(optimized_draft)
        optimized_draft["draft_text"] = self._render_draft_text(optimized_draft)
        return optimized_draft, report

    def _build_intake_summary_handoff(self, document_optimization: Optional[Dict[str, Any]]) -> Dict[str, Any]:
        optimization_report = document_optimization if isinstance(document_optimization, dict) else {}
        optimization_handoff = optimization_report.get("intake_summary_handoff")
        if isinstance(optimization_handoff, dict) and optimization_handoff:
            return dict(optimization_handoff)

        intake_status = build_intake_status_summary(self.mediator)
        status_handoff = intake_status.get("intake_summary_handoff") if isinstance(intake_status, dict) else None
        if isinstance(status_handoff, dict) and status_handoff:
            return dict(status_handoff)

        intake_case_summary = build_intake_case_review_summary(self.mediator)
        case_handoff = (
            intake_case_summary.get("intake_summary_handoff")
            if isinstance(intake_case_summary, dict)
            else None
        )
        if isinstance(case_handoff, dict) and case_handoff:
            return dict(case_handoff)

        return {}

    def _build_claim_support_temporal_handoff(self, document_optimization: Optional[Dict[str, Any]]) -> Dict[str, Any]:
        optimization_report = document_optimization if isinstance(document_optimization, dict) else {}
        optimization_handoff = optimization_report.get("claim_support_temporal_handoff")
        if isinstance(optimization_handoff, dict) and optimization_handoff:
            return dict(optimization_handoff)

        intake_case_summary = optimization_report.get("intake_case_summary")
        if not isinstance(intake_case_summary, dict) or not intake_case_summary:
            intake_case_summary = build_intake_case_review_summary(self.mediator)
        if not isinstance(intake_case_summary, dict) or not intake_case_summary:
            return {}

        packet_summary = intake_case_summary.get("claim_support_packet_summary")
        packet_summary = packet_summary if isinstance(packet_summary, dict) else {}
        alignment_tasks = intake_case_summary.get("alignment_evidence_tasks")
        alignment_tasks = alignment_tasks if isinstance(alignment_tasks, list) else []

        unresolved_temporal_issue_ids = _dedupe_text_values(
            packet_summary.get("claim_support_unresolved_temporal_issue_ids") or []
        )
        event_ids: List[str] = []
        temporal_fact_ids: List[str] = []
        temporal_relation_ids: List[str] = []
        timeline_issue_ids: List[str] = []
        temporal_issue_ids: List[str] = []
        temporal_proof_bundle_ids: List[str] = []
        temporal_proof_objectives: List[str] = []

        for task in alignment_tasks:
            if not isinstance(task, dict):
                continue
            event_ids.extend(_dedupe_text_values(task.get("event_ids") or []))
            temporal_fact_ids.extend(_dedupe_text_values(task.get("temporal_fact_ids") or []))
            temporal_relation_ids.extend(_dedupe_text_values(task.get("temporal_relation_ids") or []))
            timeline_issue_ids.extend(_dedupe_text_values(task.get("timeline_issue_ids") or []))
            temporal_issue_ids.extend(_dedupe_text_values(task.get("temporal_issue_ids") or []))
            proof_bundle_id = str(task.get("temporal_proof_bundle_id") or "").strip()
            if proof_bundle_id:
                temporal_proof_bundle_ids.append(proof_bundle_id)
            proof_objective = str(task.get("temporal_proof_objective") or "").strip()
            if proof_objective:
                temporal_proof_objectives.append(proof_objective)

        temporal_handoff = {
            "unresolved_temporal_issue_count": int(
                packet_summary.get("claim_support_unresolved_temporal_issue_count", 0) or 0
            ),
            "unresolved_temporal_issue_ids": unresolved_temporal_issue_ids,
            "chronology_task_count": int(packet_summary.get("temporal_gap_task_count", 0) or 0),
            "event_ids": _dedupe_text_values(event_ids),
            "temporal_fact_ids": _dedupe_text_values(temporal_fact_ids),
            "temporal_relation_ids": _dedupe_text_values(temporal_relation_ids),
            "timeline_issue_ids": _dedupe_text_values(timeline_issue_ids),
            "temporal_issue_ids": _dedupe_text_values(temporal_issue_ids),
            "temporal_proof_bundle_ids": _dedupe_text_values(temporal_proof_bundle_ids),
            "temporal_proof_objectives": _dedupe_text_values(temporal_proof_objectives),
        }
        if not temporal_handoff["unresolved_temporal_issue_count"] and not any(
            temporal_handoff[key]
            for key in (
                "unresolved_temporal_issue_ids",
                "event_ids",
                "temporal_fact_ids",
                "temporal_relation_ids",
                "timeline_issue_ids",
                "temporal_issue_ids",
                "temporal_proof_bundle_ids",
                "temporal_proof_objectives",
            )
        ):
            return {}
        return temporal_handoff

    def _adapt_formal_complaint_to_package_draft(self, formal_complaint: Dict[str, Any]) -> Dict[str, Any]:
        caption = formal_complaint.get("caption", {}) if isinstance(formal_complaint.get("caption"), dict) else {}
        claims_for_relief = []
        for claim in _coerce_list(formal_complaint.get("legal_claims")):
            if not isinstance(claim, dict):
                continue
            claims_for_relief.append(
                {
                    "claim_type": claim.get("claim_type") or claim.get("claim_name") or claim.get("title") or "Claim",
                    "count_title": claim.get("claim_name") or claim.get("title") or "Claim",
                    "legal_standards": _unique_preserving_order(
                        [claim.get("legal_standard", "")]
                        + [
                            f"{item.get('citation')} - {item.get('element')}"
                            if item.get("citation")
                            else str(item.get("element") or "")
                            for item in _coerce_list(claim.get("legal_standard_elements"))
                            if isinstance(item, dict) and (item.get("element") or item.get("citation"))
                        ]
                    ),
                    "supporting_facts": _unique_preserving_order(_extract_text_candidates(claim.get("supporting_facts"))),
                    "missing_elements": _unique_preserving_order(
                        _extract_text_candidates(claim.get("missing_requirements"))
                    ),
                    "partially_supported_elements": [],
                    "support_summary": {
                        "elements_satisfied": claim.get("elements_satisfied", ""),
                        "authority_count": len(_coerce_list(claim.get("supporting_authorities"))),
                    },
                    "supporting_exhibits": [
                        {
                            "label": exhibit.get("label"),
                            "title": exhibit.get("title"),
                            "link": exhibit.get("reference") or exhibit.get("source_url") or exhibit.get("link"),
                        }
                        for exhibit in _coerce_list(claim.get("supporting_exhibits"))
                        if isinstance(exhibit, dict)
                    ],
                }
            )

        legal_standards = []
        for standard in _coerce_list(formal_complaint.get("legal_standards")):
            if isinstance(standard, dict):
                claim_name = str(standard.get("claim_name") or standard.get("claim_type") or "").strip()
                body = str(standard.get("standard") or "").strip()
                citations = ", ".join(_unique_preserving_order(_extract_text_candidates(standard.get("citations"))))
                if claim_name and body and citations:
                    legal_standards.append(f"{claim_name}: {body} ({citations})")
                elif claim_name and body:
                    legal_standards.append(f"{claim_name}: {body}")
                elif body:
                    legal_standards.append(body)
            else:
                text = str(standard or "").strip()
                if text:
                    legal_standards.append(text)

        exhibits = []
        for exhibit in _coerce_list(formal_complaint.get("exhibits")):
            if not isinstance(exhibit, dict):
                continue
            exhibits.append(
                {
                    "label": exhibit.get("label"),
                    "title": exhibit.get("title") or exhibit.get("description") or "Supporting exhibit",
                    "claim_type": exhibit.get("claim_type"),
                    "kind": exhibit.get("kind") or "evidence",
                    "link": exhibit.get("reference") or exhibit.get("source_url") or exhibit.get("link") or "",
                    "source_ref": exhibit.get("cid") or exhibit.get("reference") or "",
                    "summary": exhibit.get("summary") or exhibit.get("description") or "",
                }
            )

        nature_of_action = formal_complaint.get("nature_of_action")
        if isinstance(nature_of_action, str):
            nature_of_action = [nature_of_action]

        factual_allegations = _unique_preserving_order(
            _extract_text_candidates(formal_complaint.get("factual_allegations") or formal_complaint.get("summary_of_facts"))
        )
        if not factual_allegations:
            factual_allegations = self._build_factual_allegations(
                summary_of_facts=_extract_text_candidates(formal_complaint.get("summary_of_facts")),
                claims_for_relief=claims_for_relief,
            )

        draft = {
            "court_header": formal_complaint.get("court_header", ""),
            "case_caption": {
                "plaintiffs": _coerce_list(formal_complaint.get("parties", {}).get("plaintiffs", [])) if isinstance(formal_complaint.get("parties"), dict) else [],
                "defendants": _coerce_list(formal_complaint.get("parties", {}).get("defendants", [])) if isinstance(formal_complaint.get("parties"), dict) else [],
                "case_number": caption.get("case_number") or formal_complaint.get("case_number") or "________________",
                "county": caption.get("county_line") or ((formal_complaint.get("caption") or {}).get("county_line") if isinstance(formal_complaint.get("caption"), dict) else ""),
                "lead_case_number": caption.get("lead_case_number") or ((formal_complaint.get("caption") or {}).get("lead_case_number") if isinstance(formal_complaint.get("caption"), dict) else ""),
                "related_case_number": caption.get("related_case_number") or ((formal_complaint.get("caption") or {}).get("related_case_number") if isinstance(formal_complaint.get("caption"), dict) else ""),
                "assigned_judge": caption.get("assigned_judge") or ((formal_complaint.get("caption") or {}).get("assigned_judge") if isinstance(formal_complaint.get("caption"), dict) else ""),
                "courtroom": caption.get("courtroom") or ((formal_complaint.get("caption") or {}).get("courtroom") if isinstance(formal_complaint.get("caption"), dict) else ""),
                "jury_demand_notice": caption.get("jury_demand_notice") or ((formal_complaint.get("caption") or {}).get("jury_demand_notice") if isinstance(formal_complaint.get("caption"), dict) else ""),
                "document_title": "COMPLAINT",
            },
            "title": formal_complaint.get("title") or caption.get("case_title") or "Complaint",
            "nature_of_action": _unique_preserving_order(_extract_text_candidates(nature_of_action)),
            "parties": formal_complaint.get("parties", {}),
            "jurisdiction_statement": formal_complaint.get("jurisdiction_statement", ""),
            "venue_statement": formal_complaint.get("venue_statement", ""),
            "factual_allegations": factual_allegations,
            "summary_of_facts": _unique_preserving_order(_extract_text_candidates(formal_complaint.get("summary_of_facts") or formal_complaint.get("factual_allegations"))),
            "anchored_chronology_summary": _unique_preserving_order(
                _extract_text_candidates(formal_complaint.get("anchored_chronology_summary"))
            ),
            "claims_for_relief": claims_for_relief,
            "legal_standards": _unique_preserving_order(legal_standards),
            "requested_relief": _unique_preserving_order(_extract_text_candidates(formal_complaint.get("requested_relief") or formal_complaint.get("prayer_for_relief"))),
            "jury_demand": formal_complaint.get("jury_demand", {}),
            "exhibits": exhibits,
            "signature_block": formal_complaint.get("signature_block", {}),
            "verification": formal_complaint.get("verification", {}),
            "certificate_of_service": formal_complaint.get("certificate_of_service", {}),
            "source_context": {
                "generated_at": formal_complaint.get("generated_at") or _utcnow().isoformat(),
                "district": formal_complaint.get("district") or caption.get("district") or "",
                "jurisdiction": formal_complaint.get("jurisdiction", "unknown"),
            },
        }
        self._annotate_claim_temporal_gap_hints(draft)
        self._attach_allegation_references(draft)
        self._annotate_case_caption_display(draft)
        built_affidavit = self._build_affidavit(draft)
        existing_affidavit = formal_complaint.get("affidavit", {}) if isinstance(formal_complaint.get("affidavit"), dict) else {}
        draft["affidavit"] = {**built_affidavit, **existing_affidavit}
        rendered_draft_text = self._render_draft_text(draft)
        supplied_draft_text = str(formal_complaint.get("draft_text") or "").strip()
        expected_case_line = (
            f"{draft['case_caption'].get('case_number_label', 'Civil Action No.')} "
            f"{draft['case_caption'].get('case_number', '________________')}"
        )
        draft["draft_text"] = (
            supplied_draft_text
            if supplied_draft_text and expected_case_line in supplied_draft_text
            else rendered_draft_text
        )
        return draft

    def _format_county_for_header(self, county: Optional[str]) -> str:
        county_text = str(county or "").strip().upper()
        if not county_text:
            return ""
        if county_text.startswith("COUNTY OF "):
            return county_text
        if county_text.endswith(" COUNTY"):
            return f"COUNTY OF {county_text[:-7].strip()}"
        return f"COUNTY OF {county_text}"

    def _annotate_case_caption_display(self, draft: Dict[str, Any]) -> None:
        caption = draft.get("case_caption")
        if not isinstance(caption, dict):
            return
        source_context = draft.get("source_context", {}) if isinstance(draft.get("source_context"), dict) else {}
        jurisdiction = str(source_context.get("jurisdiction") or "").strip()
        forum_type = self._infer_forum_type(
            classification={"jurisdiction": jurisdiction},
            court_name=str(draft.get("court_header") or ""),
        )
        caption["forum_type"] = forum_type
        caption["case_number_label"] = caption.get("case_number_label") or (
            "Case No." if forum_type == "state" else "Civil Action No."
        )
        caption["lead_case_number_label"] = caption.get("lead_case_number_label") or (
            "Related Proceeding No." if forum_type == "state" else "Lead Case No."
        )
        caption["related_case_number_label"] = caption.get("related_case_number_label") or (
            "Coordination No." if forum_type == "state" else "Related Case No."
        )
        caption["assigned_judge_label"] = caption.get("assigned_judge_label") or (
            "Judicial Officer" if forum_type == "state" else "Assigned Judge"
        )
        caption["courtroom_label"] = caption.get("courtroom_label") or (
            "Department" if forum_type == "state" else "Courtroom"
        )
        plaintiff_names = caption.get("plaintiffs") if isinstance(caption.get("plaintiffs"), list) else []
        defendant_names = caption.get("defendants") if isinstance(caption.get("defendants"), list) else []
        caption["plaintiff_caption_label"] = caption.get("plaintiff_caption_label") or (
            "Plaintiff" if len(plaintiff_names) == 1 else "Plaintiffs"
        )
        caption["defendant_caption_label"] = caption.get("defendant_caption_label") or (
            "Defendant" if len(defendant_names) == 1 else "Defendants"
        )
        caption["caption_party_lines"] = caption.get("caption_party_lines") or self._build_caption_party_lines(caption)

    def _build_caption_party_lines(self, caption: Dict[str, Any]) -> List[str]:
        plaintiffs = caption.get("plaintiffs") if isinstance(caption.get("plaintiffs"), list) else []
        defendants = caption.get("defendants") if isinstance(caption.get("defendants"), list) else []
        plaintiff_names = [str(name).strip() for name in plaintiffs if str(name).strip()] or ["Plaintiff"]
        defendant_names = [str(name).strip() for name in defendants if str(name).strip()] or ["Defendant"]
        plaintiff_label = str(
            caption.get("plaintiff_caption_label")
            or ("Plaintiff" if len(plaintiff_names) == 1 else "Plaintiffs")
        ).strip()
        defendant_label = str(
            caption.get("defendant_caption_label")
            or ("Defendant" if len(defendant_names) == 1 else "Defendants")
        ).strip()
        return [
            f"{'\n'.join(plaintiff_names)}, {plaintiff_label},",
            "v.",
            f"{'\n'.join(defendant_names)}, {defendant_label}.",
        ]

    def _resolve_draft_forum_type(self, draft: Dict[str, Any]) -> str:
        caption = draft.get("case_caption", {}) if isinstance(draft.get("case_caption"), dict) else {}
        forum_type = str(caption.get("forum_type") or "").strip().lower()
        if forum_type:
            return forum_type
        source_context = draft.get("source_context", {}) if isinstance(draft.get("source_context"), dict) else {}
        return self._infer_forum_type(
            classification={"jurisdiction": source_context.get("jurisdiction")},
            court_name=str(draft.get("court_header") or ""),
        )

    def _build_party_section_lines(
        self,
        *,
        plaintiffs: List[str],
        defendants: List[str],
        forum_type: str,
    ) -> List[str]:
        plaintiff_names = [str(name).strip() for name in _coerce_list(plaintiffs) if str(name).strip()] or ["Plaintiff"]
        defendant_names = [str(name).strip() for name in _coerce_list(defendants) if str(name).strip()] or ["Defendant"]
        plaintiff_label = "Plaintiff" if len(plaintiff_names) == 1 else "Plaintiffs"
        defendant_label = "Defendant" if len(defendant_names) == 1 else "Defendants"
        plaintiff_names_text = ", ".join(plaintiff_names)
        defendant_names_text = ", ".join(defendant_names)
        if forum_type == "state":
            plaintiff_verb = "is" if len(plaintiff_names) == 1 else "are"
            defendant_verb = "is" if len(defendant_names) == 1 else "are"
            return [
                f"{plaintiff_label} {plaintiff_names_text} {plaintiff_verb} a party bringing this civil action in this Court.",
                f"{defendant_label} {defendant_names_text} {defendant_verb} named as the party from whom relief is sought.",
            ]
        return [
            f"{plaintiff_label}: {plaintiff_names_text}.",
            f"{defendant_label}: {defendant_names_text}.",
        ]

    def _build_jurisdiction_statement(
        self,
        *,
        classification: Dict[str, Any],
        statutes: List[Dict[str, Any]],
        court_name: str,
    ) -> str:
        forum_type = self._infer_forum_type(classification=classification, court_name=court_name)
        first_citation = next(
            (
                str(statute.get("citation") or "").strip()
                for statute in statutes
                if isinstance(statute, dict) and statute.get("citation")
            ),
            "",
        )
        if forum_type == "federal":
            if first_citation:
                return (
                    "This Court has subject-matter jurisdiction under federal law, including "
                    f"{first_citation}, because Plaintiff alleges violations arising under the laws of the United States."
                )
            return "This Court has subject-matter jurisdiction under 28 U.S.C. § 1331 because Plaintiff alleges claims arising under federal law."
        if forum_type == "state":
            if first_citation:
                return (
                    "This Court has subject-matter jurisdiction because Plaintiff asserts claims arising under "
                    f"the governing state law, including {first_citation}, and seeks relief within this Court's authority."
                )
            return (
                "This Court has subject-matter jurisdiction because Plaintiff asserts claims arising under the "
                "governing state law and seeks relief within this Court's authority."
            )
        return "This Court has subject-matter jurisdiction because the claims arise under the governing law identified in this pleading."

    def _build_venue_statement(
        self,
        *,
        district: str,
        county: Optional[str],
        division: Optional[str],
        classification: Dict[str, Any],
        court_name: str,
    ) -> str:
        district_text = str(district or "").strip()
        county_text = str(county or "").strip()
        division_text = str(division or "").strip()
        forum_type = self._infer_forum_type(classification=classification, court_name=court_name)
        if forum_type == "state" and county_text:
            return (
                "Venue is proper in this Court because a substantial part of the events or omissions giving rise "
                f"to these claims occurred in {county_text}."
            )
        if forum_type == "federal" and district_text and division_text:
            return (
                f"Venue is proper in the {division_text} Division of the {district_text} because a substantial part of the events or omissions giving rise to these claims occurred there."
            )
        if forum_type == "federal" and district_text:
            return (
                f"Venue is proper in the {district_text} because a substantial part of the events or omissions giving rise to these claims occurred there."
            )
        if forum_type == "state" and district_text and division_text:
            return (
                "Venue is proper in this Court because a substantial part of the events or omissions giving rise "
                f"to these claims occurred in {division_text}, {district_text}."
            )
        if forum_type == "state" and district_text:
            return (
                "Venue is proper in this Court because a substantial part of the events or omissions giving rise "
                f"to these claims occurred in {district_text}."
            )
        return "Venue is proper in this Court because a substantial part of the events or omissions giving rise to these claims occurred in this judicial district."

    def _render_draft_text(self, draft: Dict[str, Any]) -> str:
        caption = draft.get("case_caption", {}) if isinstance(draft.get("case_caption"), dict) else {}
        parties = draft.get("parties", {}) if isinstance(draft.get("parties"), dict) else {}
        signature_block = draft.get("signature_block", {}) if isinstance(draft.get("signature_block"), dict) else {}
        plaintiff_list = parties.get("plaintiffs", []) or caption.get("plaintiffs", []) or ["Plaintiff"]
        defendant_list = parties.get("defendants", []) or caption.get("defendants", []) or ["Defendant"]
        plaintiffs = ", ".join(plaintiff_list)
        defendants = ", ".join(defendant_list)
        forum_type = self._resolve_draft_forum_type(draft)
        caption_party_lines = caption.get("caption_party_lines") if isinstance(caption.get("caption_party_lines"), list) else self._build_caption_party_lines(caption)
        party_section_lines = self._build_party_section_lines(
            plaintiffs=plaintiff_list,
            defendants=defendant_list,
            forum_type=forum_type,
        )
        case_number_label = str(caption.get("case_number_label") or "Civil Action No.")
        lead_case_number_label = str(caption.get("lead_case_number_label") or "Lead Case No.")
        related_case_number_label = str(caption.get("related_case_number_label") or "Related Case No.")
        assigned_judge_label = str(caption.get("assigned_judge_label") or "Assigned Judge")
        courtroom_label = str(caption.get("courtroom_label") or "Courtroom")
        lines = [
            str(draft.get("court_header") or "IN THE COURT OF COMPETENT JURISDICTION"),
            *([str(caption.get("county"))] if caption.get("county") else []),
            "",
            *caption_party_lines,
            f"{case_number_label} {caption.get('case_number', '________________')}",
            *([f"{lead_case_number_label} {caption.get('lead_case_number')}"] if caption.get('lead_case_number') else []),
            *([f"{related_case_number_label} {caption.get('related_case_number')}"] if caption.get('related_case_number') else []),
            *([f"{assigned_judge_label}: {caption.get('assigned_judge')}"] if caption.get('assigned_judge') else []),
            *([f"{courtroom_label}: {caption.get('courtroom')}"] if caption.get('courtroom') else []),
            "",
            str(caption.get("document_title") or "COMPLAINT"),
            *([str(caption.get("jury_demand_notice"))] if caption.get("jury_demand_notice") else []),
            "",
            "NATURE OF THE ACTION",
        ]
        lines.extend(self._normalize_text_lines(draft.get("nature_of_action", [])))
        lines.extend([
            "",
            "PARTIES",
            *party_section_lines,
            "",
            "JURISDICTION AND VENUE",
        ])
        if draft.get("jurisdiction_statement"):
            lines.append(str(draft["jurisdiction_statement"]))
        if draft.get("venue_statement"):
            lines.append(str(draft["venue_statement"]))
        lines.extend(["", "FACTUAL ALLEGATIONS"])
        lines.extend(self._grouped_allegation_text_lines(draft))
        chronology_lines = self._normalize_text_lines(draft.get("anchored_chronology_summary", []))
        if chronology_lines:
            lines.extend(["", "ANCHORED CHRONOLOGY"])
            lines.extend(self._numbered_lines(chronology_lines))
        claims = draft.get("claims_for_relief", []) if isinstance(draft.get("claims_for_relief"), list) else []
        if claims:
            lines.extend(["", "CLAIMS FOR RELIEF"])
        for index, claim in enumerate(claims, start=1):
            lines.extend([
                "",
                f"COUNT {_roman(index)} - {claim.get('count_title', claim.get('claim_type', 'Claim'))}",
                "Legal Standard:",
            ])
            lines.extend(self._bulletize_lines(claim.get("legal_standards", [])))
            incorporated_clause = self._format_incorporated_reference_clause(
                claim.get("allegation_references", []),
                claim.get("supporting_exhibits", []),
            )
            if incorporated_clause:
                lines.append(incorporated_clause)
            lines.append("Claim-Specific Support:")
            lines.extend(self._bulletize_lines(claim.get("supporting_facts", [])))
            missing = self._normalize_text_lines(claim.get("missing_elements", []))
            if missing:
                lines.append("Open Support Gaps:")
                lines.extend([f"- {line}" for line in missing])
        lines.extend(["", "REQUESTED RELIEF"])
        if forum_type == "state":
            lines.append("Wherefore, Plaintiff prays for judgment against Defendant as follows:")
        lines.extend(self._numbered_lines(draft.get("requested_relief", [])))
        jury_demand = draft.get("jury_demand", {}) if isinstance(draft.get("jury_demand"), dict) else {}
        if jury_demand:
            lines.extend(["", str(jury_demand.get("title") or "JURY DEMAND").upper()])
            if jury_demand.get("text"):
                lines.append(str(jury_demand.get("text")))
        exhibits = draft.get("exhibits", []) if isinstance(draft.get("exhibits"), list) else []
        if exhibits:
            lines.extend(["", "EXHIBITS"])
            for exhibit in exhibits:
                if not isinstance(exhibit, dict):
                    continue
                text = f"{exhibit.get('label', 'Exhibit')} - {exhibit.get('title', 'Supporting exhibit')}"
                if exhibit.get("link"):
                    text = f"{text} ({exhibit['link']})"
                lines.append(text)
                if exhibit.get("summary"):
                    lines.append(f"  {exhibit['summary']}")
        verification = draft.get("verification", {}) if isinstance(draft.get("verification"), dict) else {}
        if verification:
            lines.extend([
                "",
                str(verification.get("title") or "Verification").upper(),
                str(verification.get("text") or ""),
                str(verification.get("dated") or ""),
                str(verification.get("signature_line") or ""),
            ])
        certificate_of_service = draft.get("certificate_of_service", {}) if isinstance(draft.get("certificate_of_service"), dict) else {}
        if certificate_of_service:
            lines.extend([
                "",
                str(certificate_of_service.get("title") or "Certificate of Service").upper(),
                str(certificate_of_service.get("text") or ""),
                str(certificate_of_service.get("dated") or ""),
                str(certificate_of_service.get("signature_line") or ""),
            ])
        affidavit = draft.get("affidavit", {}) if isinstance(draft.get("affidavit"), dict) else {}
        if affidavit:
            lines.extend([
                "",
                str(affidavit.get("title") or "AFFIDAVIT IN SUPPORT OF COMPLAINT"),
            ])
            lines.extend(str(line) for line in _coerce_list(affidavit.get("venue_lines")) if str(line or "").strip())
            lines.extend([
                "",
                str(affidavit.get("intro") or ""),
                str(affidavit.get("knowledge_graph_note") or ""),
                "",
                "Affiant states as follows:",
            ])
            lines.extend(self._numbered_lines(affidavit.get("facts", [])))
            supporting_exhibits = affidavit.get("supporting_exhibits") if isinstance(affidavit.get("supporting_exhibits"), list) else []
            if supporting_exhibits:
                lines.extend(["", "AFFIDAVIT SUPPORTING EXHIBITS"])
                for exhibit in supporting_exhibits:
                    if not isinstance(exhibit, dict):
                        continue
                    exhibit_text = f"{exhibit.get('label', 'Exhibit')} - {exhibit.get('title', 'Supporting exhibit')}"
                    if exhibit.get("link"):
                        exhibit_text = f"{exhibit_text} ({exhibit['link']})"
                    lines.append(exhibit_text)
            lines.extend([
                "",
                str(affidavit.get("dated") or ""),
                str(affidavit.get("signature_line") or ""),
                str(affidavit.get("jurat") or ""),
            ])
            lines.extend(str(line) for line in _coerce_list(affidavit.get("notary_block")) if str(line or "").strip())
        lines.extend(["", *self._build_signature_section_lines(signature_block, forum_type)])
        return "\n".join(line for line in lines if line is not None)

    def _normalize_text_lines(self, values: Any) -> List[str]:
        normalized = []
        for value in _unique_preserving_order(_extract_text_candidates(values)):
            text = re.sub(r"\s+", " ", value).strip()
            if text:
                normalized.append(text)
        return normalized

    def _split_allegation_fragments(self, value: Any) -> List[str]:
        text = re.sub(r"\s+", " ", str(value or "")).strip(" -;")
        if not text:
            return []
        if ": " in text:
            prefix, suffix = text.split(": ", 1)
            prefix_lower = prefix.strip().lower()
            if (
                prefix.strip().endswith("?")
                or prefix_lower.startswith(("what ", "when ", "where ", "why ", "how ", "who ", "describe ", "explain "))
                or prefix_lower in {"what happened", "what relief do you want"}
            ):
                text = suffix.strip()
        parts = [
            part.strip(" -;")
            for part in re.split(r"(?<=[.!?])\s+(?=[A-Z0-9])", text)
            if part.strip(" -;")
        ]
        return parts or [text]

    def _formalize_allegation_fragment(self, value: Any) -> str:
        text = re.sub(r"\s+", " ", str(value or "")).strip(" -;")
        if not text:
            return ""
        replacements = (
            (r"^i was\b", "Plaintiff was"),
            (r"^i am\b", "Plaintiff is"),
            (r"^i need\b", "Plaintiff needs"),
            (r"^i needed\b", "Plaintiff needed"),
            (r"^i lost\b", "Plaintiff lost"),
            (r"^i asked\b", "Plaintiff asked"),
            (r"^i reported\b", "Plaintiff reported"),
            (r"^i complained\b", "Plaintiff complained"),
            (r"^i informed\b", "Plaintiff informed"),
            (r"^i notified\b", "Plaintiff notified"),
            (r"^i requested\b", "Plaintiff requested"),
            (r"^i sought\b", "Plaintiff sought"),
            (r"^i experienced\b", "Plaintiff experienced"),
            (r"^i suffered\b", "Plaintiff suffered"),
            (r"^i told\b", "Plaintiff told"),
            (r"^they\b", "Defendant"),
        )
        for pattern, replacement in replacements:
            text = re.sub(pattern, replacement, text, flags=re.IGNORECASE)
        clause_replacements = (
            (r"([,;]\s+)i was\b", r"\1Plaintiff was"),
            (r"([,;]\s+)i am\b", r"\1Plaintiff is"),
            (r"([,;]\s+)i need\b", r"\1Plaintiff needs"),
            (r"([,;]\s+)i needed\b", r"\1Plaintiff needed"),
            (r"([,;]\s+)i lost\b", r"\1Plaintiff lost"),
            (r"([,;]\s+)i asked\b", r"\1Plaintiff asked"),
            (r"([,;]\s+)i reported\b", r"\1Plaintiff reported"),
            (r"([,;]\s+)i complained\b", r"\1Plaintiff complained"),
            (r"([,;]\s+)i requested\b", r"\1Plaintiff requested"),
            (r"([,;]\s+)i informed\b", r"\1Plaintiff informed"),
            (r"([,;]\s+)i notified\b", r"\1Plaintiff notified"),
            (r"([,;]\s+)i suffered\b", r"\1Plaintiff suffered"),
            (r"([,;]\s+)i experienced\b", r"\1Plaintiff experienced"),
            (r"([,;]\s+)i told\b", r"\1Plaintiff told"),
            (r"(\band\s+)i was\b", r"\1Plaintiff was"),
            (r"(\band\s+)i am\b", r"\1Plaintiff is"),
            (r"(\band\s+)i need\b", r"\1Plaintiff needs"),
            (r"(\band\s+)i needed\b", r"\1Plaintiff needed"),
            (r"(\band\s+)i lost\b", r"\1Plaintiff lost"),
            (r"(\band\s+)i asked\b", r"\1Plaintiff asked"),
            (r"(\band\s+)i reported\b", r"\1Plaintiff reported"),
            (r"(\band\s+)i complained\b", r"\1Plaintiff complained"),
            (r"(\band\s+)i requested\b", r"\1Plaintiff requested"),
            (r"(\band\s+)i informed\b", r"\1Plaintiff informed"),
            (r"(\band\s+)i notified\b", r"\1Plaintiff notified"),
            (r"(\band\s+)i suffered\b", r"\1Plaintiff suffered"),
            (r"(\band\s+)i experienced\b", r"\1Plaintiff experienced"),
            (r"(\band\s+)i told\b", r"\1Plaintiff told"),
            (r"(\bafter\s+)i was\b", r"\1Plaintiff was"),
            (r"(\bafter\s+)i am\b", r"\1Plaintiff is"),
            (r"(\bafter\s+)i need\b", r"\1Plaintiff needs"),
            (r"(\bafter\s+)i needed\b", r"\1Plaintiff needed"),
            (r"(\bafter\s+)i lost\b", r"\1Plaintiff lost"),
            (r"(\bafter\s+)i asked\b", r"\1Plaintiff asked"),
            (r"(\bafter\s+)i reported\b", r"\1Plaintiff reported"),
            (r"(\bafter\s+)i complained\b", r"\1Plaintiff complained"),
            (r"(\bafter\s+)i requested\b", r"\1Plaintiff requested"),
            (r"(\bafter\s+)i informed\b", r"\1Plaintiff informed"),
            (r"(\bafter\s+)i notified\b", r"\1Plaintiff notified"),
            (r"(\bafter\s+)i suffered\b", r"\1Plaintiff suffered"),
            (r"(\bafter\s+)i experienced\b", r"\1Plaintiff experienced"),
            (r"(\bafter\s+)i told\b", r"\1Plaintiff told"),
            (r"(\bthat\s+)i am\b", r"\1Plaintiff is"),
            (r"(\bthat\s+)i need\b", r"\1Plaintiff needs"),
            (r"(\bthat\s+)i needed\b", r"\1Plaintiff needed"),
            (r"(\bthat\s+)i asked\b", r"\1Plaintiff asked"),
            (r"(\bthat\s+)i complained\b", r"\1Plaintiff complained"),
            (r"(\bthat\s+)i requested\b", r"\1Plaintiff requested"),
            (r"(\bthat\s+)i told\b", r"\1Plaintiff told"),
        )
        for pattern, replacement in clause_replacements:
            text = re.sub(pattern, replacement, text, flags=re.IGNORECASE)
        text = re.sub(r"\bmy\b", "Plaintiff's", text, flags=re.IGNORECASE)
        text = re.sub(r"\bmine\b", "Plaintiff's", text, flags=re.IGNORECASE)
        text = re.sub(r"\bme\b", "Plaintiff", text, flags=re.IGNORECASE)
        text = re.sub(r"\blost Plaintiff's pay and benefits\b", "lost pay and benefits", text, flags=re.IGNORECASE)
        text = re.sub(r"\blost Plaintiff's (pay|wages|salary|income|benefits)\b", r"lost \1", text, flags=re.IGNORECASE)
        if text and text[0].islower():
            text = text[0].upper() + text[1:]
        if len(text) < 12:
            return ""
        return text if text.endswith((".", "?", "!")) else f"{text}."

    def _is_factual_allegation_candidate(self, value: Any) -> bool:
        text = re.sub(r"\s+", " ", str(value or "")).strip()
        if not text:
            return False
        lowered = text.lower()
        if re.match(r"^(as to [^,]+, )?plaintiff (seeks|requests|asks|demands)\b", lowered):
            return False
        if lowered.startswith(("requested relief", "relief requested", "element supported:")):
            return False
        if lowered.startswith(("evidence shows facts supporting", "the intake record describes facts supporting")):
            return False
        if re.match(r"^(as to [^,]+, )?(title\s+[ivxlcdm0-9]+\b|\d+\s+u\.s\.c\.|\d+\s+c\.f\.r\.|[a-z]{2,6}\.\s+gov\.\s+code\b)", lowered):
            return False
        if not re.search(
            r"\b(was|were|is|are|reported|complained|terminated|fired|retaliated|denied|refused|told|informed|notified|requested|sought|experienced|suffered|lost|made|engaged|opposed|filed|sent|emailed|wrote|received|occurred|happened|subjected|demoted|suspended|disciplined|reduced)\b",
            lowered,
        ):
            return False
        return True

    def _is_generic_claim_support_text(self, value: Any) -> bool:
        lowered = re.sub(r"\s+", " ", str(value or "")).strip().lower()
        return lowered.startswith(("evidence shows facts supporting", "the intake record describes facts supporting"))

    def _expand_allegation_sources(self, values: Any, *, limit: Optional[int] = None) -> List[str]:
        expanded: List[str] = []
        for value in _extract_text_candidates(values):
            for fragment in self._split_allegation_fragments(value):
                sentence = self._formalize_allegation_fragment(fragment)
                if not sentence or not self._is_factual_allegation_candidate(sentence):
                    continue
                expanded.append(sentence)
        unique = _unique_preserving_order(expanded)
        return unique[:limit] if limit is not None else unique

    def _synthesize_narrative_allegations(self, allegations: List[str]) -> List[str]:
        cleaned = [str(item).strip() for item in allegations if str(item).strip()]
        if not cleaned:
            return []

        def _normalize_adverse_clause(clause: str) -> str:
            text = str(clause or "").strip().rstrip(".!?")
            if re.match(r"^(after|following)\b", text, flags=re.IGNORECASE) and "," in text:
                text = text.split(",", 1)[1].strip()
            return text

        def _normalize_harm_clause(clause: str) -> str:
            text = str(clause or "").strip().rstrip(".!?")
            text = re.sub(r",?\s+as a result$", "", text, flags=re.IGNORECASE)
            text = re.sub(r",?\s+as a direct result$", "", text, flags=re.IGNORECASE)
            return text.strip()

        def _pick(pattern: str, *, require_plaintiff: bool = False) -> str:
            for item in cleaned:
                lowered = item.lower()
                if require_plaintiff and "plaintiff" not in lowered:
                    continue
                if re.search(pattern, lowered):
                    return item.rstrip(".!?")
            return ""

        report_clause = _pick(r"\b(reported|complained|opposed|informed|notified|told|requested)\b", require_plaintiff=True)
        adverse_clause = _pick(r"\b(terminated|fired|demoted|suspended|disciplined|retaliated|denied)\b")
        harm_clause = _pick(r"\blost (pay|wages|salary|income|benefits)\b|\b(suffered|experienced)\b", require_plaintiff=True)
        harm_already_tied_to_adverse_action = any(
            re.search(r"\b(lost|suffered|experienced)\b", item.lower())
            and re.search(r"\b(terminated|fired|demoted|suspended|disciplined|retaliated|denied)\b", item.lower())
            for item in cleaned
        )

        synthesized: List[str] = []
        if report_clause and adverse_clause:
            synthesized.append(f"After {report_clause}, {_normalize_adverse_clause(adverse_clause)}.")
        if harm_clause and not harm_already_tied_to_adverse_action:
            normalized_harm_clause = _normalize_harm_clause(harm_clause)
            loss_match = re.search(r"\blost ([^.]+)", normalized_harm_clause, flags=re.IGNORECASE)
            if loss_match:
                synthesized.append(f"As a direct result of Defendant's conduct, Plaintiff lost {loss_match.group(1).strip()}." )
        return _unique_preserving_order(synthesized)

    def _prune_subsumed_narrative_clauses(self, allegations: List[str]) -> List[str]:
        cleaned = [str(item).strip() for item in allegations if str(item).strip()]
        if not cleaned:
            return []

        def _pick(pattern: str, *, require_plaintiff: bool = False) -> str:
            for item in cleaned:
                lowered = item.lower()
                if require_plaintiff and "plaintiff" not in lowered:
                    continue
                if re.search(pattern, lowered):
                    return item.strip()
            return ""

        report_clause = _pick(r"\b(reported|complained|opposed|informed|notified|told|requested)\b", require_plaintiff=True)
        adverse_clause = _pick(r"\b(terminated|fired|demoted|suspended|disciplined|retaliated|denied)\b")
        has_harm_tied_to_adverse_action = any(
            re.search(r"\b(lost|suffered|experienced)\b", item.lower())
            and re.search(r"\b(terminated|fired|demoted|suspended|disciplined|retaliated|denied)\b", item.lower())
            for item in cleaned
        )
        consumed = {item.lower() for item in (report_clause, adverse_clause) if item}
        if has_harm_tied_to_adverse_action:
            combined_clause = _pick(
                r"\b(reported|complained|opposed|informed|notified|told|requested)\b.*\b(terminated|fired|demoted|suspended|disciplined|retaliated|denied)\b"
                r"|\b(terminated|fired|demoted|suspended|disciplined|retaliated|denied)\b.*\b(reported|complained|opposed|informed|notified|told|requested)\b",
                require_plaintiff=True,
            )
            if combined_clause:
                consumed.add(combined_clause.lower())
        return [item for item in cleaned if item.lower() not in consumed]

    def _prune_near_duplicate_allegations(self, allegations: List[str]) -> List[str]:
        def _tokens(value: str) -> set[str]:
            scrubbed = re.sub(r"\(see exhibit [^)]+\)", "", value, flags=re.IGNORECASE)
            return {
                token
                for token in re.split(r"\W+", scrubbed.lower())
                if len(token) >= 4 and token not in {"plaintiff", "defendant", "exhibit", "after", "those", "this", "that"}
            }

        def _categories(value: str) -> set[str]:
            lowered = value.lower()
            flags = set()
            if re.search(r"\b(reported|complained|opposed|informed|notified|told|requested)\b", lowered):
                flags.add("report")
            if re.search(r"\b(terminated|fired|demoted|suspended|disciplined|retaliated|denied|removed|stripped)\b", lowered) or re.search(r"\b(end(?:ed|ing))\b[^.]{0,40}\bemployment\b", lowered):
                flags.add("adverse")
            if re.search(r"\b(lost|suffered|experienced|benefits|wages|salary|income|opportunities)\b", lowered):
                flags.add("harm")
            return flags

        def _features(value: str) -> set[str]:
            lowered = value.lower()
            flags = set()
            if re.search(r"\b(reported|complained|opposed|informed|notified|told|requested)\b", lowered):
                flags.add("report")
            if re.search(r"\b(human resources|hr)\b", lowered):
                flags.add("hr")
            if re.search(r"\bregional management|management\b", lowered):
                flags.add("management")
            if re.search(r"\b(key|major)\s+accounts?\b|\b(accounts?)\b[^.]{0,20}\b(removed|stripped|taken away)\b|\b(removed|stripped|took away)\b[^.]{0,20}\baccounts?\b", lowered):
                flags.add("accounts")
            if re.search(r"\bovertime\b", lowered):
                flags.add("overtime")
            if re.search(r"\bshift(s)?\b", lowered):
                flags.add("shifts")
            if re.search(r"\b(absences?|attendance|treatment-related absences?)\b", lowered):
                flags.add("absences")
            if re.search(r"\b(disciplined|discipline|wrote me up|write-up|write up)\b", lowered):
                flags.add("discipline")
            if re.search(r"\b(accommodation|accommodate|light duty|schedule flexibility|medical restrictions?|doctor-imposed restrictions?)\b", lowered):
                flags.add("accommodation")
            if re.search(r"\b(restrictions?|light duty|schedule flexibility)\b", lowered):
                flags.add("restrictions")
            if re.search(r"\b(terminated|fired)\b", lowered) or re.search(r"\b(end(?:ed|ing))\b[^.]{0,40}\bemployment\b", lowered):
                flags.add("termination")
            if re.search(r"\b(wages|pay|salary|income|benefits)\b", lowered):
                flags.add("economic_harm")
            if re.search(r"\b(career opportunities|future opportunities|opportunities)\b", lowered):
                flags.add("opportunities")
            return flags

        kept: List[str] = []
        for candidate in allegations:
            candidate_tokens = _tokens(candidate)
            candidate_categories = _categories(candidate)
            candidate_features = _features(candidate)
            skip = False
            for existing in kept:
                existing_tokens = _tokens(existing)
                existing_categories = _categories(existing)
                existing_features = _features(existing)
                if not candidate_tokens or not existing_tokens:
                    continue
                if not (candidate_categories & existing_categories):
                    continue
                overlap = len(candidate_tokens & existing_tokens) / max(1, min(len(candidate_tokens), len(existing_tokens)))
                shared_features = candidate_features & existing_features
                if overlap >= 0.7:
                    skip = True
                    break
                if "adverse" in candidate_categories and "adverse" in existing_categories and len(shared_features) >= 3:
                    skip = True
                    break
                if "report" in candidate_categories and "report" in existing_categories and "accommodation" in shared_features and len(shared_features) >= 2:
                    skip = True
                    break
            if not skip:
                kept.append(candidate)
        return kept

    def _is_near_duplicate_allegation(self, candidate: str, existing: List[str]) -> bool:
        if not candidate:
            return False
        pruned = self._prune_near_duplicate_allegations([*existing, candidate])
        return len(pruned) == len(existing)

    def _build_factual_allegations(
        self,
        *,
        summary_of_facts: Any,
        claims_for_relief: List[Dict[str, Any]],
    ) -> List[str]:
        base_allegations = list(self._expand_allegation_sources(summary_of_facts, limit=14))
        allegations = list(self._synthesize_narrative_allegations(base_allegations))
        for item in self._prune_subsumed_narrative_clauses(base_allegations):
            if item.lower() not in {entry.lower() for entry in allegations}:
                allegations.append(item)
        seen = {entry.lower() for entry in allegations}

        for claim in _coerce_list(claims_for_relief):
            if not isinstance(claim, dict):
                continue
            count_title = str(claim.get("count_title") or claim.get("claim_type") or "Claim").strip()
            for fact in self._expand_allegation_sources(claim.get("supporting_facts", []), limit=10):
                if not fact:
                    continue
                if self._is_near_duplicate_allegation(fact, allegations):
                    continue
                prefixed_fact = fact
                if count_title and not fact.lower().startswith("as to ") and fact.lower() not in seen:
                    lowered = fact
                    if not re.match(r"^(Plaintiff|Defendant)\b", fact):
                        lowered = fact[0].lower() + fact[1:] if len(fact) > 1 and fact[0].isalpha() else fact
                    prefixed_fact = f"As to {count_title}, {lowered}"
                    if not prefixed_fact.endswith((".", "?", "!")):
                        prefixed_fact = f"{prefixed_fact}."
                key = prefixed_fact.lower()
                if key in seen:
                    continue
                seen.add(key)
                allegations.append(prefixed_fact)
                if len(allegations) >= 24:
                    return self._prune_near_duplicate_allegations(allegations)

        pruned = self._prune_near_duplicate_allegations(allegations)
        if pruned and not any(_contains_hearing_timing_marker(line) for line in pruned):
            pruned.append(
                "Plaintiff requested an informal review or hearing on [date], and the complaint should state when that request was made in relation to each adverse-action step."
            )
        if pruned and not any(_contains_response_date_marker(line) for line in pruned):
            pruned.append(
                "HACC response dates for notice, hearing/review requests, and final decision communications should be identified with exact dates."
            )
        if pruned and not any(_contains_staff_identity_marker(line) for line in pruned):
            pruned.append(
                "For each key event, the complaint should identify the HACC staff member by name and title, or by the best-known title if the name is not yet confirmed."
            )
        if pruned and not any("days after" in str(line).lower() or "weeks after" in str(line).lower() for line in pruned):
            pruned.append(
                "The complaint should describe the sequence between protected activity and adverse treatment using concrete timing, including whether action occurred days or weeks after protected activity."
            )
        return pruned[:24] or ["Additional factual development is required before filing."]

    def _attach_allegation_references(self, draft: Dict[str, Any]) -> None:
        allegation_lines = self._normalize_text_lines(
            draft.get("factual_allegations") or draft.get("summary_of_facts", [])
        )
        paragraph_entries = [
            {
                "number": index,
                "text": text,
            }
            for index, text in enumerate(allegation_lines, start=1)
        ]
        draft["factual_allegations"] = allegation_lines
        draft["factual_allegation_paragraphs"] = paragraph_entries
        draft["factual_allegation_groups"] = self._build_factual_allegation_groups(paragraph_entries)

        claims = draft.get("claims_for_relief") if isinstance(draft.get("claims_for_relief"), list) else []
        for claim in claims:
            if not isinstance(claim, dict):
                continue
            claim["allegation_references"] = self._select_allegation_references_for_claim(
                claim=claim,
                allegation_paragraphs=paragraph_entries,
            )

    def _build_factual_allegation_groups(self, allegation_paragraphs: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        ordered_titles = [
            "Protected Activity and Complaints",
            "Adverse Action and Retaliatory Conduct",
            "Damages and Resulting Harm",
            "Additional Factual Support",
        ]
        groups: Dict[str, List[Dict[str, Any]]] = {title: [] for title in ordered_titles}

        for paragraph in allegation_paragraphs:
            if not isinstance(paragraph, dict):
                continue
            text = str(paragraph.get("text") or "").strip()
            lowered = text.lower()
            if re.search(r"\b(reported|complained|opposed|informed|notified|told|requested)\b", lowered):
                title = "Protected Activity and Complaints"
            elif re.search(r"\b(terminated|fired|demoted|suspended|disciplined|retaliated|denied)\b", lowered):
                title = "Adverse Action and Retaliatory Conduct"
            elif re.search(r"\b(lost|damages|harm|injur|suffered|experienced|benefits|wages|salary|income)\b", lowered):
                title = "Damages and Resulting Harm"
            else:
                title = "Additional Factual Support"
            groups[title].append(paragraph)

        return [
            {"title": title, "paragraphs": groups[title]}
            for title in ordered_titles
            if groups[title]
        ]

    def _grouped_allegation_text_lines(self, draft: Dict[str, Any]) -> List[str]:
        groups = draft.get("factual_allegation_groups") if isinstance(draft.get("factual_allegation_groups"), list) else []
        if not groups:
            return self._numbered_lines(draft.get("factual_allegations") or draft.get("summary_of_facts", []))

        lines: List[str] = []
        for group in groups:
            if not isinstance(group, dict):
                continue
            title = str(group.get("title") or "").strip()
            paragraphs = group.get("paragraphs") if isinstance(group.get("paragraphs"), list) else []
            if not paragraphs:
                continue
            if title:
                lines.append(title.upper())
            for paragraph in paragraphs:
                if not isinstance(paragraph, dict):
                    continue
                number = paragraph.get("number")
                text = str(paragraph.get("text") or "").strip()
                if text:
                    lines.append(f"{number}. {text}" if number else text)
        return lines

    def _select_allegation_references_for_claim(
        self,
        *,
        claim: Dict[str, Any],
        allegation_paragraphs: List[Dict[str, Any]],
    ) -> List[int]:
        references: List[int] = []
        supporting_facts = self._normalize_text_lines(claim.get("supporting_facts", []))
        count_title = str(claim.get("count_title") or claim.get("claim_type") or "").strip().lower()

        for fact in supporting_facts:
            fact_tokens = self._text_tokens(fact)
            if not fact_tokens:
                continue
            best_number: Optional[int] = None
            best_score = 0
            fact_lower = fact.lower()
            for paragraph in allegation_paragraphs:
                if not isinstance(paragraph, dict):
                    continue
                paragraph_text = str(paragraph.get("text") or "").strip()
                paragraph_lower = paragraph_text.lower()
                paragraph_tokens = self._text_tokens(paragraph_text)
                score = len(fact_tokens & paragraph_tokens)
                if fact_lower in paragraph_lower:
                    score += 100
                if count_title and count_title in paragraph_lower:
                    score += 5
                if score > best_score:
                    best_score = score
                    best_number = int(paragraph.get("number", 0) or 0)
            if best_number and best_number not in references:
                references.append(best_number)
                if len(references) >= 6:
                    break

        if references:
            return references

        fallback = []
        for paragraph in allegation_paragraphs:
            paragraph_text = str(paragraph.get("text") or "").lower()
            if count_title and count_title in paragraph_text:
                fallback.append(int(paragraph.get("number", 0) or 0))
        return fallback[:4]

    def _format_paragraph_reference_clause(self, references: Any) -> str:
        values = []
        for value in _coerce_list(references):
            try:
                number = int(value)
            except (TypeError, ValueError):
                continue
            if number > 0 and number not in values:
                values.append(number)
        if not values:
            return ""
        citation = self._format_paragraph_citation(values)
        return f"Plaintiff repeats and realleges {citation} as if fully set forth herein."

    def _format_incorporated_reference_clause(self, references: Any, exhibits: Any) -> str:
        paragraph_citation = self._format_paragraph_citation(references)
        exhibit_phrase = self._format_exhibit_reference_phrase(exhibits)
        if paragraph_citation and exhibit_phrase:
            return (
                f"Plaintiff repeats and realleges {paragraph_citation} and incorporates {exhibit_phrase} "
                "as if fully set forth herein."
            )
        if paragraph_citation:
            return f"Plaintiff repeats and realleges {paragraph_citation} as if fully set forth herein."
        if exhibit_phrase:
            return f"Plaintiff incorporates {exhibit_phrase} as if fully set forth herein."
        return ""

    def _format_paragraph_citation(self, references: Any) -> str:
        values = []
        for value in _coerce_list(references):
            try:
                number = int(value)
            except (TypeError, ValueError):
                continue
            if number > 0 and number not in values:
                values.append(number)
        if not values:
            return ""
        values.sort()
        ranges: List[str] = []
        range_start = values[0]
        range_end = values[0]
        for number in values[1:]:
            if number == range_end + 1:
                range_end = number
                continue
            ranges.append(self._format_paragraph_range(range_start, range_end))
            range_start = number
            range_end = number
        ranges.append(self._format_paragraph_range(range_start, range_end))
        marker = "¶" if len(values) == 1 else "¶¶"
        return f"{marker} {', '.join(ranges)}"

    def _format_exhibit_reference_phrase(self, exhibits: Any) -> str:
        labels = []
        for exhibit in _coerce_list(exhibits):
            if not isinstance(exhibit, dict):
                continue
            label = str(exhibit.get("label") or "").strip()
            if label and label not in labels:
                labels.append(label)
        if not labels:
            return ""
        if len(labels) == 1:
            return labels[0]
        if len(labels) == 2:
            return f"{labels[0]} and {labels[1]}"
        return f"{', '.join(labels[:-1])}, and {labels[-1]}"

    def _format_paragraph_range(self, start: int, end: int) -> str:
        return str(start) if start == end else f"{start}-{end}"

    def _numbered_lines(self, values: Any) -> List[str]:
        return [f"{index}. {line}" for index, line in enumerate(self._normalize_text_lines(values), start=1)]

    def _bulletize_lines(self, values: Any) -> List[str]:
        return [f"- {line}" for line in self._normalize_text_lines(values)]

    def _build_signature_block(
        self,
        plaintiffs: List[str],
        *,
        signer_name: Optional[str] = None,
        signer_title: Optional[str] = None,
        signer_firm: Optional[str] = None,
        signer_bar_number: Optional[str] = None,
        signer_contact: Optional[str] = None,
        additional_signers: Optional[List[Dict[str, str]]] = None,
        signature_date: Optional[str] = None,
    ) -> Dict[str, Any]:
        plaintiff_name = str(signer_name or "").strip() or (plaintiffs or ["Plaintiff"])[0]
        return {
            "name": plaintiff_name,
            "signature_line": f"/s/ {plaintiff_name}",
            "title": str(signer_title or "").strip() or "Plaintiff, Pro Se",
            "firm": str(signer_firm or "").strip() or "",
            "bar_number": str(signer_bar_number or "").strip(),
            "contact": str(signer_contact or "").strip() or "Mailing address, telephone number, and email address to be completed before filing.",
            "additional_signers": self._normalize_additional_signers(additional_signers),
            "dated": self._format_dated_line("Dated", signature_date),
        }

    def _normalize_additional_signers(self, values: Any) -> List[Dict[str, str]]:
        normalized: List[Dict[str, str]] = []
        seen: set[tuple[str, str, str, str, str]] = set()
        for item in _coerce_list(values):
            if not isinstance(item, dict):
                continue
            name = str(item.get("name") or item.get("signer_name") or "").strip()
            title = str(item.get("title") or item.get("signer_title") or "").strip()
            firm = str(item.get("firm") or item.get("signer_firm") or "").strip()
            bar_number = str(item.get("bar_number") or item.get("signer_bar_number") or "").strip()
            contact = str(item.get("contact") or item.get("signer_contact") or "").strip()
            if not any((name, title, firm, bar_number, contact)):
                continue
            key = (name, title, firm, bar_number, contact)
            if key in seen:
                continue
            seen.add(key)
            normalized.append(
                {
                    "name": name or "Additional Counsel",
                    "signature_line": f"/s/ {name}" if name else "",
                    "title": title,
                    "firm": firm,
                    "bar_number": bar_number,
                    "contact": contact,
                }
            )
        return normalized

    def _signature_block_lines(self, signature_block: Dict[str, Any], *, include_dated: bool = True) -> List[str]:
        lines: List[str] = [
            str(signature_block.get("signature_line") or "/s/ Plaintiff"),
            str(signature_block.get("name") or "Plaintiff"),
        ]
        for key in ("title", "firm"):
            if signature_block.get(key):
                lines.append(str(signature_block[key]))
        if signature_block.get("bar_number"):
            lines.append(f"Bar No. {signature_block['bar_number']}")
        if signature_block.get("contact"):
            lines.append(str(signature_block["contact"]))
        for signer in _coerce_list(signature_block.get("additional_signers")):
            if not isinstance(signer, dict):
                continue
            lines.append("")
            if signer.get("signature_line"):
                lines.append(str(signer["signature_line"]))
            lines.append(str(signer.get("name") or "Additional Counsel"))
            for key in ("title", "firm"):
                if signer.get(key):
                    lines.append(str(signer[key]))
            if signer.get("bar_number"):
                lines.append(f"Bar No. {signer['bar_number']}")
            if signer.get("contact"):
                lines.append(str(signer["contact"]))
        if include_dated and signature_block.get("dated"):
            lines.append(str(signature_block["dated"]))
        return lines

    def _build_signature_section_lines(self, signature_block: Dict[str, Any], forum_type: str) -> List[str]:
        if forum_type == "state":
            lines: List[str] = []
            if signature_block.get("dated"):
                lines.append(str(signature_block["dated"]))
            lines.extend(["", "Respectfully submitted,", *self._signature_block_lines(signature_block, include_dated=False)])
            return lines
        return ["Respectfully submitted,", *self._signature_block_lines(signature_block)]

    def _build_jury_demand(
        self,
        *,
        jury_demand: Optional[bool] = None,
        jury_demand_text: Optional[str] = None,
    ) -> Dict[str, str]:
        text = str(jury_demand_text or "").strip()
        if text:
            return {
                "title": "Jury Demand",
                "text": text if text.endswith((".", "?", "!")) else f"{text}.",
            }
        if jury_demand:
            return {
                "title": "Jury Demand",
                "text": "Plaintiff demands a trial by jury on all issues so triable.",
            }
        return {}

    def _build_verification(
        self,
        plaintiffs: List[str],
        *,
        declarant_name: Optional[str] = None,
        signer_name: Optional[str] = None,
        verification_date: Optional[str] = None,
        jurisdiction: Optional[str] = None,
    ) -> Dict[str, str]:
        plaintiff_name = str(declarant_name or "").strip() or str(signer_name or "").strip() or (plaintiffs or ["Plaintiff"])[0]
        is_state = str(jurisdiction or "").strip().lower() == "state"
        return {
            "title": "Verification",
            "text": (
                f"I, {plaintiff_name}, verify that I have reviewed this Complaint and know its contents. "
                "The facts stated in this Complaint are true of my own knowledge, except as to those matters "
                "stated on information and belief, and as to those matters I believe them to be true."
                if is_state
                else (
                    f"I, {plaintiff_name}, declare under penalty of perjury that I have reviewed this Complaint "
                    "and that the factual allegations stated in it are true and correct to the best of my knowledge, "
                    "information, and belief."
                )
            ),
            "dated": self._format_dated_line("Verified on" if is_state else "Executed on", verification_date),
            "signature_line": f"/s/ {plaintiff_name}",
        }

    def _build_certificate_of_service(
        self,
        plaintiffs: List[str],
        defendants: List[str],
        *,
        signer_name: Optional[str] = None,
        service_method: Optional[str] = None,
        service_recipients: Optional[List[str]] = None,
        service_recipient_details: Optional[List[Dict[str, str]]] = None,
        service_date: Optional[str] = None,
        jurisdiction: Optional[str] = None,
    ) -> Dict[str, Any]:
        plaintiff_name = str(signer_name or "").strip() or (plaintiffs or ["Plaintiff"])[0]
        recipient_details = self._normalize_service_recipient_details(service_recipient_details)
        detail_recipients = [detail["recipient"] for detail in recipient_details if detail.get("recipient")]
        recipients_list = _unique_preserving_order([str(item or "").strip() for item in _coerce_list(service_recipients)] + detail_recipients) or defendants or ["all defendants"]
        recipients = ", ".join(recipients_list)
        method_text = str(service_method or "").strip() or "a method authorized by the applicable rules of civil procedure"
        detail_lines = [self._format_service_recipient_detail(detail) for detail in recipient_details]
        is_state = str(jurisdiction or "").strip().lower() == "state"
        return {
            "title": "Proof of Service" if is_state else "Certificate of Service",
            "text": (
                ("I declare that a true and correct copy of this Complaint will be served promptly after filing on the following recipients."
                if is_state
                else "I certify that a true and correct copy of this Complaint will be served promptly after filing on the following recipients.")
                if detail_lines
                else (("I declare that a true and correct copy of this Complaint will be served on "
                if is_state else "I certify that a true and correct copy of this Complaint will be served on ")
                + f"{recipients} using {method_text} promptly after filing.")
            ),
            "recipients": recipients_list,
            "recipient_details": recipient_details,
            "detail_lines": detail_lines,
            "dated": self._format_dated_line("Service date", service_date),
            "signature_line": f"/s/ {plaintiff_name}",
        }

    def _normalize_service_recipient_details(self, values: Any) -> List[Dict[str, str]]:
        details: List[Dict[str, str]] = []
        seen = set()
        for item in _coerce_list(values):
            if not isinstance(item, dict):
                continue
            detail = {
                "recipient": str(item.get("recipient") or "").strip(),
                "method": str(item.get("method") or "").strip(),
                "address": str(item.get("address") or "").strip(),
                "notes": str(item.get("notes") or "").strip(),
            }
            if not any(detail.values()):
                continue
            key = (detail["recipient"], detail["method"], detail["address"], detail["notes"])
            if key in seen:
                continue
            seen.add(key)
            details.append(detail)
        return details

    def _format_service_recipient_detail(self, detail: Dict[str, str]) -> str:
        segments = [detail.get("recipient") or "Recipient"]
        if detail.get("method"):
            segments.append(f"Method: {detail['method']}")
        if detail.get("address"):
            segments.append(f"Address: {detail['address']}")
        if detail.get("notes"):
            segments.append(f"Notes: {detail['notes']}")
        return " | ".join(segment for segment in segments if segment)

    def _format_dated_line(self, label: str, value: Optional[str]) -> str:
        cleaned = str(value or "").strip()
        return f"{label}: {cleaned}" if cleaned else f"{label}: __________________"

    def render_artifacts(
        self,
        draft: Dict[str, Any],
        *,
        output_dir: Optional[str],
        output_formats: List[str],
    ) -> Dict[str, Dict[str, Any]]:
        output_root = Path(output_dir).expanduser() if output_dir else DEFAULT_OUTPUT_DIR
        output_root.mkdir(parents=True, exist_ok=True)
        timestamp = _utcnow().strftime("%Y%m%dT%H%M%SZ")
        file_stem = f"{_slugify(draft.get('title') or 'complaint')}-{timestamp}"
        artifacts: Dict[str, Dict[str, Any]] = {}

        for output_format in output_formats:
            if output_format == "packet":
                continue
            path = self._artifact_path(output_root, file_stem, output_format)
            if output_format == "docx":
                self._render_docx(draft, path)
                affidavit_path = self._artifact_path(output_root, file_stem, output_format, document_kind="affidavit")
                self._render_affidavit_docx(draft, affidavit_path)
                artifacts["affidavit_docx"] = {
                    "path": str(affidavit_path),
                    "filename": affidavit_path.name,
                    "size_bytes": affidavit_path.stat().st_size,
                }
            elif output_format == "pdf":
                self._render_pdf(draft, path)
                affidavit_path = self._artifact_path(output_root, file_stem, output_format, document_kind="affidavit")
                self._render_affidavit_pdf(draft, affidavit_path)
                artifacts["affidavit_pdf"] = {
                    "path": str(affidavit_path),
                    "filename": affidavit_path.name,
                    "size_bytes": affidavit_path.stat().st_size,
                }
            elif output_format == "txt":
                self._render_txt(draft, path)
                affidavit_path = self._artifact_path(output_root, file_stem, output_format, document_kind="affidavit")
                self._render_affidavit_txt(draft, affidavit_path)
                artifacts["affidavit_txt"] = {
                    "path": str(affidavit_path),
                    "filename": affidavit_path.name,
                    "size_bytes": affidavit_path.stat().st_size,
                }
            elif output_format == "checklist":
                self._render_checklist_txt(draft, path)
            artifacts[output_format] = {
                "path": str(path),
                "filename": path.name,
                "size_bytes": path.stat().st_size,
            }

        if "packet" in output_formats:
            path = self._artifact_path(output_root, file_stem, "packet")
            self._render_packet_json(draft, path, artifacts=artifacts)
            artifacts["packet"] = {
                "path": str(path),
                "filename": path.name,
                "size_bytes": path.stat().st_size,
            }

        return artifacts

    def _resolve_user_id(self, user_id: Optional[str]) -> str:
        if user_id:
            return user_id
        state = getattr(self.mediator, "state", None)
        return (
            getattr(state, "username", None)
            or getattr(state, "hashed_username", None)
            or "anonymous"
        )

    def _normalize_formats(self, output_formats: Optional[List[str]]) -> List[str]:
        values = output_formats or ["docx", "pdf"]
        normalized = []
        for value in values:
            current = str(value or "").strip().lower()
            if current in {"docx", "pdf", "txt", "checklist", "packet"} and current not in normalized:
                normalized.append(current)
        return normalized or ["docx", "pdf"]

    def _build_affidavit_overrides(
        self,
        *,
        affidavit_title: Optional[str],
        affidavit_intro: Optional[str],
        affidavit_facts: Optional[List[str]],
        affidavit_supporting_exhibits: Optional[List[Dict[str, str]]],
        affidavit_include_complaint_exhibits: Optional[bool],
        affidavit_venue_lines: Optional[List[str]],
        affidavit_jurat: Optional[str],
        affidavit_notary_block: Optional[List[str]],
    ) -> Dict[str, Any]:
        normalized_override_facts = []
        for value in affidavit_facts or []:
            cleaned = self._sanitize_affidavit_fact(value)
            if cleaned:
                normalized_override_facts.append(cleaned)
        normalized_supporting_exhibits = []
        for exhibit in _coerce_list(affidavit_supporting_exhibits):
            if not isinstance(exhibit, dict):
                continue
            normalized = {
                "label": str(exhibit.get("label") or "Exhibit").strip(),
                "title": str(exhibit.get("title") or exhibit.get("summary") or "Supporting exhibit").strip(),
                "link": str(exhibit.get("link") or exhibit.get("reference") or "").strip(),
                "summary": str(exhibit.get("summary") or "").strip(),
            }
            if any(normalized.values()):
                normalized_supporting_exhibits.append(normalized)
        return {
            "title": str(affidavit_title or "").strip() or None,
            "intro": str(affidavit_intro or "").strip() or None,
            "facts": normalized_override_facts,
            "supporting_exhibits": normalized_supporting_exhibits,
            "include_complaint_exhibits": affidavit_include_complaint_exhibits,
            "venue_lines": self._normalize_text_lines(affidavit_venue_lines or []),
            "jurat": str(affidavit_jurat or "").strip() or None,
            "notary_block": self._normalize_text_lines(affidavit_notary_block or []),
        }

    def _build_affidavit(self, draft: Dict[str, Any]) -> Dict[str, Any]:
        verification = draft.get("verification", {}) if isinstance(draft.get("verification"), dict) else {}
        signature_block = draft.get("signature_block", {}) if isinstance(draft.get("signature_block"), dict) else {}
        case_caption = draft.get("case_caption", {}) if isinstance(draft.get("case_caption"), dict) else {}
        affidavit_overrides = draft.get("affidavit_overrides", {}) if isinstance(draft.get("affidavit_overrides"), dict) else {}
        declarant_name = self._derive_affidavit_declarant_name(draft)
        is_state = self._resolve_draft_forum_type(draft) == "state"
        exhibits = []
        for exhibit in _coerce_list(draft.get("exhibits")):
            if not isinstance(exhibit, dict):
                continue
            exhibits.append(
                {
                    "label": str(exhibit.get("label") or "Exhibit").strip(),
                    "title": str(exhibit.get("title") or exhibit.get("summary") or "Supporting exhibit").strip(),
                    "link": str(exhibit.get("link") or exhibit.get("reference") or "").strip(),
                    "summary": str(exhibit.get("summary") or "").strip(),
                }
            )
        return {
            "title": str(affidavit_overrides.get("title") or f"AFFIDAVIT OF {declarant_name.upper()} IN SUPPORT OF COMPLAINT"),
            "declarant_name": declarant_name,
            "intro": str(
                affidavit_overrides.get("intro")
                or (
                    (
                        f"I, {declarant_name}, being duly sworn, state that I am competent to testify to the matters stated below, "
                        "that these statements are based on my personal knowledge and the complaint intake knowledge graph assembled from the facts, records, and exhibits provided in support of this action, and that the following facts are true and correct."
                    )
                    if is_state
                    else (
                        f"I, {declarant_name}, declare under penalty of perjury that I am competent to testify to the matters stated below, "
                        "that these statements are based on my personal knowledge and the complaint intake knowledge graph assembled from the facts, records, and exhibits provided in support of this action, and that the following facts are true and correct."
                    )
                )
            ),
            "knowledge_graph_note": "This affidavit is generated from the complaint intake knowledge graph and supporting records rather than a turn-by-turn chat transcript.",
            "venue_lines": list(affidavit_overrides.get("venue_lines") or self._build_affidavit_venue_lines(draft)),
            "facts": list(affidavit_overrides.get("facts") or self._collect_affidavit_facts(draft)),
            "supporting_exhibits": list(
                affidavit_overrides.get("supporting_exhibits")
                or ([] if affidavit_overrides.get("include_complaint_exhibits") is False else exhibits)
            ),
            "dated": str(verification.get("dated") or signature_block.get("dated") or self._format_dated_line("Verified on" if is_state else "Executed on", None)),
            "signature_line": str(verification.get("signature_line") or signature_block.get("signature_line") or f"/s/ {declarant_name}"),
            "jurat": str(
                affidavit_overrides.get("jurat")
                or (
                    f"Subscribed and sworn to before me on __________________ by {declarant_name}."
                    if is_state
                    else f"Subscribed and sworn to (or affirmed) before me on __________________ by {declarant_name}."
                )
            ),
            "notary_block": list(
                affidavit_overrides.get("notary_block")
                or [
                    "__________________________________",
                    "Notary Public",
                    "My commission expires: __________________",
                ]
            ),
            "case_number": str(case_caption.get("case_number") or "________________"),
        }

    def _derive_affidavit_declarant_name(self, draft: Dict[str, Any]) -> str:
        verification = draft.get("verification", {}) if isinstance(draft.get("verification"), dict) else {}
        signature_block = draft.get("signature_block", {}) if isinstance(draft.get("signature_block"), dict) else {}
        parties = draft.get("parties", {}) if isinstance(draft.get("parties"), dict) else {}
        signature_line = str(verification.get("signature_line") or signature_block.get("signature_line") or "").strip()
        if signature_line.startswith("/s/ "):
            return signature_line[4:].strip() or str(signature_block.get("name") or "Plaintiff")
        plaintiffs = [str(name).strip() for name in _coerce_list(parties.get("plaintiffs")) if str(name).strip()]
        return str(signature_block.get("name") or (plaintiffs[0] if plaintiffs else "Plaintiff")).strip() or "Plaintiff"

    def _build_affidavit_venue_lines(self, draft: Dict[str, Any]) -> List[str]:
        caption = draft.get("case_caption", {}) if isinstance(draft.get("case_caption"), dict) else {}
        source_context = draft.get("source_context", {}) if isinstance(draft.get("source_context"), dict) else {}
        county = str(caption.get("county") or "").strip()
        district = str(source_context.get("district") or "").strip()
        jurisdiction = str(source_context.get("jurisdiction") or caption.get("forum_type") or "").strip().lower()
        lines: List[str] = []
        if district:
            lines.append(f"State/District: {district}")
        elif jurisdiction == "federal":
            lines.append("State/District: __________________")
        if county:
            lines.append(f"County: {county.title()}")
        elif jurisdiction == "state":
            lines.append("County: __________________")
        return lines or ["Venue: __________________"]

    def _collect_affidavit_facts(self, draft: Dict[str, Any]) -> List[str]:
        candidates: List[str] = []
        parties = draft.get("parties", {}) if isinstance(draft.get("parties"), dict) else {}
        plaintiffs = [str(name).strip() for name in _coerce_list(parties.get("plaintiffs")) if str(name).strip()]
        if plaintiffs:
            candidates.append(f"I am {plaintiffs[0]}, the plaintiff in this action.")
        candidates.extend(self._normalize_text_lines(draft.get("factual_allegations", [])))

        facts: List[str] = []
        seen: set[str] = set()
        for candidate in candidates:
            cleaned = self._sanitize_affidavit_fact(candidate)
            if not cleaned:
                continue
            key = cleaned.lower()
            if key in seen:
                continue
            seen.add(key)
            facts.append(cleaned)
            if len(facts) >= 12:
                break
        return facts or ["Additional fact development is required before the affidavit can be finalized."]

    def _sanitize_affidavit_fact(self, value: str) -> str:
        text = re.sub(r"\s+", " ", str(value or "")).strip()
        if not text:
            return ""
        text = re.sub(r"^As to [^,]+,\s*", "", text, flags=re.IGNORECASE)
        if ": " in text:
            prefix, suffix = text.split(": ", 1)
            prefix_lower = prefix.strip().lower()
            if (
                prefix.strip().endswith("?")
                or prefix_lower.startswith(("what ", "when ", "where ", "why ", "how ", "who ", "describe ", "explain "))
                or prefix_lower in {"what happened", "what relief do you want"}
            ):
                text = suffix.strip()
        lowered = text.lower()
        if lowered.startswith("plaintiff repeats and realleges"):
            return ""
        if not self._is_factual_allegation_candidate(text) and not lowered.startswith("i am "):
            return ""
        if len(text) < 12:
            return ""
        if text and text[0].islower():
            text = text[0].upper() + text[1:]
        if text[-1] not in ".!?":
            text = f"{text}."
        return text

    def _render_txt(self, draft: Dict[str, Any], path: Path) -> None:
        path.write_text(str(draft.get("draft_text") or self._render_draft_text(draft)), encoding="utf-8")

    def _render_affidavit_txt(self, draft: Dict[str, Any], path: Path) -> None:
        affidavit = draft.get("affidavit", {}) if isinstance(draft.get("affidavit"), dict) else self._build_affidavit(draft)
        path.write_text(self._render_affidavit_text(draft, affidavit), encoding="utf-8")

    def _render_checklist_txt(self, draft: Dict[str, Any], path: Path) -> None:
        checklist = draft.get("filing_checklist") if isinstance(draft.get("filing_checklist"), list) else []
        title = str(draft.get("title") or draft.get("case_caption", {}).get("document_title") or "Complaint").strip()
        lines = [
            f"PRE-FILING CHECKLIST: {title}",
            "",
        ]
        if not checklist:
            lines.append("No pre-filing checklist items were generated.")
        else:
            for index, item in enumerate(checklist, start=1):
                if not isinstance(item, dict):
                    continue
                scope = str(item.get("scope") or "item").strip().upper()
                title_text = str(item.get("title") or "Checklist Item").strip()
                status = str(item.get("status") or "ready").strip().upper()
                summary = str(item.get("summary") or "").strip()
                detail = str(item.get("detail") or "").strip()
                review_url = str(item.get("review_url") or "").strip()
                lines.append(f"{index}. [{status}] {scope}: {title_text}")
                if summary:
                    lines.append(f"   Summary: {summary}")
                if detail:
                    lines.append(f"   Detail: {detail}")
                if review_url:
                    lines.append(f"   Review URL: {review_url}")
                lines.append("")
        path.write_text("\n".join(lines).rstrip() + "\n", encoding="utf-8")

    def _artifact_path(self, output_root: Path, file_stem: str, output_format: str, document_kind: str = "complaint") -> Path:
        suffix = "-affidavit" if document_kind == "affidavit" else ""
        if output_format == "checklist":
            return output_root / f"{file_stem}{suffix}-checklist.txt"
        if output_format == "packet":
            return output_root / f"{file_stem}-packet.json"
        return output_root / f"{file_stem}{suffix}.{output_format}"

    def _render_packet_json(
        self,
        draft: Dict[str, Any],
        path: Path,
        *,
        artifacts: Dict[str, Dict[str, Any]],
    ) -> None:
        payload = self._build_filing_packet_payload(draft, artifacts=artifacts)
        path.write_text(json.dumps(payload, indent=2, sort_keys=True), encoding="utf-8")

    def _build_filing_packet_payload(
        self,
        draft: Dict[str, Any],
        *,
        artifacts: Dict[str, Dict[str, Any]],
    ) -> Dict[str, Any]:
        case_caption = draft.get("case_caption", {}) if isinstance(draft.get("case_caption"), dict) else {}
        source_context = draft.get("source_context", {}) if isinstance(draft.get("source_context"), dict) else {}
        packet_artifacts = {
            key: {
                "filename": value.get("filename"),
                "path": value.get("path"),
                "size_bytes": value.get("size_bytes"),
            }
            for key, value in artifacts.items()
            if isinstance(value, dict)
        }
        return {
            "title": draft.get("title"),
            "court_header": draft.get("court_header"),
            "generated_at": source_context.get("generated_at") or _utcnow().isoformat(),
            "claim_support_temporal_handoff": dict(source_context.get("claim_support_temporal_handoff") or {}) if isinstance(source_context.get("claim_support_temporal_handoff"), dict) else {},
            "source_context": source_context,
            "case_caption": {
                "plaintiffs": case_caption.get("plaintiffs", []),
                "defendants": case_caption.get("defendants", []),
                "case_number": case_caption.get("case_number"),
                "document_title": case_caption.get("document_title"),
                "jury_demand_notice": case_caption.get("jury_demand_notice"),
            },
            "sections": {
                "nature_of_action": draft.get("nature_of_action", []),
                "summary_of_facts": draft.get("summary_of_facts", []),
                "factual_allegations": draft.get("factual_allegations", []),
                "anchored_chronology_summary": draft.get("anchored_chronology_summary", []),
                "claims_for_relief": draft.get("claims_for_relief", []),
                "legal_standards": draft.get("legal_standards", []),
                "requested_relief": draft.get("requested_relief", []),
            },
            "affidavit": draft.get("affidavit", {}),
            "verification": draft.get("verification", {}),
            "certificate_of_service": draft.get("certificate_of_service", {}),
            "exhibits": draft.get("exhibits", []),
            "filing_checklist": draft.get("filing_checklist", []),
            "drafting_readiness": draft.get("drafting_readiness", {}),
            "artifacts": packet_artifacts,
        }

    def _render_affidavit_text(self, draft: Dict[str, Any], affidavit: Dict[str, Any]) -> str:
        caption = draft.get("case_caption", {}) if isinstance(draft.get("case_caption"), dict) else {}
        caption_party_lines = caption.get("caption_party_lines") if isinstance(caption.get("caption_party_lines"), list) else self._build_caption_party_lines(caption)
        lines = [
            str(draft.get("court_header") or "IN THE COURT OF COMPETENT JURISDICTION"),
            *([str(caption.get("county"))] if caption.get("county") else []),
            "",
            *caption_party_lines,
            f"{caption.get('case_number_label', 'Civil Action No.')} {caption.get('case_number', '________________')}",
            "",
            str(affidavit.get("title") or "AFFIDAVIT IN SUPPORT OF COMPLAINT"),
            *[str(line) for line in _coerce_list(affidavit.get("venue_lines")) if str(line or "").strip()],
            "",
            str(affidavit.get("intro") or ""),
            str(affidavit.get("knowledge_graph_note") or ""),
            "",
            "Affiant states as follows:",
            *self._numbered_lines(affidavit.get("facts", [])),
        ]
        exhibits = affidavit.get("supporting_exhibits") if isinstance(affidavit.get("supporting_exhibits"), list) else []
        if exhibits:
            lines.extend(["", "SUPPORTING EXHIBITS"])
            for exhibit in exhibits:
                if not isinstance(exhibit, dict):
                    continue
                exhibit_text = f"{exhibit.get('label', 'Exhibit')} - {exhibit.get('title', 'Supporting exhibit')}"
                if exhibit.get("link"):
                    exhibit_text = f"{exhibit_text} ({exhibit['link']})"
                lines.append(exhibit_text)
        lines.extend(["", str(affidavit.get("dated") or ""), str(affidavit.get("signature_line") or ""), str(affidavit.get("jurat") or "")])
        lines.extend(str(line) for line in _coerce_list(affidavit.get("notary_block")) if str(line or "").strip())
        return "\n".join(line for line in lines if line is not None)

    def _render_affidavit_docx(self, draft: Dict[str, Any], path: Path) -> None:
        from docx import Document

        document = Document()
        for line in self._render_affidavit_text(
            draft,
            draft.get("affidavit", {}) if isinstance(draft.get("affidavit"), dict) else self._build_affidavit(draft),
        ).split("\n"):
            document.add_paragraph(line)
        document.save(path)

    def _render_affidavit_pdf(self, draft: Dict[str, Any], path: Path) -> None:
        from reportlab.lib.pagesizes import LETTER
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.lib.units import inch
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

        styles = getSampleStyleSheet()
        story = []
        for line in self._render_affidavit_text(
            draft,
            draft.get("affidavit", {}) if isinstance(draft.get("affidavit"), dict) else self._build_affidavit(draft),
        ).split("\n"):
            story.append(Paragraph(escape(line or "&nbsp;"), styles["Normal"]))
            story.append(Spacer(1, 4))
        doc = SimpleDocTemplate(
            str(path),
            pagesize=LETTER,
            topMargin=inch,
            bottomMargin=inch,
            leftMargin=inch,
            rightMargin=inch,
        )
        doc.build(story)

    def _get_existing_formal_complaint(self) -> Dict[str, Any]:
        phase_manager = getattr(self.mediator, "phase_manager", None)
        if phase_manager is None:
            return {}
        existing = _safe_call(
            phase_manager,
            "get_phase_data",
            ComplaintPhase.FORMALIZATION,
            "formal_complaint",
        )
        if isinstance(existing, dict) and existing:
            return existing
        return {}

    def _derive_claim_types(
        self,
        generated_complaint: Dict[str, Any],
        classification: Dict[str, Any],
        support_claims: Dict[str, Any],
        requirements: Dict[str, Any],
    ) -> List[str]:
        claim_names = []
        claim_names.extend(_coerce_list(classification.get("claim_types")))
        claim_names.extend(list(support_claims.keys()))
        claim_names.extend(list(requirements.keys()))
        for claim in _coerce_list(generated_complaint.get("legal_claims")):
            if isinstance(claim, dict):
                claim_names.append(claim.get("title"))
        return _unique_preserving_order(claim_names) or ["General civil action"]

    def _derive_parties(
        self,
        generated_complaint: Dict[str, Any],
        *,
        plaintiff_names: Optional[List[str]],
        defendant_names: Optional[List[str]],
    ) -> tuple[List[str], List[str]]:
        parties = generated_complaint.get("parties", {}) if isinstance(generated_complaint, dict) else {}
        plaintiffs = _unique_preserving_order(
            list(plaintiff_names or []) + list(parties.get("plaintiffs", []) or [])
        ) or ["Plaintiff"]
        defendants = _unique_preserving_order(
            list(defendant_names or []) + list(parties.get("defendants", []) or [])
        ) or ["Defendant"]
        return plaintiffs, defendants

    def _derive_title(self, plaintiffs: List[str], defendants: List[str]) -> str:
        return f"{plaintiffs[0]} v. {defendants[0]}"

    def _build_court_header(
        self,
        *,
        court_name: str,
        district: str,
        county: Optional[str],
        division: Optional[str],
        override: Optional[str],
    ) -> str:
        if override:
            return override.strip().upper()
        court = str(court_name or "United States District Court").strip().upper()
        parts = [f"IN THE {court}"]
        forum_type = self._infer_forum_type(classification={}, court_name=court_name)
        county_text = self._format_county_for_header(county)
        if county_text and forum_type == "state":
            parts.append(f"FOR THE {county_text}")
        elif district:
            parts.append(f"FOR THE {str(district).strip().upper()}")
        if division:
            parts.append(str(division).strip().upper())
        return " ".join(parts)

    def _infer_forum_type(
        self,
        *,
        classification: Dict[str, Any],
        court_name: str,
    ) -> str:
        jurisdiction = str(classification.get("jurisdiction") or "").strip().lower()
        if jurisdiction in {"federal", "us", "united states"}:
            return "federal"
        if jurisdiction in {"state", "state court", "county", "local"}:
            return "state"

        court_name_text = str(court_name or "").strip().lower()
        if "united states" in court_name_text or "u.s." in court_name_text:
            return "federal"
        if any(
            marker in court_name_text
            for marker in ("superior court", "circuit court", "common pleas", "state of", "county")
        ):
            return "state"
        return "unknown"

    def _build_nature_of_action(
        self,
        *,
        claim_types: List[str],
        classification: Dict[str, Any],
        statutes: List[Dict[str, Any]],
        court_name: str,
    ) -> List[str]:
        claim_phrase = ", ".join(claim_types)
        legal_areas = ", ".join(_coerce_list(classification.get("legal_areas")))
        jurisdiction = str(classification.get("jurisdiction") or "the applicable court")
        forum_type = self._infer_forum_type(classification=classification, court_name=court_name)
        statute_refs = _unique_preserving_order(
            [s.get("citation") for s in statutes if isinstance(s, dict) and s.get("citation")]
        )
        if forum_type == "federal":
            paragraphs = [
                (
                    "This is a civil action arising under federal law and the facts disclosed during the "
                    f"complaint intake process. Plaintiff seeks relief for {claim_phrase} within {jurisdiction} jurisdiction."
                )
            ]
        elif forum_type == "state":
            paragraphs = [
                (
                    "This is a civil action brought in state court arising from the facts disclosed during "
                    f"the complaint intake process. Plaintiff seeks relief for {claim_phrase} under the governing state law."
                )
            ]
        else:
            paragraphs = [
                (
                    "This is a civil action arising from the facts disclosed during the complaint intake "
                    f"process. Plaintiff seeks relief for {claim_phrase} within {jurisdiction} jurisdiction."
                )
            ]
        if legal_areas:
            paragraphs.append(
                f"The action implicates the following areas of law: {legal_areas}."
            )
        if statute_refs:
            paragraphs.append(
                "The draft relies on the following principal legal authorities: "
                f"{', '.join(statute_refs[:5])}."
            )
        return paragraphs

    def _collect_general_facts(
        self,
        generated_complaint: Dict[str, Any],
        classification: Dict[str, Any],
        state: Any,
    ) -> List[str]:
        facts: List[str] = []
        for allegation in _coerce_list(generated_complaint.get("factual_allegations")):
            facts.extend(_extract_text_candidates(allegation))
        facts.extend(_extract_text_candidates(classification.get("key_facts")))
        for inquiry in _coerce_list(getattr(state, "inquiries", []) if state is not None else []):
            if isinstance(inquiry, dict):
                question = str(inquiry.get("question") or "").strip()
                answer = str(inquiry.get("answer") or "").strip()
                if answer:
                    if question:
                        facts.append(f"{question}: {answer}")
                    else:
                        facts.append(answer)
        complaint_text = getattr(state, "complaint", None) if state is not None else None
        original_text = getattr(state, "original_complaint", None) if state is not None else None
        if isinstance(complaint_text, dict):
            explicit_facts = _extract_text_candidates(complaint_text.get("facts"))
            if explicit_facts:
                facts.extend(explicit_facts)
            else:
                facts.extend(_extract_text_candidates(complaint_text.get("summary") or complaint_text))
        elif complaint_text:
            facts.extend(_extract_text_candidates(complaint_text))
        elif original_text:
            facts.extend(_extract_text_candidates(original_text))

        normalized = []
        for item in _unique_preserving_order(facts):
            text = re.sub(r"\s+", " ", item).strip()
            if len(text) < 12:
                continue
            normalized.append(text)
            if len(normalized) >= 12:
                break
        return normalized or ["Additional factual development is required before filing."]

    def _build_claims_for_relief(
        self,
        *,
        user_id: str,
        claim_types: List[str],
        requirements: Dict[str, Any],
        statutes: List[Dict[str, Any]],
        support_claims: Dict[str, Any],
        exhibits: List[Dict[str, Any]],
    ) -> List[Dict[str, Any]]:
        claims: List[Dict[str, Any]] = []
        for claim_type in claim_types:
            support_claim = support_claims.get(claim_type, {}) if isinstance(support_claims, dict) else {}
            overview = _safe_call(
                self.mediator,
                "get_claim_overview",
                claim_type=claim_type,
                user_id=user_id,
                required_support_kinds=["evidence", "authority"],
            ) or {}
            overview_claim = overview.get("claims", {}).get(claim_type, {}) if isinstance(overview, dict) else {}
            related_exhibits = [
                exhibit for exhibit in exhibits if not exhibit.get("claim_type") or exhibit.get("claim_type") == claim_type
            ]
            claim_facts = self._collect_claim_facts(claim_type, user_id, support_claim)
            claim_facts = self._annotate_lines_with_exhibits(claim_facts, related_exhibits)
            source_context = self._extract_support_source_context_counts(support_claim)
            claims.append(
                {
                    "claim_type": claim_type,
                    "count_title": claim_type.title(),
                    "legal_standards": self._build_claim_legal_standards(
                        claim_type=claim_type,
                        requirements=requirements,
                        statutes=statutes,
                    ),
                    "supporting_facts": claim_facts,
                    "missing_elements": self._extract_overview_elements(overview_claim.get("missing")),
                    "partially_supported_elements": self._extract_overview_elements(
                        overview_claim.get("partially_supported")
                    ),
                    "support_summary": {
                        "total_elements": support_claim.get("total_elements", 0),
                        "covered_elements": support_claim.get("covered_elements", 0),
                        "uncovered_elements": support_claim.get("uncovered_elements", 0),
                        "support_by_kind": support_claim.get("support_by_kind", {}),
                        "support_by_source": source_context["support_by_source"],
                        "source_family_counts": source_context["source_family_counts"],
                        "record_scope_counts": source_context["record_scope_counts"],
                        "artifact_family_counts": source_context["artifact_family_counts"],
                        "corpus_family_counts": source_context["corpus_family_counts"],
                        "content_origin_counts": source_context["content_origin_counts"],
                    },
                    "supporting_exhibits": [
                        {
                            "label": exhibit.get("label"),
                            "title": exhibit.get("title"),
                            "link": exhibit.get("link"),
                        }
                        for exhibit in related_exhibits[:8]
                    ],
                }
            )
        return claims

    def _extract_support_source_context_counts(self, support_claim: Dict[str, Any]) -> Dict[str, Dict[str, int]]:
        packet_summary = (
            support_claim.get("support_packet_summary", {})
            if isinstance(support_claim, dict) and isinstance(support_claim.get("support_packet_summary"), dict)
            else {}
        )

        def _normalized_counts(key: str) -> Dict[str, int]:
            primary = support_claim.get(key, {}) if isinstance(support_claim, dict) else {}
            fallback = packet_summary.get(key, {})
            source = primary if isinstance(primary, dict) and primary else fallback
            if not isinstance(source, dict):
                return {}
            counts: Dict[str, int] = {}
            for label, value in source.items():
                normalized_label = str(label or "").strip()
                if not normalized_label:
                    continue
                count = int(value or 0)
                if count <= 0:
                    continue
                counts[normalized_label] = count
            return counts

        return {
            "support_by_source": _normalized_counts("support_by_source"),
            "source_family_counts": _normalized_counts("source_family_counts"),
            "record_scope_counts": _normalized_counts("record_scope_counts"),
            "artifact_family_counts": _normalized_counts("artifact_family_counts"),
            "corpus_family_counts": _normalized_counts("corpus_family_counts"),
            "content_origin_counts": _normalized_counts("content_origin_counts"),
        }

    def _collect_claim_facts(
        self,
        claim_type: str,
        user_id: str,
        support_claim: Dict[str, Any],
    ) -> List[str]:
        facts: List[str] = []
        fact_rows = _safe_call(
            self.mediator,
            "get_claim_support_facts",
            claim_type=claim_type,
            user_id=user_id,
        ) or []
        for row in _coerce_list(fact_rows):
            facts.extend(_extract_text_candidates(row))

        for element in _coerce_list(support_claim.get("elements") if isinstance(support_claim, dict) else []):
            if not isinstance(element, dict):
                continue
            element_text = str(element.get("element_text") or element.get("claim_element") or "").strip()
            if element_text:
                facts.append(f"Element supported: {element_text}")
            for link in _coerce_list(element.get("links")):
                if isinstance(link, dict):
                    facts.extend(_extract_text_candidates(link))

        normalized = []
        for item in _unique_preserving_order(facts):
            text = re.sub(r"\s+", " ", item).strip()
            if len(text) < 10 or self._is_generic_claim_support_text(text):
                continue
            normalized.append(text)
            if len(normalized) >= 8:
                break
        chronology_support = self._build_claim_chronology_support(claim_type=claim_type, claim_name=claim_type.title())
        combined = _unique_preserving_order(chronology_support + normalized)
        return combined or [f"The intake record describes facts supporting the {claim_type} claim."]

    def _build_claim_chronology_support(self, *, claim_type: str, claim_name: str, limit: int = 2) -> List[str]:
        phase_manager = getattr(self.mediator, "phase_manager", None)
        intake_case_file = phase_manager.get_phase_data(ComplaintPhase.INTAKE, "intake_case_file") if phase_manager else None
        if not isinstance(intake_case_file, dict):
            return []

        facts = [dict(item) for item in list(intake_case_file.get("canonical_facts") or []) if isinstance(item, dict)]
        relations = [dict(item) for item in list(intake_case_file.get("timeline_relations") or []) if isinstance(item, dict)]
        if not facts or not relations:
            return []

        combined = " ".join([str(claim_type or ""), str(claim_name or "")]).strip().lower()
        if any(token in combined for token in ("retaliat", "reprisal", "protected activity")):
            focus_families = {"protected_activity", "adverse_action", "causation"}
        elif any(token in combined for token in ("due process", "grievance", "hearing", "appeal", "review", "notice")):
            focus_families = {"notice_chain", "hearing_process", "response_timeline", "adverse_action"}
        elif any(token in combined for token in ("accommodation", "disabil", "fair housing", "discrimination", "termination", "denial")):
            focus_families = {"adverse_action", "response_timeline", "notice_chain"}
        else:
            focus_families = set()

        fact_by_id = {
            str(fact.get("fact_id") or "").strip(): fact
            for fact in facts
            if str(fact.get("fact_id") or "").strip()
        }
        relation_records = []
        for relation in relations:
            if str(relation.get("relation_type") or "").strip().lower() != "before":
                continue
            source_id = str(relation.get("source_fact_id") or "").strip()
            target_id = str(relation.get("target_fact_id") or "").strip()
            source_fact = fact_by_id.get(source_id)
            target_fact = fact_by_id.get(target_id)
            if not source_fact or not target_fact:
                continue
            source_date = _format_timeline_date((source_fact.get("temporal_context") or {}).get("start_date") or relation.get("source_start_date"))
            target_date = _format_timeline_date((target_fact.get("temporal_context") or {}).get("start_date") or relation.get("target_start_date"))
            if not source_date or not target_date:
                continue
            relation_records.append(
                {
                    "key": (source_id, target_id),
                    "source_id": source_id,
                    "target_id": target_id,
                    "source_fact": source_fact,
                    "target_fact": target_fact,
                    "source_date": source_date,
                    "target_date": target_date,
                    "source_family": str(source_fact.get("predicate_family") or "").strip().lower(),
                    "target_family": str(target_fact.get("predicate_family") or "").strip().lower(),
                }
            )
        if not relation_records:
            return []

        filtered_records = [
            record for record in relation_records
            if not focus_families or ({record['source_family'], record['target_family']} & focus_families)
        ]
        if not filtered_records:
            filtered_records = relation_records

        outgoing: Dict[str, List[Dict[str, Any]]] = {}
        incoming_count: Dict[str, int] = {}
        for record in filtered_records:
            outgoing.setdefault(record["source_id"], []).append(record)
            incoming_count[record["target_id"]] = incoming_count.get(record["target_id"], 0) + 1
            incoming_count.setdefault(record["source_id"], incoming_count.get(record["source_id"], 0))

        lines: List[str] = []
        fallback_lines: List[str] = []
        seen = set()
        used_keys = set()
        for record in filtered_records:
            if len(lines) >= limit:
                break
            if record["key"] in used_keys:
                continue
            if incoming_count.get(record["source_id"], 0) != 0 or len(outgoing.get(record["source_id"], [])) != 1:
                continue
            chain = [record]
            next_id = record["target_id"]
            temp_used = {record["key"]}
            while len(outgoing.get(next_id, [])) == 1 and incoming_count.get(next_id, 0) == 1:
                next_record = outgoing[next_id][0]
                if next_record["key"] in temp_used:
                    break
                chain.append(next_record)
                temp_used.add(next_record["key"])
                next_id = next_record["target_id"]
            if len(chain) < 2:
                continue
            segments = [
                f"{_chronology_fact_label(chain[0]['source_fact']).lower()} on {chain[0]['source_date']}"
            ]
            segments.extend(
                f"{_chronology_fact_label(item['target_fact']).lower()} on {item['target_date']}"
                for item in chain
            )
            line = f"The chronology shows {_join_chronology_segments(segments)} in sequence."
            last_target = chain[-1]["target_fact"]
            target_context = last_target.get("temporal_context") if isinstance(last_target.get("temporal_context"), dict) else {}
            if target_context.get("derived_from_relative_anchor"):
                relative_markers = [str(item) for item in list(target_context.get("relative_markers") or []) if str(item)]
                if relative_markers:
                    line = line.rstrip(".") + f". The later date is derived from reported timing ({relative_markers[0]})."
            key = line.lower()
            if key in seen:
                continue
            seen.add(key)
            used_keys.update(temp_used)
            lines.append(line)

        for record in filtered_records:
            if record["key"] in used_keys:
                continue
            source_label = _chronology_fact_label(record["source_fact"])
            target_label = _chronology_fact_label(record["target_fact"]).lower()
            line = f"The chronology shows {source_label.lower()} on {record['source_date']} before {target_label} on {record['target_date']}."
            target_context = record["target_fact"].get("temporal_context") if isinstance(record["target_fact"].get("temporal_context"), dict) else {}
            if target_context.get("derived_from_relative_anchor"):
                relative_markers = [str(item) for item in list(target_context.get("relative_markers") or []) if str(item)]
                if relative_markers:
                    line = line.rstrip(".") + f". The later date is derived from reported timing ({relative_markers[0]})."
            key = line.lower()
            if key in seen:
                continue
            seen.add(key)
            fallback_lines.append(line)
        if lines:
            return lines
        if fallback_lines:
            return fallback_lines[:1]
        return []

    def _claim_temporal_gap_focus(self, claim_type: str, claim_name: str) -> Dict[str, set[str]]:
        combined = " ".join([str(claim_type or ""), str(claim_name or "")]).strip().lower()
        issue_families = {"timeline"}
        element_tags = {"timeline"}
        objectives = {"timeline", "exact_dates"}
        if any(token in combined for token in ("retaliat", "reprisal", "protected activity")):
            issue_families.update({"causation", "protected_activity", "adverse_action"})
            element_tags.update({"causation", "protected_activity", "adverse_action"})
            objectives.update({"causation_link", "causation_sequence", "anchor_adverse_action"})
        if any(token in combined for token in ("due process", "grievance", "hearing", "appeal", "review", "notice")):
            issue_families.update({"notice_chain", "hearing_process", "response_timeline"})
            element_tags.update({"notice", "hearing", "appeal", "response", "review"})
            objectives.update({"anchor_appeal_rights", "hearing_request_timing", "response_dates"})
        if any(token in combined for token in ("accommodation", "disabil", "discrimination", "termination", "denial")):
            issue_families.update({"adverse_action", "notice_chain", "response_timeline"})
            element_tags.update({"adverse_action", "response", "notice"})
            objectives.update({"response_dates", "anchor_adverse_action"})
        return {
            "issue_families": issue_families,
            "element_tags": element_tags,
            "objectives": objectives,
        }

    def _build_claim_temporal_gap_hints(self, claim_type: str, claim_name: str, *, limit: int = 3) -> List[str]:
        phase_manager = getattr(self.mediator, "phase_manager", None)
        intake_case_file = phase_manager.get_phase_data(ComplaintPhase.INTAKE, "intake_case_file") if phase_manager else None
        if not isinstance(intake_case_file, dict):
            return []

        focus = self._claim_temporal_gap_focus(claim_type, claim_name)
        normalized_claim_type = str(claim_type or "").strip().lower()
        hints: List[str] = []

        for issue in _coerce_list(intake_case_file.get("temporal_issue_registry")):
            if not isinstance(issue, dict):
                continue
            status = str(issue.get("status") or "open").strip().lower()
            if status not in {"open", "blocking", "warning"}:
                continue
            issue_claim_types = {str(item).strip().lower() for item in _coerce_list(issue.get("claim_types")) if str(item).strip()}
            issue_element_tags = {str(item).strip().lower() for item in _coerce_list(issue.get("element_tags")) if str(item).strip()}
            if issue_claim_types:
                if normalized_claim_type not in issue_claim_types:
                    continue
            elif issue_element_tags and not (issue_element_tags & focus["element_tags"]):
                continue
            summary = str(issue.get("summary") or "").strip()
            if summary:
                hints.append(f"Chronology gap: {summary}")

        blocker_summary = intake_case_file.get("blocker_follow_up_summary") if isinstance(intake_case_file.get("blocker_follow_up_summary"), dict) else {}
        for blocker in _coerce_list(blocker_summary.get("blocking_items")):
            if not isinstance(blocker, dict):
                continue
            issue_family = str(blocker.get("issue_family") or "").strip().lower()
            primary_objective = str(blocker.get("primary_objective") or "").strip().lower()
            blocker_objectives = {str(item).strip().lower() for item in _coerce_list(blocker.get("blocker_objectives")) if str(item).strip()}
            matched_objectives = ({primary_objective} if primary_objective else set()) | blocker_objectives
            if issue_family:
                if issue_family not in focus["issue_families"] and not (matched_objectives & focus["objectives"]):
                    continue
            elif not (matched_objectives & focus["objectives"]):
                continue
            reason = str(blocker.get("reason") or "").strip()
            if reason:
                hints.append(f"Chronology gap: {reason}")

        return _unique_preserving_order(hints)[:limit]

    def _annotate_claim_temporal_gap_hints(self, draft: Dict[str, Any]) -> None:
        claims = draft.get("claims_for_relief") if isinstance(draft.get("claims_for_relief"), list) else []
        for claim in claims:
            if not isinstance(claim, dict):
                continue
            claim_type = str(claim.get("claim_type") or claim.get("count_title") or "").strip()
            claim_name = str(claim.get("count_title") or claim.get("claim_type") or "").strip()
            hints = self._build_claim_temporal_gap_hints(claim_type, claim_name)
            if not hints:
                continue
            claim["missing_elements"] = _unique_preserving_order(
                list(_coerce_list(claim.get("missing_elements"))) + hints
            )
            support_summary = claim.get("support_summary") if isinstance(claim.get("support_summary"), dict) else {}
            claim["support_summary"] = {
                **support_summary,
                "temporal_gap_hint_count": len(hints),
            }

    def _build_claim_legal_standards(
        self,
        *,
        claim_type: str,
        requirements: Dict[str, Any],
        statutes: List[Dict[str, Any]],
    ) -> List[str]:
        standards = _unique_preserving_order(_extract_text_candidates(requirements.get(claim_type, [])))
        related_statutes = self._select_statutes_for_claim(claim_type, statutes)
        for statute in related_statutes:
            citation = statute.get("citation")
            title = statute.get("title")
            relevance = statute.get("relevance")
            parts = [part for part in [citation, title, relevance] if part]
            if parts:
                standards.append(" - ".join(parts))
        return standards or [f"Plaintiff must prove the elements of {claim_type} under the applicable law."]

    def _build_legal_standards_summary(
        self,
        *,
        statutes: List[Dict[str, Any]],
        requirements: Dict[str, Any],
    ) -> List[str]:
        summary = []
        for claim_type, elements in requirements.items():
            summary.append(
                f"{claim_type.title()}: {', '.join(_unique_preserving_order(_extract_text_candidates(elements))[:4])}"
            )
        for statute in statutes[:5]:
            if isinstance(statute, dict):
                parts = [statute.get("citation"), statute.get("title"), statute.get("relevance")]
                text = " - ".join([part for part in parts if part])
                if text:
                    summary.append(text)
        return _unique_preserving_order(summary)

    def _safe_mediator_dict(self, method_name: str, **kwargs: Any) -> Dict[str, Any]:
        method = getattr(self.mediator, method_name, None)
        if not callable(method):
            return {}
        try:
            result = method(**kwargs)
        except Exception:
            return {}
        return result if isinstance(result, dict) else {}

    def _extract_blocker_follow_up_signals(self, optimization_report: Dict[str, Any]) -> Dict[str, Any]:
        report = optimization_report if isinstance(optimization_report, dict) else {}
        intake_case_summary = (
            report.get("intake_case_summary")
            if isinstance(report.get("intake_case_summary"), dict)
            else build_intake_case_review_summary(self.mediator)
        )
        intake_case_summary = intake_case_summary if isinstance(intake_case_summary, dict) else {}
        blocker_follow_up_summary = (
            intake_case_summary.get("blocker_follow_up_summary")
            if isinstance(intake_case_summary.get("blocker_follow_up_summary"), dict)
            else {}
        )
        blocker_items = [
            dict(item)
            for item in list(blocker_follow_up_summary.get("blocking_items") or [])
            if isinstance(item, dict)
        ]
        open_items = [
            dict(item)
            for item in list(intake_case_summary.get("open_items") or [])
            if isinstance(item, dict) and str(item.get("kind") or "").strip().lower() == "blocker_follow_up"
        ]
        combined_items = blocker_items + open_items
        issue_types = _dedupe_text_values(
            str(item.get("issue_type") or item.get("type") or item.get("gap_type") or "").strip().lower()
            for item in combined_items
        )
        extraction_targets = _dedupe_text_values(
            list(blocker_follow_up_summary.get("extraction_targets") or [])
            + [
                target
                for item in combined_items
                for target in list(item.get("extraction_targets") or [])
            ]
        )
        workflow_phases = _dedupe_text_values(
            list(blocker_follow_up_summary.get("workflow_phases") or [])
            + [str(item.get("workflow_phase") or "").strip() for item in combined_items]
        )
        follow_up_questions = _dedupe_text_values(
            str(item.get("suggested_question") or item.get("next_question_template") or "").strip()
            for item in combined_items
        )

        def _item_has_confirmation_placeholder(item: Dict[str, Any]) -> bool:
            return any(
                _contains_confirmation_placeholder(value)
                for value in (
                    item.get("summary"),
                    item.get("reason"),
                    item.get("suggested_question"),
                    item.get("next_question_template"),
                    item.get("description"),
                )
            )

        placeholder_items = [
            item
            for item in combined_items
            if "confirmation_placeholder" in str(item.get("issue_type") or "").strip().lower()
            or _item_has_confirmation_placeholder(item)
        ]
        decision_maker_items = [
            item
            for item in combined_items
            if "decision_maker" in str(item.get("issue_type") or "").strip().lower()
            or any(
                target in {"decision_maker", "actor_name", "actor_role"}
                for target in list(item.get("extraction_targets") or [])
            )
        ]
        causation_items = [
            item
            for item in combined_items
            if (
                "causation" in str(item.get("issue_type") or "").strip().lower()
                or "retaliation_missing_sequence" in str(item.get("issue_type") or "").strip().lower()
                or "retaliation_missing_decision_maker" in str(item.get("issue_type") or "").strip().lower()
                or "causation_link" in list(item.get("extraction_targets") or [])
            )
        ]
        document_anchor_items = [
            item
            for item in combined_items
            if any(
                target in {"document_type", "document_date", "document_owner", "evidence_record", "verification_source"}
                for target in list(item.get("extraction_targets") or [])
            )
            or "document_anchor" in str(item.get("issue_type") or "").strip().lower()
        ]

        return {
            "blocker_count": int(blocker_follow_up_summary.get("blocking_item_count") or len(blocker_items) or 0),
            "issue_types": issue_types,
            "extraction_targets": extraction_targets,
            "workflow_phases": workflow_phases,
            "follow_up_questions": follow_up_questions,
            "confirmation_placeholder_count": len(placeholder_items),
            "decision_maker_probe_count": len(decision_maker_items),
            "causation_probe_count": len(causation_items),
            "document_anchor_probe_count": len(document_anchor_items),
            "needs_confirmation_follow_up": bool(placeholder_items),
            "needs_decision_maker_follow_up": bool(decision_maker_items),
            "needs_causation_follow_up": bool(causation_items),
            "needs_document_anchor_follow_up": bool(document_anchor_items),
        }

    def _extract_actor_critic_priority_metrics(self, optimization_report: Dict[str, Any]) -> Dict[str, float]:
        if not isinstance(optimization_report, dict):
            return {}
        final_review = optimization_report.get("final_review") if isinstance(optimization_report.get("final_review"), dict) else {}
        dimension_scores = final_review.get("dimension_scores") if isinstance(final_review.get("dimension_scores"), dict) else {}
        adversarial_batch = (
            optimization_report.get("adversarial_batch")
            if isinstance(optimization_report.get("adversarial_batch"), dict)
            else {}
        )
        metric_candidates = [
            optimization_report.get("actor_critic_metrics"),
            optimization_report.get("adversarial_batch_metrics"),
            optimization_report.get("latest_adversarial_batch_metrics"),
            optimization_report.get("priority_metrics"),
            optimization_report.get("baseline_metrics"),
            optimization_report.get("metrics"),
            adversarial_batch.get("metrics") if isinstance(adversarial_batch, dict) else {},
            (optimization_report.get("actor_critic_optimizer") or {}).get("metrics")
            if isinstance(optimization_report.get("actor_critic_optimizer"), dict)
            else {},
        ]
        normalized: Dict[str, float] = {}
        aliases = {
            "empathy": "empathy",
            "empathy_avg": "empathy",
            "avg_empathy": "empathy",
            "question_quality": "question_quality",
            "question_quality_avg": "question_quality",
            "avg_question_quality": "question_quality",
            "question_quality_score": "question_quality",
            "information_extraction": "information_extraction",
            "information_extraction_avg": "information_extraction",
            "avg_information_extraction": "information_extraction",
            "coverage": "coverage",
            "coverage_avg": "coverage",
            "avg_coverage": "coverage",
            "efficiency": "efficiency",
            "efficiency_avg": "efficiency",
            "avg_efficiency": "efficiency",
        }
        for candidate in metric_candidates:
            if not isinstance(candidate, dict):
                continue
            for key, value in candidate.items():
                metric_name = aliases.get(str(key).strip().lower())
                if not metric_name:
                    continue
                normalized[metric_name] = max(0.0, min(1.0, _safe_float(value, DEFAULT_ACTOR_CRITIC_PRIORITY_METRICS.get(metric_name, 0.0))))
        if "question_quality" not in normalized:
            normalized["question_quality"] = max(
                _safe_float(dimension_scores.get("coherence"), DEFAULT_ACTOR_CRITIC_PRIORITY_METRICS["question_quality"]),
                DEFAULT_ACTOR_CRITIC_PRIORITY_METRICS["question_quality"],
            )
        if "coverage" not in normalized:
            normalized["coverage"] = max(
                _safe_float(dimension_scores.get("completeness"), DEFAULT_ACTOR_CRITIC_PRIORITY_METRICS["coverage"]),
                DEFAULT_ACTOR_CRITIC_PRIORITY_METRICS["coverage"],
            )
        if "information_extraction" not in normalized:
            normalized["information_extraction"] = max(
                _safe_float(dimension_scores.get("grounding"), DEFAULT_ACTOR_CRITIC_PRIORITY_METRICS["information_extraction"]),
                DEFAULT_ACTOR_CRITIC_PRIORITY_METRICS["information_extraction"],
            )
        if "efficiency" not in normalized:
            normalized["efficiency"] = max(
                _safe_float(dimension_scores.get("procedural"), DEFAULT_ACTOR_CRITIC_PRIORITY_METRICS["efficiency"]),
                DEFAULT_ACTOR_CRITIC_PRIORITY_METRICS["efficiency"],
            )
        if "empathy" not in normalized:
            normalized["empathy"] = DEFAULT_ACTOR_CRITIC_PRIORITY_METRICS["empathy"]
        return normalized

    def _extract_actor_critic_priority_value(self, optimization_report: Dict[str, Any]) -> int:
        report = optimization_report if isinstance(optimization_report, dict) else {}
        priority_candidates = [
            report.get("priority"),
            report.get("actor_critic_priority"),
            (report.get("actor_critic_optimizer") or {}).get("priority")
            if isinstance(report.get("actor_critic_optimizer"), dict)
            else None,
        ]
        for candidate in priority_candidates:
            if candidate is None:
                continue
            try:
                return max(1, min(100, int(candidate)))
            except Exception:
                continue
        return 70

    def _resolve_router_backed_question_quality(self, optimization_report: Dict[str, Any]) -> bool:
        report = optimization_report if isinstance(optimization_report, dict) else {}
        router_status = dict(report.get("router_status") or {})
        router_usage = dict(report.get("router_usage") or {})
        actor_critic_optimizer = (
            report.get("actor_critic_optimizer")
            if isinstance(report.get("actor_critic_optimizer"), dict)
            else {}
        )
        adversarial_batch = report.get("adversarial_batch") if isinstance(report.get("adversarial_batch"), dict) else {}
        candidate_flags = [
            str(router_status.get("llm_router") or "").strip().lower() == "available",
            _coerce_bool(router_status.get("llm_router_available"), default=False),
            _coerce_bool(router_status.get("available"), default=False),
            _coerce_bool(router_usage.get("llm_router_available"), default=False),
            _coerce_bool(router_usage.get("router_backed_question_quality"), default=False),
            _coerce_bool(report.get("router_backed_question_quality"), default=False),
            _coerce_bool(actor_critic_optimizer.get("router_backed_question_quality"), default=False),
            _coerce_bool(adversarial_batch.get("router_backed_question_quality"), default=False),
        ]
        return any(candidate_flags)

    def _extract_actor_critic_guidance(self, optimization_report: Dict[str, Any]) -> Dict[str, Any]:
        report = optimization_report if isinstance(optimization_report, dict) else {}
        optimization_method = str(report.get("optimization_method") or "").strip().lower()
        method_name = str(report.get("method") or "").strip().lower()
        if "actor_critic" not in optimization_method and "actor_critic" not in method_name:
            return {}

        provided_order = [
            str(item).strip()
            for item in list(report.get("phase_focus_order") or [])
            if str(item).strip()
        ]
        phase_focus_order = [name for name in ACTOR_CRITIC_PHASE_FOCUS_ORDER if name in provided_order] or list(ACTOR_CRITIC_PHASE_FOCUS_ORDER)
        phase_focus_order.extend(name for name in provided_order if name not in phase_focus_order)
        final_review = report.get("final_review") if isinstance(report.get("final_review"), dict) else {}
        section_scores = final_review.get("section_scores") if isinstance(final_review.get("section_scores"), dict) else {}

        metrics = self._extract_actor_critic_priority_metrics(report)
        priority = self._extract_actor_critic_priority_value(report)
        intake_score = _safe_float(section_scores.get("intake_questioning"), 0.0)
        priority_findings = _dedupe_text_values(
            _extract_latest_adversarial_priority_findings(
                {
                    "priorities": report.get("priorities"),
                    "priority_findings": report.get("priority_findings"),
                    "latest_adversarial_batch": report.get("latest_adversarial_batch"),
                    "adversarial_batch": report.get("adversarial_batch"),
                    "latest_batch_priorities": report.get("latest_batch_priorities"),
                }
            )
        )
        if intake_score < 0.85 or priority >= 70:
            forced_order = [name for name in ACTOR_CRITIC_PHASE_FOCUS_ORDER if name in phase_focus_order]
            forced_order.extend(name for name in phase_focus_order if name not in forced_order)
            phase_focus_order = forced_order

        return {
            "optimization_method": optimization_method or "actor_critic",
            "phase_focus_order": phase_focus_order,
            "priority": priority,
            "metrics": metrics,
            "intake_questioning_score": intake_score,
            "router_backed_question_quality": self._resolve_router_backed_question_quality(report),
            "priority_findings": priority_findings,
            "needs_chronology_closure": _has_chronology_gap_priority(priority_findings),
            "needs_decision_document_precision": _has_decision_or_document_precision_priority(priority_findings),
        }

    def _extract_adversarial_session_flow_signals(self, optimization_report: Dict[str, Any]) -> Dict[str, Any]:
        report = optimization_report if isinstance(optimization_report, dict) else {}
        if not report:
            return {
                "available": False,
                "successful_session_count": 0,
                "session_count": 0,
                "assessment_blocked": False,
            }

        adversarial_batch = report.get("adversarial_batch") if isinstance(report.get("adversarial_batch"), dict) else {}
        actor_critic_optimizer = (
            report.get("actor_critic_optimizer")
            if isinstance(report.get("actor_critic_optimizer"), dict)
            else {}
        )
        review_metadata = report.get("review_metadata") if isinstance(report.get("review_metadata"), dict) else {}

        def _first_count(candidates: List[Any], *, default: int = 0) -> int:
            for candidate in candidates:
                if candidate is None:
                    continue
                try:
                    if isinstance(candidate, list):
                        return max(0, len(candidate))
                    return max(0, int(candidate))
                except Exception:
                    continue
            return default

        successful_session_count = _first_count(
            [
                report.get("successful_session_count"),
                report.get("successful_sessions"),
                adversarial_batch.get("successful_session_count"),
                adversarial_batch.get("successful_sessions"),
                actor_critic_optimizer.get("successful_session_count"),
                actor_critic_optimizer.get("successful_sessions"),
                review_metadata.get("successful_session_count"),
                review_metadata.get("successful_sessions"),
                report.get("accepted_iterations"),
            ]
        )
        session_count = _first_count(
            [
                report.get("session_count"),
                report.get("total_session_count"),
                report.get("adversarial_session_count"),
                adversarial_batch.get("session_count"),
                adversarial_batch.get("total_session_count"),
                adversarial_batch.get("adversarial_session_count"),
                adversarial_batch.get("sessions"),
                actor_critic_optimizer.get("session_count"),
                actor_critic_optimizer.get("sessions"),
                report.get("section_history"),
            ]
        )
        has_adversarial_evidence = any(
            (
                bool(adversarial_batch),
                bool(report.get("latest_adversarial_batch")),
                bool(report.get("latest_batch_priorities")),
                bool(report.get("priority_findings")),
                bool(report.get("adversarial_batch_metrics")),
            )
        )
        available = bool(report)
        assessment_blocked = available and successful_session_count <= 0 and (
            session_count > 0 or has_adversarial_evidence
        )
        return {
            "available": available,
            "successful_session_count": successful_session_count,
            "session_count": session_count,
            "assessment_blocked": assessment_blocked,
        }

    def _refresh_drafting_readiness_workflow_warnings(
        self,
        *,
        drafting_readiness: Dict[str, Any],
        workflow_phase_plan: Dict[str, Any],
    ) -> None:
        if not isinstance(drafting_readiness, dict):
            return
        all_existing_warnings = [
            dict(item) for item in list(drafting_readiness.get("warnings") or []) if isinstance(item, dict)
        ]
        prior_workflow_warning_count = sum(
            1
            for item in all_existing_warnings
            if str(item.get("code") or "").strip().lower().startswith("workflow_")
        )
        existing_warnings = [
            item
            for item in all_existing_warnings
            if not str(item.get("code") or "").strip().lower().startswith("workflow_")
        ]
        workflow_warnings = self._build_workflow_phase_warning_entries(workflow_phase_plan)
        combined_warnings = existing_warnings + workflow_warnings
        base_warning_count = max(
            0,
            int(drafting_readiness.get("warning_count") or 0) - prior_workflow_warning_count,
        )
        if combined_warnings:
            drafting_readiness["warnings"] = combined_warnings
            drafting_readiness["warning_count"] = base_warning_count + len(workflow_warnings)
            for warning in workflow_warnings:
                drafting_readiness["status"] = _merge_status(
                    str(drafting_readiness.get("status") or "ready"),
                    str(warning.get("severity") or "ready"),
                )
        elif "warnings" in drafting_readiness:
            drafting_readiness.pop("warnings", None)
            drafting_readiness["warning_count"] = base_warning_count

    def _build_runtime_workflow_phase_plan(
        self,
        *,
        drafting_readiness: Dict[str, Any],
        document_optimization: Optional[Dict[str, Any]] = None,
    ) -> Dict[str, Any]:
        phase_manager = getattr(self.mediator, "phase_manager", None)
        optimization_report = document_optimization if isinstance(document_optimization, dict) else {}
        if phase_manager is None and not optimization_report:
            return {}

        phases: Dict[str, Dict[str, Any]] = {}
        actor_critic_guidance = self._extract_actor_critic_guidance(optimization_report)
        adversarial_flow_signals = self._extract_adversarial_session_flow_signals(optimization_report)
        graph_phase = self._build_graph_analysis_phase_guidance(
            phase_manager,
            document_optimization=optimization_report,
        )
        if graph_phase:
            phases["graph_analysis"] = graph_phase

        document_phase = self._build_document_generation_phase_guidance(
            drafting_readiness=drafting_readiness,
            document_optimization=optimization_report,
        )
        graph_status = str(graph_phase.get("status") or "ready").strip().lower() if graph_phase else "ready"
        graph_signals = (
            dict(graph_phase.get("signals") or {})
            if isinstance(graph_phase.get("signals"), dict)
            else {}
        )
        readiness_graph_signals = (
            dict(drafting_readiness.get("graph_completeness_signals") or {})
            if isinstance(drafting_readiness.get("graph_completeness_signals"), dict)
            else {}
        )
        unresolved_factual_gaps = [
            str(item).strip()
            for item in list(drafting_readiness.get("unresolved_factual_gaps") or [])
            if str(item).strip()
        ]
        unresolved_legal_gaps = [
            str(item).strip()
            for item in list(drafting_readiness.get("unresolved_legal_gaps") or [])
            if str(item).strip()
        ]
        uncovered_intake_objectives = [
            str(item).strip()
            for item in list(drafting_readiness.get("uncovered_intake_objectives") or [])
            if str(item).strip()
        ]
        missing_required_intake_objectives = [
            str(item).strip()
            for item in list(drafting_readiness.get("missing_required_intake_objectives") or [])
            if str(item).strip()
        ]
        structured_handoff_signals = (
            dict(drafting_readiness.get("structured_intake_handoff_signals") or {})
            if isinstance(drafting_readiness.get("structured_intake_handoff_signals"), dict)
            else {}
        )
        structured_handoff_gap_count = int(structured_handoff_signals.get("gap_count", 0) or 0)
        weak_complaint_types = [
            str(item).strip()
            for item in list(drafting_readiness.get("weak_complaint_types") or [])
            if str(item).strip()
        ]
        weak_evidence_modalities = [
            str(item).strip()
            for item in list(drafting_readiness.get("weak_evidence_modalities") or [])
            if str(item).strip()
        ]
        targeted_weak_complaint_types = [
            item
            for item in weak_complaint_types
            if item.lower() in {"housing_discrimination", "hacc_research_engine"}
        ]
        targeted_weak_evidence_modalities = [
            item
            for item in weak_evidence_modalities
            if item.lower() in {"policy_document", "file_evidence"}
        ]
        graph_remaining_gap_count = max(
            int(graph_signals.get("remaining_gap_count", 0) or 0),
            int(graph_signals.get("current_gap_count", 0) or 0),
            int(readiness_graph_signals.get("remaining_gap_count", 0) or 0),
            int(readiness_graph_signals.get("current_gap_count", 0) or 0),
        )
        graph_gate_active = (
            graph_status != "ready"
            or graph_remaining_gap_count > 0
            or not _coerce_bool(
                readiness_graph_signals.get("knowledge_graph_available", graph_signals.get("knowledge_graph_available", True)),
                default=True,
            )
            or not _coerce_bool(
                readiness_graph_signals.get("dependency_graph_available", graph_signals.get("dependency_graph_available", True)),
                default=True,
            )
            )
        if document_phase:
            updated_document_phase = dict(document_phase)
            if graph_gate_active:
                gate_status = "blocked"
                updated_document_phase["status"] = _merge_status(
                    str(updated_document_phase.get("status") or "ready"),
                    gate_status,
                )
                summary = str(updated_document_phase.get("summary") or "").strip()
                gate_summary = (
                    "Document generation is gated on graph completeness and should not be treated as final until graph blockers are resolved."
                )
                updated_document_phase["summary"] = f"{summary} {gate_summary}".strip()
                actions = [str(item).strip() for item in list(updated_document_phase.get("recommended_actions") or []) if str(item).strip()]
                actions.append(
                    "Resolve graph completeness blockers (knowledge/dependency graph availability and unresolved graph gaps) before formalization."
                )
                if unresolved_factual_gaps:
                    actions.append(
                        "Close unresolved factual gaps before formalization: "
                        + "; ".join(unresolved_factual_gaps[:3])
                    )
                if unresolved_legal_gaps:
                    actions.append(
                        "Close unresolved legal gaps before formalization: "
                        + "; ".join(unresolved_legal_gaps[:3])
                    )
                if uncovered_intake_objectives or missing_required_intake_objectives:
                    actions.append(
                        "Resolve uncovered intake objectives before formalization: "
                        + ", ".join((missing_required_intake_objectives or uncovered_intake_objectives)[:4])
                    )
                if structured_handoff_gap_count > 0:
                    actions.append(
                        "Promote structured intake facts, date/actor anchors, and evidence references directly into summary-of-facts and claim-support paragraphs before document optimization runs."
                    )
                if targeted_weak_complaint_types:
                    actions.append(
                        "Generalize drafting quality for weak complaint types before formalization: "
                        + ", ".join(targeted_weak_complaint_types)
                    )
                if targeted_weak_evidence_modalities:
                    actions.append(
                        "Strengthen allegation support for weak evidence modalities before formalization: "
                        + ", ".join(targeted_weak_evidence_modalities)
                    )
                updated_document_phase["recommended_actions"] = _dedupe_text_values(actions)
            if bool(adversarial_flow_signals.get("assessment_blocked")):
                gate_status = "blocked" if str(drafting_readiness.get("phase_status") or "").strip().lower() == "critical" else "warning"
                updated_document_phase["status"] = _merge_status(
                    str(updated_document_phase.get("status") or "ready"),
                    gate_status,
                )
                summary = str(updated_document_phase.get("summary") or "").strip()
                session_summary = (
                    "No successful adversarial sessions were available to assess drafting handoff quality."
                )
                updated_document_phase["summary"] = f"{summary} {session_summary}".strip()
                actions = [str(item).strip() for item in list(updated_document_phase.get("recommended_actions") or []) if str(item).strip()]
                actions.append(
                    "Restore a stable adversarial session flow before tuning document-generation handoffs."
                )
                updated_document_phase["recommended_actions"] = _dedupe_text_values(actions)
            signals = dict(updated_document_phase.get("signals") or {})
            signals["gate_on_graph_completeness"] = bool(graph_gate_active)
            signals["graph_phase_status"] = graph_status
            signals["graph_remaining_gap_count"] = int(graph_remaining_gap_count)
            signals["drafting_coverage"] = _safe_float(drafting_readiness.get("coverage"), 0.0)
            signals["unresolved_factual_gap_count"] = len(unresolved_factual_gaps)
            signals["unresolved_legal_gap_count"] = len(unresolved_legal_gaps)
            signals["uncovered_intake_objective_count"] = len(uncovered_intake_objectives)
            signals["missing_required_intake_objective_count"] = len(missing_required_intake_objectives)
            signals["structured_intake_handoff_gap_count"] = int(structured_handoff_gap_count)
            signals["uncovered_intake_objectives"] = uncovered_intake_objectives[:8]
            signals["unresolved_factual_gaps"] = unresolved_factual_gaps[:6]
            signals["unresolved_legal_gaps"] = unresolved_legal_gaps[:6]
            signals["weak_complaint_types"] = weak_complaint_types
            signals["weak_evidence_modalities"] = weak_evidence_modalities
            signals["targeted_weak_complaint_types"] = targeted_weak_complaint_types
            signals["targeted_weak_evidence_modalities"] = targeted_weak_evidence_modalities
            signals["ready_for_formalization"] = not bool(
                graph_gate_active
                or unresolved_factual_gaps
                or unresolved_legal_gaps
                or uncovered_intake_objectives
                or missing_required_intake_objectives
                or structured_handoff_gap_count > 0
                or targeted_weak_complaint_types
                or targeted_weak_evidence_modalities
            )
            signals["adversarial_session_flow_available"] = bool(adversarial_flow_signals.get("available"))
            signals["adversarial_session_count"] = int(adversarial_flow_signals.get("session_count") or 0)
            signals["adversarial_successful_session_count"] = int(
                adversarial_flow_signals.get("successful_session_count") or 0
            )
            signals["adversarial_session_flow_stable"] = not bool(adversarial_flow_signals.get("assessment_blocked"))
            updated_document_phase["signals"] = signals
            phases["document_generation"] = updated_document_phase
        intake_phase = self._build_intake_questioning_phase_guidance(
            drafting_readiness=drafting_readiness,
            document_optimization=optimization_report,
        )
        if intake_phase:
            phases["intake_questioning"] = intake_phase

        plan = build_workflow_phase_plan(phases)
        if not plan:
            return {}
        preferred_order = (
            list(actor_critic_guidance.get("phase_focus_order") or [])
            if actor_critic_guidance
            else ["graph_analysis", "document_generation", "intake_questioning"]
        )
        if bool(adversarial_flow_signals.get("assessment_blocked")):
            preferred_order = ["graph_analysis", "intake_questioning", "document_generation"]
        ordered = [name for name in preferred_order if name in phases]
        ordered.extend(
            name
            for name in list(plan.get("recommended_order") or [])
            if name in phases and name not in ordered
        )
        plan["recommended_order"] = ordered
        if actor_critic_guidance:
            plan["actor_critic_guidance"] = actor_critic_guidance
        return plan

    def _build_graph_analysis_phase_guidance(
        self,
        phase_manager: Any,
        *,
        document_optimization: Dict[str, Any],
    ) -> Dict[str, Any]:
        phase = build_graph_analysis_phase_guidance(phase_manager, audience="drafting")
        if not phase:
            return {}
        updated = dict(phase)
        signals = dict(updated.get("signals") or {})
        optimization_report = document_optimization if isinstance(document_optimization, dict) else {}
        temporal_handoff = (
            optimization_report.get("claim_support_temporal_handoff")
            if isinstance(optimization_report.get("claim_support_temporal_handoff"), dict)
            else {}
        )
        unresolved_temporal_count = int(temporal_handoff.get("unresolved_temporal_issue_count", 0) or 0)
        chronology_tasks = int(temporal_handoff.get("chronology_task_count", 0) or 0)
        actor_critic_guidance = self._extract_actor_critic_guidance(optimization_report)
        blocker_signals = self._extract_blocker_follow_up_signals(optimization_report)
        needs_chronology_closure = bool(actor_critic_guidance.get("needs_chronology_closure")) or bool(
            blocker_signals.get("needs_causation_follow_up")
        )
        if unresolved_temporal_count > 0 or chronology_tasks > 0:
            updated["status"] = "warning" if str(updated.get("status") or "").lower() == "ready" else updated.get("status")
            summary = str(updated.get("summary") or "").strip()
            suffix = (
                f" Temporal graph alignment still has {unresolved_temporal_count} unresolved chronology issue(s) "
                f"across {chronology_tasks} chronology task(s)."
            )
            updated["summary"] = f"{summary}{suffix}".strip()
            actions = [str(item) for item in list(updated.get("recommended_actions") or []) if str(item).strip()]
            actions.append(
                "Resolve chronology edges for protected activity, hearing/review requests, response dates, and adverse-action outcomes before finalizing the complaint timeline."
            )
            updated["recommended_actions"] = _dedupe_text_values(actions)
        if needs_chronology_closure:
            updated["status"] = "warning" if str(updated.get("status") or "").lower() == "ready" else updated.get("status")
            actions = [str(item) for item in list(updated.get("recommended_actions") or []) if str(item).strip()]
            actions.append(
                "Close critical chronology gaps by confirming exact date anchors for protected activity, each notice/response event, and each adverse-action step in sequence."
            )
            updated["recommended_actions"] = _dedupe_text_values(actions)
        signals["unresolved_temporal_issue_count"] = unresolved_temporal_count
        signals["chronology_task_count"] = chronology_tasks
        signals["needs_chronology_closure"] = needs_chronology_closure
        signals["decision_maker_probe_count"] = int(blocker_signals.get("decision_maker_probe_count", 0) or 0)
        signals["document_anchor_probe_count"] = int(blocker_signals.get("document_anchor_probe_count", 0) or 0)
        updated["signals"] = signals
        return updated

    def _build_document_generation_phase_guidance(
        self,
        *,
        drafting_readiness: Dict[str, Any],
        document_optimization: Dict[str, Any],
    ) -> Dict[str, Any]:
        phase = build_drafting_document_generation_phase_guidance(
            drafting_readiness=drafting_readiness,
            document_optimization=document_optimization,
        )
        if not phase:
            return {}
        updated = dict(phase)
        optimization_report = document_optimization if isinstance(document_optimization, dict) else {}
        actor_critic_guidance = self._extract_actor_critic_guidance(optimization_report)
        blocker_signals = self._extract_blocker_follow_up_signals(optimization_report)
        final_review = optimization_report.get("final_review") if isinstance(optimization_report.get("final_review"), dict) else {}
        section_scores = final_review.get("section_scores") if isinstance(final_review.get("section_scores"), dict) else {}
        intake_score = float(section_scores.get("intake_questioning") or 0.0)
        if intake_score < 0.8:
            updated["status"] = "warning" if str(updated.get("status") or "").lower() == "ready" else updated.get("status")
            summary = str(updated.get("summary") or "").strip()
            updated["summary"] = (
                f"{summary} Document generation should preserve patchability while improving fact sequencing for retaliation causation and adverse-action chronology."
            ).strip()
            actions = [str(item) for item in list(updated.get("recommended_actions") or []) if str(item).strip()]
            actions.append(
                "Strengthen factual paragraphs so each adverse action is paired with exact date anchors, named/titled staff actors, hearing-request timing, response dates, and causation sequencing."
            )
            actions.append(
                "Preserve patchability by keeping one material fact per sentence and explicitly anchoring each sentence to dates, actors, and response events extracted during intake questioning."
            )
            updated["recommended_actions"] = _dedupe_text_values(actions)
        if actor_critic_guidance and bool(actor_critic_guidance.get("router_backed_question_quality")):
            actions = [str(item) for item in list(updated.get("recommended_actions") or []) if str(item).strip()]
            actions.append(
                "Use router-backed drafting passes to convert unresolved intake objectives into concrete chronology-aligned allegations and claim-support paragraphs."
            )
            actions.append(
                "Preserve patchability by emitting single-sentence allegation units with explicit date, actor/title, and source anchors for each router-refined paragraph."
            )
            updated["recommended_actions"] = _dedupe_text_values(actions)
        if bool(actor_critic_guidance.get("needs_decision_document_precision")) or bool(
            blocker_signals.get("needs_document_anchor_follow_up")
        ):
            actions = [str(item) for item in list(updated.get("recommended_actions") or []) if str(item).strip()]
            actions.append(
                "Increase precision in adverse-action allegations by naming the decision-maker (or known title), the specific decision communicated, and the controlling documentary artifact for each step."
            )
            actions.append(
                "Keep each allegation patchable by isolating one decision event per sentence with explicit fields for date, actor/title, adverse action detail, and source document anchor."
            )
            updated["recommended_actions"] = _dedupe_text_values(actions)
        signals = dict(updated.get("signals") or {})
        signals["intake_questioning_score"] = intake_score
        if actor_critic_guidance:
            signals["actor_critic_priority"] = int(actor_critic_guidance.get("priority") or 70)
            signals["router_backed_question_quality"] = bool(actor_critic_guidance.get("router_backed_question_quality"))
            signals["needs_decision_document_precision"] = bool(
                actor_critic_guidance.get("needs_decision_document_precision")
            )
        signals["decision_maker_probe_count"] = int(blocker_signals.get("decision_maker_probe_count", 0) or 0)
        signals["document_anchor_probe_count"] = int(blocker_signals.get("document_anchor_probe_count", 0) or 0)
        updated["signals"] = signals
        return updated

    def _build_intake_questioning_phase_guidance(
        self,
        *,
        drafting_readiness: Dict[str, Any],
        document_optimization: Dict[str, Any],
    ) -> Dict[str, Any]:
        optimization_report = document_optimization if isinstance(document_optimization, dict) else {}
        actor_critic_guidance = self._extract_actor_critic_guidance(optimization_report)
        actor_critic_metrics = dict(actor_critic_guidance.get("metrics") or {}) if actor_critic_guidance else {}
        blocker_signals = self._extract_blocker_follow_up_signals(optimization_report)
        needs_chronology_closure = bool(actor_critic_guidance.get("needs_chronology_closure")) or bool(
            blocker_signals.get("needs_causation_follow_up")
        )
        needs_decision_document_precision = bool(actor_critic_guidance.get("needs_decision_document_precision")) or bool(
            blocker_signals.get("needs_decision_maker_follow_up")
        ) or bool(blocker_signals.get("needs_document_anchor_follow_up"))
        intake_status = (
            optimization_report.get("intake_status")
            if isinstance(optimization_report.get("intake_status"), dict)
            else build_intake_status_summary(self.mediator)
        )
        intake_handoff = (
            intake_status.get("intake_summary_handoff")
            if isinstance(intake_status.get("intake_summary_handoff"), dict)
            else {}
        )
        confirmation = (
            intake_handoff.get("complainant_summary_confirmation")
            if isinstance(intake_handoff.get("complainant_summary_confirmation"), dict)
            else {}
        )
        confirmation_snapshot = (
            confirmation.get("confirmed_summary_snapshot")
            if isinstance(confirmation.get("confirmed_summary_snapshot"), dict)
            else {}
        )
        intake_priority_summary = (
            confirmation_snapshot.get("adversarial_intake_priority_summary")
            if isinstance(confirmation_snapshot.get("adversarial_intake_priority_summary"), dict)
            else {}
        )
        objective_aliases = {
            "staff_names": "staff_names_titles",
            "staff_titles": "staff_names_titles",
            "hearing_timing": "hearing_request_timing",
            "response_timing": "response_dates",
            "causation": "causation_link",
            "adverse_action": "anchor_adverse_action",
            "appeal_rights": "anchor_appeal_rights",
        }

        def _normalize_objective(value: Any) -> str:
            text = str(value or "").strip().lower()
            if not text:
                return ""
            return objective_aliases.get(text, text)

        uncovered_objectives = _dedupe_text_values(
            _normalize_objective(item)
            for item in (intake_priority_summary.get("uncovered_objectives") or [])
            if _normalize_objective(item)
        )
        objective_question_counts = {
            _normalize_objective(key): int(value or 0)
            for key, value in dict(intake_priority_summary.get("objective_question_counts") or {}).items()
            if _normalize_objective(key)
        }
        required_objectives = (
            "timeline",
            "actors",
            "staff_names_titles",
            "causation_link",
            "anchor_adverse_action",
            "anchor_appeal_rights",
            "hearing_request_timing",
            "response_dates",
        )
        missing_required = [
            objective
            for objective in required_objectives
            if objective in uncovered_objectives or int(objective_question_counts.get(objective, 0)) <= 0
        ]
        objective_actions = {
            "timeline": "Capture exact dates for complaint activity, notices, hearing/review requests, and adverse action outcomes.",
            "actors": "Identify who at HACC made, communicated, and carried out each key decision.",
            "staff_names_titles": "Capture each HACC staff member name and title, or best-known title when name is unknown.",
            "causation_link": "Document direct causation facts linking protected activity to adverse action.",
            "anchor_adverse_action": "Confirm the exact denial, termination, or threatened loss of assistance and its communication date.",
            "anchor_appeal_rights": "Confirm whether written notice, informal review, grievance hearing, and appeal rights were provided, requested, denied, or ignored.",
            "hearing_request_timing": "Capture when hearing/review was requested and whether the request timing was acknowledged.",
            "response_dates": "Capture exact response dates for HACC notices, hearing/review responses, and final decision outcomes.",
        }
        if missing_required:
            status = "warning"
            summary = (
                "Intake questioning still needs closure on key blockers (exact dates, staff names/titles, hearing request timing, response dates, "
                "and causation links between protected activity and adverse treatment)."
            )
            actions = _dedupe_text_values(
                objective_actions.get(objective, "")
                for objective in missing_required
                if objective_actions.get(objective, "")
            )
        else:
            status = "ready"
            summary = (
                "Intake questioning currently includes timeline anchors, staff identification, hearing/response timing, and causation probes."
            )
            actions = []

        empathy_score = _safe_float(actor_critic_metrics.get("empathy"), DEFAULT_ACTOR_CRITIC_PRIORITY_METRICS["empathy"])
        question_quality_score = _safe_float(actor_critic_metrics.get("question_quality"), DEFAULT_ACTOR_CRITIC_PRIORITY_METRICS["question_quality"])
        information_extraction_score = _safe_float(actor_critic_metrics.get("information_extraction"), DEFAULT_ACTOR_CRITIC_PRIORITY_METRICS["information_extraction"])
        coverage_score = _safe_float(actor_critic_metrics.get("coverage"), DEFAULT_ACTOR_CRITIC_PRIORITY_METRICS["coverage"])
        efficiency_score = _safe_float(actor_critic_metrics.get("efficiency"), DEFAULT_ACTOR_CRITIC_PRIORITY_METRICS["efficiency"])
        router_backed_question_quality = bool(actor_critic_guidance.get("router_backed_question_quality")) if actor_critic_guidance else False

        if empathy_score < 0.5:
            actions.append(
                "Lead each intake follow-up with a short empathy frame before asking for chronology or decision details."
            )
            actions.append(
                "Use an empathy-forward transition format ('I hear the impact this had. To keep your record accurate...') before requesting exact dates or staff identifiers."
            )
        if question_quality_score < 0.7:
            actions.append(
                "Upgrade question quality by asking one targeted objective per question with explicit date, actor/title, and decision-anchor prompts."
            )
            actions.append(
                "Require each follow-up question to include one objective tag plus a concrete answer target (date, actor/title, notice/response event, or quoted decision language)."
            )
        if information_extraction_score < 0.7:
            actions.append(
                "Improve information extraction by capturing exact event sequence, communication channel, and quoted decision language where available."
            )
        if coverage_score < 0.7:
            actions.append(
                "Increase coverage by ensuring every required objective has at least one answered question and no zero-count objective remains."
            )
        if efficiency_score < 0.75:
            actions.append(
                "Improve intake efficiency by removing multi-part questions and collapsing duplicates to the shortest sufficient prompt."
            )
        if router_backed_question_quality:
            actions.append(
                "Use router-backed question generation to refine prompts for specificity while preserving patchability and objective coverage."
            )
            if question_quality_score < 0.75:
                actions.append(
                    "Route low-quality follow-ups through the llm router with an objective-specific schema so each prompt asks for one verifiable fact and one temporal anchor."
                )
            if information_extraction_score < 0.75:
                actions.append(
                    "Use router-backed extraction checks to normalize each answer into date, actor/title, event, and source fields before drafting updates."
                )
        if needs_chronology_closure:
            actions.append(
                "Close chronology gaps with follow-up questions that require exact dates, elapsed response timing, and event sequence ordering before drafting updates."
            )
            if router_backed_question_quality:
                actions.append(
                    "Use router-backed single-objective prompts to ask one chronology gap at a time (event date, response date, or sequence delta) and reject answers without explicit date anchors."
                )
        if needs_decision_document_precision:
            actions.append(
                "Ask decision-precision follow-ups that capture who made each adverse decision, their title, the exact decision communicated, and the document or notice that records it."
            )
            if router_backed_question_quality:
                actions.append(
                    "Use router-backed extraction checks to normalize adverse-action answers into decision-maker, decision detail, communication channel, document artifact, and date fields."
                )
        actions.append(
            "Keep follow-up prompts patchable by using short, single-objective question templates that can be edited independently without changing global intake flow."
        )
        actions = _dedupe_text_values(actions)
        if actions and status == "ready":
            status = "warning"
        if actions and actor_critic_guidance:
            summary = (
                f"{summary} Actor-critic optimization highlights low empathy/question quality/extraction/coverage/efficiency signals that should be addressed before final drafting."
            ).strip()

        return {
            "priority": 2,
            "status": status,
            "summary": summary,
            "signals": {
                "uncovered_objectives": uncovered_objectives,
                "objective_question_counts": objective_question_counts,
                "missing_required_objectives": missing_required,
                "drafting_status": str(drafting_readiness.get("status") or ""),
                "actor_critic_priority": int(actor_critic_guidance.get("priority") or 70) if actor_critic_guidance else None,
                "actor_critic_metrics": actor_critic_metrics if actor_critic_guidance else {},
                "router_backed_question_quality": router_backed_question_quality,
                "blocker_follow_up_signals": blocker_signals,
                "needs_chronology_closure": needs_chronology_closure,
                "needs_decision_document_precision": needs_decision_document_precision,
                "priority_findings": list(actor_critic_guidance.get("priority_findings") or []) if actor_critic_guidance else [],
            },
            "recommended_actions": actions,
        }

    def _build_workflow_phase_warning_entries(self, workflow_phase_plan: Dict[str, Any]) -> List[Dict[str, Any]]:
        return build_workflow_phase_warning_entries(workflow_phase_plan)

    def _build_formalization_gate_payload(self, drafting_readiness: Dict[str, Any]) -> Dict[str, Any]:
        if not isinstance(drafting_readiness, dict):
            return {}
        phase_status = str(drafting_readiness.get("phase_status") or drafting_readiness.get("status") or "ready").strip().lower() or "ready"
        blockers = [
            str(item).strip()
            for item in list(drafting_readiness.get("blockers") or [])
            if str(item).strip()
        ]
        unresolved_factual_gaps = [
            str(item).strip()
            for item in list(drafting_readiness.get("unresolved_factual_gaps") or [])
            if str(item).strip()
        ]
        unresolved_legal_gaps = [
            str(item).strip()
            for item in list(drafting_readiness.get("unresolved_legal_gaps") or [])
            if str(item).strip()
        ]
        return {
            "ready_for_formalization": phase_status == "ready" and not blockers,
            "phase_status": phase_status,
            "coverage": _safe_float(drafting_readiness.get("coverage"), 0.0),
            "blockers": blockers,
            "unresolved_factual_gaps": unresolved_factual_gaps[:6],
            "unresolved_legal_gaps": unresolved_legal_gaps[:6],
            "weak_complaint_types": list(drafting_readiness.get("weak_complaint_types") or []),
            "weak_evidence_modalities": list(drafting_readiness.get("weak_evidence_modalities") or []),
        }

    def _build_graph_completeness_signals(self, phase_manager: Any) -> Dict[str, Any]:
        if phase_manager is None:
            return {
                "status": "warning",
                "knowledge_graph_available": False,
                "dependency_graph_available": False,
                "remaining_gap_count": 0,
                "current_gap_count": 0,
                "knowledge_graph_enhanced": False,
            }
        graph_phase = build_graph_analysis_phase_guidance(phase_manager, audience="drafting")
        signals = dict(graph_phase.get("signals") or {}) if isinstance(graph_phase, dict) else {}
        return {
            "status": str(graph_phase.get("status") or "warning").strip().lower() if isinstance(graph_phase, dict) else "warning",
            "knowledge_graph_available": _coerce_bool(signals.get("knowledge_graph_available"), default=False),
            "dependency_graph_available": _coerce_bool(signals.get("dependency_graph_available"), default=False),
            "remaining_gap_count": int(signals.get("remaining_gap_count", 0) or 0),
            "current_gap_count": int(signals.get("current_gap_count", 0) or 0),
            "knowledge_graph_enhanced": _coerce_bool(signals.get("knowledge_graph_enhanced"), default=False),
        }

    def _collect_evidence_modality_signals(
        self,
        *,
        draft: Dict[str, Any],
        claim_readiness: List[Dict[str, Any]],
    ) -> Dict[str, Any]:
        policy_count = 0
        file_count = 0

        for claim in claim_readiness:
            if not isinstance(claim, dict):
                continue
            for key, value in dict(claim.get("artifact_family_counts") or {}).items():
                count = int(value or 0)
                token = str(key or "").strip().lower()
                if count <= 0:
                    continue
                if "policy" in token:
                    policy_count += count
                if any(marker in token for marker in ("file", "upload", "document", "pdf", "exhibit", "record")):
                    file_count += count
            for key, value in dict(claim.get("support_by_source") or {}).items():
                count = int(value or 0)
                token = str(key or "").strip().lower()
                if count <= 0:
                    continue
                if "policy_document" in token:
                    policy_count += count
                if "file_evidence" in token:
                    file_count += count

        for exhibit in _coerce_list(draft.get("exhibits")):
            if not isinstance(exhibit, dict):
                continue
            type_text = " ".join(
                str(exhibit.get(key) or "").strip().lower()
                for key in ("type", "evidence_type", "source_type", "title", "description", "link", "path")
            )
            if any(marker in type_text for marker in ("policy", "administrative plan", "acop", "grievance")):
                policy_count += 1
            if any(marker in type_text for marker in ("file", "upload", ".pdf", "exhibit", "record", "document")):
                file_count += 1

        weak_modalities: List[str] = []
        if policy_count <= 0:
            weak_modalities.append("policy_document")
        if file_count <= 0:
            weak_modalities.append("file_evidence")
        return {
            "modalities": {
                "policy_document": int(policy_count),
                "file_evidence": int(file_count),
            },
            "weak_modalities": weak_modalities,
        }

    def _collect_unresolved_readiness_gaps(
        self,
        *,
        claim_readiness: List[Dict[str, Any]],
        sections: Dict[str, Dict[str, Any]],
        graph_signals: Dict[str, Any],
    ) -> Dict[str, List[str]]:
        factual_gaps: List[str] = []
        legal_gaps: List[str] = []

        for claim in claim_readiness:
            if not isinstance(claim, dict):
                continue
            claim_type = str(claim.get("claim_type") or "claim").strip()
            for warning in _coerce_list(claim.get("warnings")):
                if not isinstance(warning, dict):
                    continue
                code = str(warning.get("code") or "").strip().lower()
                message = " ".join(str(warning.get("message") or "").split()).strip()
                if not message:
                    continue
                if code in {"proof_gaps_present", "unresolved_elements", "claim_contradicted"}:
                    factual_gaps.append(f"{claim_type}: {message}")
                elif code in {"adverse_authority_present", "authority_reliability_uncertain"}:
                    legal_gaps.append(f"{claim_type}: {message}")
                else:
                    factual_gaps.append(f"{claim_type}: {message}")

        for section_key, section in sections.items():
            if not isinstance(section, dict):
                continue
            section_name = str(section.get("title") or section_key or "section").strip()
            for warning in _coerce_list(section.get("warnings")):
                if not isinstance(warning, dict):
                    continue
                code = str(warning.get("code") or "").strip().lower()
                message = " ".join(str(warning.get("message") or "").split()).strip()
                if not message:
                    continue
                if code in {"procedural_prerequisites_identified", "jurisdiction_or_venue_incomplete"}:
                    legal_gaps.append(f"{section_name}: {message}")
                else:
                    factual_gaps.append(f"{section_name}: {message}")

        graph_status = str(graph_signals.get("status") or "ready").strip().lower()
        graph_gap_count = max(
            int(graph_signals.get("remaining_gap_count", 0) or 0),
            int(graph_signals.get("current_gap_count", 0) or 0),
        )
        if graph_status != "ready":
            factual_gaps.append(
                "Graph analysis is not ready and chronology/cross-claim dependencies should be closed before formalization."
            )
        if graph_gap_count > 0:
            factual_gaps.append(
                f"Graph analysis still has {graph_gap_count} unresolved intake gap(s) that affect drafting completeness."
            )

        return {
            "factual_gaps": _dedupe_text_values(factual_gaps)[:6],
            "legal_gaps": _dedupe_text_values(legal_gaps)[:6],
        }

    def _build_drafting_readiness(
        self,
        *,
        user_id: str,
        draft: Dict[str, Any],
    ) -> Dict[str, Any]:
        support_summary = self._safe_mediator_dict("summarize_claim_support", user_id=user_id)
        gap_summary = self._safe_mediator_dict("get_claim_support_gaps", user_id=user_id)
        validation_summary = self._safe_mediator_dict("get_claim_support_validation", user_id=user_id)

        support_claims = support_summary.get("claims", {}) if isinstance(support_summary.get("claims"), dict) else {}
        gap_claims = gap_summary.get("claims", {}) if isinstance(gap_summary.get("claims"), dict) else {}
        validation_claims = validation_summary.get("claims", {}) if isinstance(validation_summary.get("claims"), dict) else {}

        claim_types = _unique_preserving_order(
            _extract_text_candidates((draft.get("source_context") or {}).get("claim_types"))
            + list(support_claims.keys())
            + list(validation_claims.keys())
            + [
                str(claim.get("claim_type") or "").strip()
                for claim in _coerce_list(draft.get("claims_for_relief"))
                if isinstance(claim, dict)
            ]
        )

        claim_readiness: List[Dict[str, Any]] = []
        aggregate_warning_count = 0
        overall_status = "ready"

        for claim_type in claim_types:
            support_claim = support_claims.get(claim_type, {}) if isinstance(support_claims.get(claim_type), dict) else {}
            gap_claim = gap_claims.get(claim_type, {}) if isinstance(gap_claims.get(claim_type), dict) else {}
            validation_claim = validation_claims.get(claim_type, {}) if isinstance(validation_claims.get(claim_type), dict) else {}
            overview_payload = self._safe_mediator_dict(
                "get_claim_overview",
                claim_type=claim_type,
                user_id=user_id,
                required_support_kinds=["evidence", "authority"],
            )
            overview_claim = overview_payload.get("claims", {}).get(claim_type, {}) if isinstance(overview_payload.get("claims"), dict) else {}
            treatment_summary = support_claim.get("authority_treatment_summary", {}) if isinstance(support_claim.get("authority_treatment_summary"), dict) else {}
            rule_summary = support_claim.get("authority_rule_candidate_summary", {}) if isinstance(support_claim.get("authority_rule_candidate_summary"), dict) else {}
            source_context = self._extract_support_source_context_counts(support_claim)

            claim_status = "ready"
            warnings: List[Dict[str, Any]] = []

            validation_status = str(validation_claim.get("validation_status") or "")
            if validation_status == "contradicted":
                claim_status = _merge_status(claim_status, "blocked")
                warnings.append(
                    {
                        "code": "claim_contradicted",
                        "severity": "blocked",
                        "message": f"{claim_type.title()} has contradiction signals that should be resolved before filing.",
                    }
                )
            elif validation_status in {"missing", "incomplete"}:
                claim_status = _merge_status(claim_status, "warning")

            if int(validation_claim.get("proof_gap_count", 0) or 0) > 0:
                claim_status = _merge_status(claim_status, "warning")
                warnings.append(
                    {
                        "code": "proof_gaps_present",
                        "severity": "warning",
                        "message": f"{claim_type.title()} still has proof or failed-premise gaps.",
                    }
                )

            if int(treatment_summary.get("adverse_authority_link_count", 0) or 0) > 0:
                claim_status = _merge_status(claim_status, "warning")
                warnings.append(
                    {
                        "code": "adverse_authority_present",
                        "severity": "warning",
                        "message": f"{claim_type.title()} includes adverse or limiting authority that should be reviewed before relying on it in the draft.",
                    }
                )

            uncertain_authority_count = int(treatment_summary.get("uncertain_authority_link_count", 0) or 0)
            uncertain_treatment_types = sorted(
                str(name)
                for name in (treatment_summary.get("treatment_type_counts", {}) or {}).keys()
                if str(name) in {"questioned", "limits", "superseded", "good_law_unconfirmed"}
            )
            if uncertain_authority_count > 0 or uncertain_treatment_types:
                claim_status = _merge_status(claim_status, "warning")
                warnings.append(
                    {
                        "code": "authority_reliability_uncertain",
                        "severity": "warning",
                        "message": f"{claim_type.title()} has authority support with unresolved treatment or good-law uncertainty.",
                    }
                )

            unresolved_elements = int(gap_claim.get("unresolved_count", 0) or 0)
            if unresolved_elements == 0:
                unresolved_elements = len(_coerce_list(overview_claim.get("missing"))) + len(_coerce_list(overview_claim.get("partially_supported")))
            if unresolved_elements > 0:
                claim_status = _merge_status(claim_status, "warning")
                warnings.append(
                    {
                        "code": "unresolved_elements",
                        "severity": "warning",
                        "message": f"{claim_type.title()} still has {unresolved_elements} unresolved claim element(s).",
                    }
                )

            claim_entry = {
                "claim_type": claim_type,
                "status": claim_status,
                "validation_status": validation_status or ("supported" if claim_status == "ready" else "incomplete"),
                "covered_elements": int(support_claim.get("covered_elements", 0) or 0),
                "total_elements": int(support_claim.get("total_elements", 0) or 0),
                "unresolved_element_count": unresolved_elements,
                "proof_gap_count": int(validation_claim.get("proof_gap_count", 0) or 0),
                "contradiction_candidate_count": int(validation_claim.get("contradiction_candidate_count", 0) or 0),
                "support_by_kind": support_claim.get("support_by_kind", {}),
                "support_by_source": source_context["support_by_source"],
                "source_family_counts": source_context["source_family_counts"],
                "record_scope_counts": source_context["record_scope_counts"],
                "artifact_family_counts": source_context["artifact_family_counts"],
                "corpus_family_counts": source_context["corpus_family_counts"],
                "content_origin_counts": source_context["content_origin_counts"],
                "authority_treatment_summary": treatment_summary,
                "authority_rule_candidate_summary": rule_summary,
                "warnings": warnings,
            }
            aggregate_warning_count += len(warnings)
            overall_status = _merge_status(overall_status, claim_status)
            claim_readiness.append(claim_entry)

        claims_section_status = "ready"
        for claim_entry in claim_readiness:
            claims_section_status = _merge_status(claims_section_status, claim_entry.get("status", "ready"))

        total_fact_count = sum(int(claim.get("total_facts", 0) or 0) for claim in support_claims.values() if isinstance(claim, dict))
        if total_fact_count <= 0:
            total_fact_count = sum(
                len(self._normalize_text_lines(claim.get("supporting_facts", [])))
                for claim in _coerce_list(draft.get("claims_for_relief"))
                if isinstance(claim, dict)
            )
        summary_fact_count = len(self._normalize_text_lines(draft.get("summary_of_facts", [])))
        exhibits = _coerce_list(draft.get("exhibits"))
        relief_items = self._normalize_text_lines(draft.get("requested_relief", []))

        sections: Dict[str, Dict[str, Any]] = {}

        facts_status = "ready" if total_fact_count > 0 and summary_fact_count > 0 else "warning"
        facts_warnings: List[Dict[str, Any]] = []
        if facts_status != "ready":
            facts_warnings.append(
                {
                    "code": "fact_support_thin",
                    "severity": "warning",
                    "message": "The factual allegations section has limited fact-backed support and should be reviewed before filing.",
                }
            )
        sections["summary_of_facts"] = {
            "title": "Summary of Facts",
            "status": facts_status,
            "metrics": {
                "summary_fact_count": summary_fact_count,
                "support_fact_count": total_fact_count,
            },
            "warnings": facts_warnings,
        }

        jurisdiction_status = "ready" if draft.get("jurisdiction_statement") and draft.get("venue_statement") else "warning"
        jurisdiction_warnings: List[Dict[str, Any]] = []
        procedural_rule_count = sum(
            int((entry.get("authority_rule_candidate_summary", {}).get("rule_type_counts", {}) or {}).get("procedural_prerequisite", 0) or 0)
            for entry in claim_readiness
            if isinstance(entry, dict)
        )
        if jurisdiction_status != "ready":
            jurisdiction_warnings.append(
                {
                    "code": "jurisdiction_or_venue_incomplete",
                    "severity": "warning",
                    "message": "Jurisdiction or venue language is incomplete and should be confirmed before export.",
                }
            )
        if procedural_rule_count > 0:
            jurisdiction_status = _merge_status(jurisdiction_status, "warning")
            jurisdiction_warnings.append(
                {
                    "code": "procedural_prerequisites_identified",
                    "severity": "warning",
                    "message": "Authority-derived procedural prerequisites were identified and should be checked against the current facts before filing.",
                }
            )
        sections["jurisdiction_and_venue"] = {
            "title": "Jurisdiction and Venue",
            "status": jurisdiction_status,
            "metrics": {
                "procedural_rule_count": procedural_rule_count,
            },
            "warnings": jurisdiction_warnings,
        }

        sections["claims_for_relief"] = {
            "title": "Claims for Relief",
            "status": claims_section_status,
            "metrics": {
                "claim_count": len(claim_readiness),
                "blocked_claim_count": len([entry for entry in claim_readiness if entry.get("status") == "blocked"]),
                "warning_claim_count": len([entry for entry in claim_readiness if entry.get("status") == "warning"]),
            },
            "warnings": [
                warning
                for entry in claim_readiness
                for warning in entry.get("warnings", [])
                if isinstance(warning, dict)
            ],
        }

        exhibits_status = "ready" if exhibits else "warning"
        exhibits_warnings: List[Dict[str, Any]] = []
        if not exhibits:
            exhibits_warnings.append(
                {
                    "code": "no_exhibits",
                    "severity": "warning",
                    "message": "No exhibits are currently attached to the draft package.",
                }
            )
        sections["exhibits"] = {
            "title": "Exhibits",
            "status": exhibits_status,
            "metrics": {
                "exhibit_count": len(exhibits),
            },
            "warnings": exhibits_warnings,
        }

        relief_status = "ready" if relief_items else "warning"
        relief_warnings: List[Dict[str, Any]] = []
        if not relief_items:
            relief_warnings.append(
                {
                    "code": "requested_relief_missing",
                    "severity": "warning",
                    "message": "Requested relief should be confirmed before filing.",
                }
            )
        sections["requested_relief"] = {
            "title": "Requested Relief",
            "status": relief_status,
            "metrics": {
                "requested_relief_count": len(relief_items),
            },
            "warnings": relief_warnings,
        }

        for section in sections.values():
            overall_status = _merge_status(overall_status, str(section.get("status") or "ready"))
            aggregate_warning_count += len(section.get("warnings", []) or [])

        graph_signals = self._build_graph_completeness_signals(getattr(self.mediator, "phase_manager", None))
        gap_signals = self._collect_unresolved_readiness_gaps(
            claim_readiness=claim_readiness,
            sections=sections,
            graph_signals=graph_signals,
        )
        intake_status = build_intake_status_summary(self.mediator)
        intake_handoff = (
            intake_status.get("intake_summary_handoff")
            if isinstance(intake_status, dict) and isinstance(intake_status.get("intake_summary_handoff"), dict)
            else {}
        )
        confirmation = (
            intake_handoff.get("complainant_summary_confirmation")
            if isinstance(intake_handoff.get("complainant_summary_confirmation"), dict)
            else {}
        )
        confirmation_snapshot = (
            confirmation.get("confirmed_summary_snapshot")
            if isinstance(confirmation.get("confirmed_summary_snapshot"), dict)
            else {}
        )
        intake_priority_summary = (
            confirmation_snapshot.get("adversarial_intake_priority_summary")
            if isinstance(confirmation_snapshot.get("adversarial_intake_priority_summary"), dict)
            else {}
        )
        objective_aliases = {
            "staff_names": "staff_names_titles",
            "staff_titles": "staff_names_titles",
            "hearing_timing": "hearing_request_timing",
            "response_timing": "response_dates",
            "causation": "causation_link",
            "adverse_action": "anchor_adverse_action",
            "appeal_rights": "anchor_appeal_rights",
        }

        def _normalize_objective(value: Any) -> str:
            objective_text = str(value or "").strip().lower()
            if not objective_text:
                return ""
            return objective_aliases.get(objective_text, objective_text)

        uncovered_intake_objectives = _dedupe_text_values(
            _normalize_objective(item)
            for item in list(intake_priority_summary.get("uncovered_objectives") or [])
            if _normalize_objective(item)
        )
        objective_question_counts = {
            _normalize_objective(key): int(value or 0)
            for key, value in dict(intake_priority_summary.get("objective_question_counts") or {}).items()
            if _normalize_objective(key)
        }
        required_intake_objectives = (
            "timeline",
            "actors",
            "staff_names_titles",
            "causation_link",
            "anchor_adverse_action",
            "anchor_appeal_rights",
            "hearing_request_timing",
            "response_dates",
        )
        missing_required_intake_objectives = [
            objective
            for objective in required_intake_objectives
            if objective in uncovered_intake_objectives or int(objective_question_counts.get(objective, 0)) <= 0
        ]

        summary_fact_lines = self._normalize_text_lines(draft.get("summary_of_facts", []))
        claim_support_lines = [
            line
            for claim in _coerce_list(draft.get("claims_for_relief"))
            if isinstance(claim, dict)
            for line in self._normalize_text_lines(claim.get("supporting_facts", []))
        ]

        def _has_date_anchor(text: str) -> bool:
            candidate = str(text or "").strip().lower()
            if not candidate:
                return False
            return bool(
                re.search(r"\b\d{1,2}[/-]\d{1,2}(?:[/-]\d{2,4})?\b", candidate)
                or re.search(r"\b\d{4}-\d{2}-\d{2}\b", candidate)
                or re.search(r"\b(?:jan|feb|mar|apr|may|jun|jul|aug|sep|sept|oct|nov|dec)[a-z]*\b", candidate)
            )

        def _has_actor_anchor(text: str) -> bool:
            candidate = str(text or "").strip().lower()
            if not candidate:
                return False
            return bool(
                re.search(r"\b(manager|director|coordinator|officer|supervisor|caseworker|staff|agent|representative|administrator)\b", candidate)
                or re.search(r"\b(hacc|housing authority)\b", candidate)
            )

        def _has_evidence_anchor(text: str) -> bool:
            candidate = str(text or "").strip().lower()
            if not candidate:
                return False
            return bool(
                re.search(r"\b(exhibit|policy|document|notice|letter|email|record|file|pdf|attachment|cid|http)\b", candidate)
            )

        summary_date_anchor_count = sum(1 for line in summary_fact_lines if _has_date_anchor(line))
        summary_actor_anchor_count = sum(1 for line in summary_fact_lines if _has_actor_anchor(line))
        summary_evidence_anchor_count = sum(1 for line in summary_fact_lines if _has_evidence_anchor(line))
        claim_support_date_anchor_count = sum(1 for line in claim_support_lines if _has_date_anchor(line))
        claim_support_actor_anchor_count = sum(1 for line in claim_support_lines if _has_actor_anchor(line))
        claim_support_evidence_anchor_count = sum(1 for line in claim_support_lines if _has_evidence_anchor(line))
        structured_handoff_gaps: List[str] = []
        if summary_fact_lines and summary_date_anchor_count <= 0:
            structured_handoff_gaps.append("summary_of_facts_missing_date_anchors")
        if summary_fact_lines and summary_actor_anchor_count <= 0:
            structured_handoff_gaps.append("summary_of_facts_missing_actor_anchors")
        if summary_fact_lines and summary_evidence_anchor_count <= 0:
            structured_handoff_gaps.append("summary_of_facts_missing_evidence_anchors")
        if claim_support_lines and claim_support_date_anchor_count <= 0:
            structured_handoff_gaps.append("claim_support_missing_date_anchors")
        if claim_support_lines and claim_support_actor_anchor_count <= 0:
            structured_handoff_gaps.append("claim_support_missing_actor_anchors")
        if claim_support_lines and claim_support_evidence_anchor_count <= 0:
            structured_handoff_gaps.append("claim_support_missing_evidence_anchors")
        structured_handoff_signals = {
            "summary_fact_count": len(summary_fact_lines),
            "claim_support_fact_count": len(claim_support_lines),
            "summary_date_anchor_count": int(summary_date_anchor_count),
            "summary_actor_anchor_count": int(summary_actor_anchor_count),
            "summary_evidence_anchor_count": int(summary_evidence_anchor_count),
            "claim_support_date_anchor_count": int(claim_support_date_anchor_count),
            "claim_support_actor_anchor_count": int(claim_support_actor_anchor_count),
            "claim_support_evidence_anchor_count": int(claim_support_evidence_anchor_count),
            "gap_codes": _dedupe_text_values(structured_handoff_gaps),
            "gap_count": len(_dedupe_text_values(structured_handoff_gaps)),
        }
        unresolved_factual_gaps = list(gap_signals.get("factual_gaps") or [])
        unresolved_legal_gaps = list(gap_signals.get("legal_gaps") or [])
        if uncovered_intake_objectives:
            unresolved_factual_gaps.insert(
                0,
                "Intake coverage remains incomplete for one or more required objectives; chronology/actor/notice details should be closed before formalization.",
            )
        if missing_required_intake_objectives:
            unresolved_factual_gaps.insert(
                0,
                "Required intake objectives remain uncovered: " + ", ".join(missing_required_intake_objectives[:5]),
            )
        if int(structured_handoff_signals.get("gap_count", 0) or 0) > 0:
            unresolved_factual_gaps.append(
                "Structured drafting handoff is incomplete; promote intake facts, date/actor anchors, and evidence references into summary-of-facts and claim-support before optimization."
            )
        unresolved_factual_gaps = _dedupe_text_values(unresolved_factual_gaps)[:6]
        evidence_modality_signals = self._collect_evidence_modality_signals(
            draft=draft,
            claim_readiness=claim_readiness,
        )
        weak_evidence_modalities = list(evidence_modality_signals.get("weak_modalities") or [])
        weak_complaint_types = [
            claim_type
            for claim_type in claim_types
            if str(claim_type or "").strip().lower() in {"housing_discrimination", "hacc_research_engine"}
        ]
        targeted_weak_evidence_modalities = [
            modality
            for modality in weak_evidence_modalities
            if str(modality or "").strip().lower() in {"policy_document", "file_evidence"}
        ]
        total_elements = sum(int(entry.get("total_elements", 0) or 0) for entry in claim_readiness if isinstance(entry, dict))
        covered_elements = sum(int(entry.get("covered_elements", 0) or 0) for entry in claim_readiness if isinstance(entry, dict))
        claim_coverage = (
            (float(covered_elements) / float(total_elements))
            if total_elements > 0
            else (1.0 if summary_fact_count > 0 else 0.0)
        )
        ready_section_count = sum(
            1
            for section in sections.values()
            if isinstance(section, dict) and str(section.get("status") or "ready").strip().lower() == "ready"
        )
        section_coverage = (float(ready_section_count) / float(len(sections))) if sections else 0.0
        graph_gap_count = max(
            int(graph_signals.get("remaining_gap_count", 0) or 0),
            int(graph_signals.get("current_gap_count", 0) or 0),
        )
        graph_coverage = 1.0
        if str(graph_signals.get("status") or "ready").strip().lower() != "ready":
            graph_coverage -= 0.05
        if graph_gap_count > 0:
            graph_coverage -= min(0.25, 0.02 * float(graph_gap_count))
        if not _coerce_bool(graph_signals.get("knowledge_graph_enhanced"), default=True):
            graph_coverage -= 0.05
        graph_coverage = max(0.0, min(1.0, graph_coverage))
        graph_gate_active = (
            str(graph_signals.get("status") or "ready").strip().lower() != "ready"
            or graph_gap_count > 0
            or not _coerce_bool(graph_signals.get("knowledge_graph_available"), default=True)
            or not _coerce_bool(graph_signals.get("dependency_graph_available"), default=True)
        )
        coverage = max(
            0.0,
            min(
                1.0,
                round(
                    (0.5 * claim_coverage) + (0.25 * section_coverage) + (0.25 * graph_coverage),
                    3,
                ),
            ),
        )
        phase_status = str(overall_status or "ready").strip().lower() or "ready"
        if str(graph_signals.get("status") or "ready").strip().lower() != "ready":
            phase_status = _merge_status(phase_status, "warning")
        if graph_gate_active:
            phase_status = _merge_status(phase_status, "blocked")
        if (
            coverage < 0.98
            or unresolved_factual_gaps
            or unresolved_legal_gaps
            or missing_required_intake_objectives
            or targeted_weak_evidence_modalities
            or weak_complaint_types
        ):
            phase_status = _merge_status(phase_status, "warning")
        if coverage <= 0.05 and (
            graph_gap_count > 0
            or not claim_readiness
            or unresolved_factual_gaps
            or unresolved_legal_gaps
        ):
            phase_status = "critical"

        blockers: List[str] = []
        if graph_gate_active:
            blockers.append("graph_analysis_not_ready")
        if phase_status in {"warning", "blocked", "critical"}:
            blockers.append("document_generation_not_ready")
        if phase_status == "critical":
            blockers.append("document_generation_critical")
        if unresolved_factual_gaps:
            blockers.append("unresolved_factual_gaps_not_closed")
        if unresolved_legal_gaps:
            blockers.append("unresolved_legal_gaps_not_closed")
        if uncovered_intake_objectives or missing_required_intake_objectives:
            blockers.append("uncovered_intake_objectives")
        if int(structured_handoff_signals.get("gap_count", 0) or 0) > 0:
            blockers.append("structured_intake_handoff_incomplete")
        if weak_complaint_types:
            blockers.append("weak_complaint_type_generalization_needed")
        if targeted_weak_evidence_modalities:
            blockers.append("weak_evidence_modality_support_needed")
        blockers = _dedupe_text_values(blockers)

        readiness_payload = {
            "status": overall_status,
            "claim_types": claim_types,
            "warning_count": aggregate_warning_count,
            "claims": claim_readiness,
            "sections": sections,
            "coverage": coverage,
            "phase_status": phase_status,
            "blockers": blockers,
            "unresolved_factual_gaps": unresolved_factual_gaps,
            "unresolved_legal_gaps": unresolved_legal_gaps,
            "uncovered_intake_objectives": uncovered_intake_objectives,
            "missing_required_intake_objectives": missing_required_intake_objectives,
            "objective_question_counts": objective_question_counts,
            "structured_intake_handoff_signals": structured_handoff_signals,
            "weak_complaint_types": weak_complaint_types,
            "evidence_modalities": dict(evidence_modality_signals.get("modalities") or {}),
            "weak_evidence_modalities": weak_evidence_modalities,
            "graph_completeness_signals": graph_signals,
            "drafting_handoff": {
                "gate_on_graph_completeness": bool(graph_gate_active),
                "graph_phase_status": str(graph_signals.get("status") or "ready").strip().lower() or "ready",
                "graph_remaining_gap_count": int(graph_gap_count),
                "coverage": float(coverage),
                "ready_for_formalization": phase_status == "ready" and not blockers,
                "blockers": list(blockers),
                "uncovered_intake_objectives": uncovered_intake_objectives[:8],
                "missing_required_intake_objectives": missing_required_intake_objectives[:8],
                "unresolved_factual_gaps": unresolved_factual_gaps[:6],
                "unresolved_legal_gaps": unresolved_legal_gaps[:6],
                "targeted_weak_complaint_types": weak_complaint_types[:4],
                "targeted_weak_evidence_modalities": targeted_weak_evidence_modalities[:4],
                "structured_intake_handoff_signals": structured_handoff_signals,
            },
        }
        workflow_phase_plan = self._build_runtime_workflow_phase_plan(
            drafting_readiness=readiness_payload,
            document_optimization=None,
        )
        workflow_warnings = self._build_workflow_phase_warning_entries(workflow_phase_plan)
        if workflow_phase_plan:
            readiness_payload["workflow_phase_plan"] = workflow_phase_plan
        if workflow_warnings:
            readiness_payload["warnings"] = workflow_warnings
            readiness_payload["warning_count"] = int(readiness_payload.get("warning_count") or 0) + len(workflow_warnings)
            for warning in workflow_warnings:
                readiness_payload["status"] = _merge_status(
                    str(readiness_payload.get("status") or "ready"),
                    str(warning.get("severity") or "ready"),
                )
        gap_warnings: List[Dict[str, Any]] = []
        if unresolved_factual_gaps:
            gap_warnings.append(
                {
                    "code": "unresolved_factual_gaps",
                    "severity": "warning",
                    "message": "Unresolved factual gaps remain and should be closed before formalization.",
                    "gaps": unresolved_factual_gaps[:5],
                }
            )
        if unresolved_legal_gaps:
            gap_warnings.append(
                {
                    "code": "unresolved_legal_gaps",
                    "severity": "warning",
                    "message": "Unresolved legal gaps remain and should be closed before formalization.",
                    "gaps": unresolved_legal_gaps[:5],
                }
            )
        if weak_complaint_types:
            gap_warnings.append(
                {
                    "code": "weak_complaint_type_generalization_needed",
                    "severity": "warning",
                    "message": "Target complaint types still need stronger drafting generalization before formalization.",
                    "claim_types": weak_complaint_types[:4],
                }
            )
        if uncovered_intake_objectives or missing_required_intake_objectives:
            gap_warnings.append(
                {
                    "code": "uncovered_intake_objectives",
                    "severity": "warning",
                    "message": "Intake objectives remain uncovered and should be resolved before formalization.",
                    "uncovered_objectives": uncovered_intake_objectives[:8],
                    "missing_required_objectives": missing_required_intake_objectives[:8],
                }
            )
        if int(structured_handoff_signals.get("gap_count", 0) or 0) > 0:
            gap_warnings.append(
                {
                    "code": "structured_intake_handoff_incomplete",
                    "severity": "warning",
                    "message": (
                        "Structured intake handoff is incomplete; promote facts, date/actor anchors, and evidence references "
                        "into summary-of-facts and claim-support generation before document optimization."
                    ),
                    "gap_codes": list(structured_handoff_signals.get("gap_codes") or [])[:8],
                }
            )
        if targeted_weak_evidence_modalities:
            gap_warnings.append(
                {
                    "code": "weak_evidence_modality_support_needed",
                    "severity": "warning",
                    "message": "Weak evidence modalities require stronger source-anchored allegations before formalization.",
                    "modalities": targeted_weak_evidence_modalities[:4],
                }
            )
        if gap_warnings:
            existing_warnings = [
                dict(item)
                for item in list(readiness_payload.get("warnings") or [])
                if isinstance(item, dict)
            ]
            readiness_payload["warnings"] = existing_warnings + gap_warnings
            readiness_payload["warning_count"] = int(readiness_payload.get("warning_count") or 0) + len(gap_warnings)
            readiness_payload["status"] = _merge_status(str(readiness_payload.get("status") or "ready"), "warning")
            readiness_payload["phase_status"] = _merge_status(str(readiness_payload.get("phase_status") or "ready"), "warning")

        return readiness_payload

    def _build_filing_checklist(self, drafting_readiness: Dict[str, Any]) -> List[Dict[str, Any]]:
        if not isinstance(drafting_readiness, dict):
            return []

        checklist: List[Dict[str, Any]] = []
        sections = drafting_readiness.get("sections") if isinstance(drafting_readiness.get("sections"), dict) else {}
        claims = drafting_readiness.get("claims") if isinstance(drafting_readiness.get("claims"), list) else []

        for section_key, section in sections.items():
            if not isinstance(section, dict):
                continue
            status = str(section.get("status") or "ready")
            title = str(section.get("title") or section_key or "Section").strip()
            warnings = section.get("warnings") if isinstance(section.get("warnings"), list) else []
            metrics = section.get("metrics") if isinstance(section.get("metrics"), dict) else {}
            if status == "ready":
                checklist.append(
                    {
                        "scope": "section",
                        "key": str(section_key),
                        "title": title,
                        "status": "ready",
                        "summary": f"{title} is ready for filing review.",
                        "detail": self._summarize_metrics(metrics),
                    }
                )
                continue
            primary_warning = warnings[0] if warnings and isinstance(warnings[0], dict) else {}
            checklist.append(
                {
                    "scope": "section",
                    "key": str(section_key),
                    "title": title,
                    "status": status,
                    "summary": str(primary_warning.get("message") or f"Review {title} before filing."),
                    "detail": self._summarize_metrics(metrics),
                }
            )

        for claim in claims:
            if not isinstance(claim, dict):
                continue
            status = str(claim.get("status") or "ready")
            claim_type = str(claim.get("claim_type") or "claim").strip()
            warnings = claim.get("warnings") if isinstance(claim.get("warnings"), list) else []
            metrics = {
                "covered_elements": claim.get("covered_elements"),
                "total_elements": claim.get("total_elements"),
                "unresolved_element_count": claim.get("unresolved_element_count"),
                "proof_gap_count": claim.get("proof_gap_count"),
            }
            if status == "ready":
                checklist.append(
                    {
                        "scope": "claim",
                        "key": claim_type,
                        "title": claim_type.title(),
                        "status": "ready",
                        "summary": f"{claim_type.title()} is ready for filing review.",
                        "detail": self._summarize_metrics(metrics),
                    }
                )
                continue
            primary_warning = warnings[0] if warnings and isinstance(warnings[0], dict) else {}
            checklist.append(
                {
                    "scope": "claim",
                    "key": claim_type,
                    "title": claim_type.title(),
                    "status": status,
                    "summary": str(primary_warning.get("message") or f"Review {claim_type.title()} before filing."),
                    "detail": self._summarize_metrics(metrics),
                }
            )

        checklist.sort(key=lambda item: {"blocked": 0, "warning": 1, "ready": 2}.get(str(item.get("status")), 3))
        return checklist

    def _annotate_filing_checklist_review_links(
        self,
        *,
        filing_checklist: List[Dict[str, Any]],
        drafting_readiness: Dict[str, Any],
        user_id: Optional[str],
    ) -> None:
        if not filing_checklist or not isinstance(drafting_readiness, dict):
            return

        claim_map: Dict[str, Dict[str, Any]] = {}
        for claim in _coerce_list(drafting_readiness.get("claims")):
            if not isinstance(claim, dict):
                continue
            claim_type = str(claim.get("claim_type") or "").strip()
            if not claim_type:
                continue
            claim_map[claim_type] = {
                "review_url": self._build_review_url(user_id=user_id, claim_type=claim_type),
                "review_context": {
                    "user_id": user_id,
                    "claim_type": claim_type,
                },
            }

        section_map: Dict[str, Dict[str, Any]] = {}
        for section_key, section in (drafting_readiness.get("sections") or {}).items():
            if not isinstance(section, dict):
                continue
            resolved_key = str(section_key or "").strip()
            if not resolved_key:
                continue
            section_map[resolved_key] = {
                "review_url": self._build_review_url(user_id=user_id, section=resolved_key),
                "review_context": {
                    "user_id": user_id,
                    "section": resolved_key,
                    "claim_type": None,
                },
            }

        dashboard_url = self._build_review_url(user_id=user_id)
        for item in filing_checklist:
            if not isinstance(item, dict):
                continue
            scope = str(item.get("scope") or "").strip().lower()
            key = str(item.get("key") or "").strip()
            target = None
            if scope == "claim":
                target = claim_map.get(key)
            elif scope == "section":
                target = section_map.get(key)
            if target:
                item["review_url"] = target["review_url"]
                item["review_context"] = target["review_context"]
            else:
                item["review_url"] = dashboard_url
                item["review_context"] = {"user_id": user_id}

    def _build_review_url(
        self,
        *,
        user_id: Optional[str] = None,
        claim_type: Optional[str] = None,
        section: Optional[str] = None,
    ) -> str:
        params = {}
        if user_id:
            params["user_id"] = user_id
        if claim_type:
            params["claim_type"] = claim_type
        if section:
            params["section"] = section
        query = urlencode(params)
        return f"/claim-support-review?{query}" if query else "/claim-support-review"

    def _summarize_metrics(self, metrics: Dict[str, Any]) -> str:
        parts = []
        for key, value in metrics.items():
            if value in (None, "", []):
                continue
            parts.append(f"{key.replace('_', ' ')}={value}")
            if len(parts) >= 3:
                break
        return "; ".join(parts)

    def _select_statutes_for_claim(
        self,
        claim_type: str,
        statutes: List[Dict[str, Any]],
    ) -> List[Dict[str, Any]]:
        tokens = {token for token in re.split(r"\W+", claim_type.lower()) if token}
        scored = []
        for statute in statutes:
            if not isinstance(statute, dict):
                continue
            haystack = " ".join(
                str(statute.get(field) or "") for field in ("citation", "title", "relevance")
            ).lower()
            score = sum(1 for token in tokens if token in haystack)
            scored.append((score, statute))
        scored.sort(key=lambda item: item[0], reverse=True)
        selected = [statute for score, statute in scored if score > 0][:3]
        return selected or [statute for _, statute in scored[:3]]

    def _extract_overview_elements(self, elements: Any) -> List[str]:
        names = []
        for element in _coerce_list(elements):
            if isinstance(element, dict):
                names.extend(_extract_text_candidates(element.get("element_text") or element.get("claim_element") or element))
            else:
                names.extend(_extract_text_candidates(element))
        return _unique_preserving_order(names)

    def _extract_requested_relief_from_facts(self, facts: List[str]) -> List[str]:
        remedies = []
        for fact in facts:
            lower = fact.lower()
            if "reinstat" in lower:
                remedies.append("Reinstatement or front pay in lieu of reinstatement.")
            if "back pay" in lower or "lost wages" in lower:
                remedies.append("Back pay, front pay, and lost benefits.")
            if "injunct" in lower:
                remedies.append("Injunctive relief to prevent continuing violations.")
        return remedies

    def _collect_exhibits(
        self,
        *,
        user_id: str,
        claim_types: List[str],
        support_claims: Dict[str, Any],
    ) -> List[Dict[str, Any]]:
        exhibits: List[Dict[str, Any]] = []
        evidence_records = _safe_call(self.mediator, "get_user_evidence", user_id=user_id) or []
        for record in _coerce_list(evidence_records):
            if not isinstance(record, dict):
                continue
            claim_type = record.get("claim_type")
            if claim_type and claim_types and claim_type not in claim_types:
                continue
            exhibits.append(
                {
                    "label": f"Exhibit {chr(65 + len(exhibits))}",
                    "title": record.get("description") or record.get("type") or record.get("cid") or "Supporting exhibit",
                    "claim_type": claim_type,
                    "kind": "evidence",
                    "link": self._build_exhibit_link(record),
                    "source_ref": record.get("cid") or record.get("source_url") or "",
                    "summary": record.get("parsed_text_preview") or record.get("description") or "",
                }
            )

        for claim_type, claim_summary in (support_claims or {}).items():
            if not isinstance(claim_summary, dict):
                continue
            for element in _coerce_list(claim_summary.get("elements")):
                if not isinstance(element, dict):
                    continue
                for link in _coerce_list(element.get("links")):
                    if not isinstance(link, dict):
                        continue
                    support_kind = str(link.get("support_kind") or "").strip().lower()
                    if support_kind != "authority":
                        continue
                    link_url = self._build_exhibit_link(link)
                    title = link.get("support_label") or link.get("title") or link.get("citation") or element.get("element_text")
                    source_ref = link.get("support_ref") or link.get("citation") or link_url or ""
                    if not title and not source_ref:
                        continue
                    exhibits.append(
                        {
                            "label": f"Exhibit {chr(65 + len(exhibits))}",
                            "title": title or "Authority support",
                            "claim_type": claim_type,
                            "kind": "authority",
                            "link": link_url,
                            "source_ref": source_ref,
                            "summary": link.get("relevance") or link.get("description") or "",
                        }
                    )

        deduped: List[Dict[str, Any]] = []
        seen = set()
        for exhibit in exhibits:
            key = (exhibit.get("kind"), exhibit.get("title"), exhibit.get("source_ref"))
            if key in seen:
                continue
            seen.add(key)
            deduped.append(exhibit)
            if len(deduped) >= 20:
                break
        return deduped

    def _build_exhibit_link(self, record: Dict[str, Any]) -> str:
        source_url = str(record.get("source_url") or "").strip()
        if source_url:
            return source_url
        support_ref = str(record.get("support_ref") or record.get("cid") or "").strip()
        if support_ref.startswith("http://") or support_ref.startswith("https://"):
            return support_ref
        if support_ref:
            return f"https://ipfs.io/ipfs/{support_ref}"
        return ""

    def _annotate_lines_with_exhibits(
        self,
        lines: List[str],
        exhibits: List[Dict[str, Any]],
    ) -> List[str]:
        if not lines or not exhibits:
            return lines
        annotated: List[str] = []
        for index, line in enumerate(lines):
            exhibit = self._select_exhibit_for_line(line, exhibits)
            if exhibit is None and index == 0:
                exhibit = exhibits[0]
            annotated.append(self._append_exhibit_citation(line, exhibit))
        return annotated

    def _select_exhibit_for_line(
        self,
        line: str,
        exhibits: List[Dict[str, Any]],
    ) -> Optional[Dict[str, Any]]:
        line_tokens = self._text_tokens(line)
        if not line_tokens:
            return exhibits[0] if exhibits else None

        best_match: Optional[Dict[str, Any]] = None
        best_score = 0
        for exhibit in exhibits:
            if not isinstance(exhibit, dict):
                continue
            exhibit_tokens = self._text_tokens(
                " ".join(
                    str(exhibit.get(field) or "")
                    for field in ("title", "summary", "source_ref", "claim_type")
                )
            )
            score = len(line_tokens & exhibit_tokens)
            if score > best_score:
                best_score = score
                best_match = exhibit

        return best_match if best_score > 0 else None

    def _append_exhibit_citation(
        self,
        line: str,
        exhibit: Optional[Dict[str, Any]],
    ) -> str:
        text = str(line or "").strip()
        if not text or exhibit is None:
            return text
        label = str(exhibit.get("label") or "").strip()
        if not label:
            return text
        if label.lower() in text.lower():
            return text
        punctuation = "." if text.endswith(".") else ""
        base = text[:-1] if punctuation else text
        return f"{base} (See {label}){punctuation}"

    def _text_tokens(self, value: str) -> set[str]:
        return {
            token
            for token in re.split(r"\W+", str(value or "").lower())
            if len(token) >= 4
        }

    def _render_docx(self, draft: Dict[str, Any], path: Path) -> None:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.opc.constants import RELATIONSHIP_TYPE
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Inches, Pt, RGBColor

        document = Document()
        section = document.sections[0]
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

        normal_style = document.styles["Normal"]
        normal_style.font.name = "Times New Roman"
        normal_style.font.size = Pt(12)

        heading = document.add_paragraph()
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = heading.add_run(draft.get("court_header", ""))
        run.bold = True
        run.font.size = Pt(12)

        case_caption = draft.get("case_caption", {}) if isinstance(draft.get("case_caption"), dict) else {}
        caption_party_lines = case_caption.get("caption_party_lines") if isinstance(case_caption.get("caption_party_lines"), list) else self._build_caption_party_lines(case_caption)
        caption = document.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.add_run("\n\n".join(caption_party_lines) + "\n")

        case_no = document.add_paragraph()
        case_no.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        case_no.add_run(
            f"{case_caption.get('case_number_label', 'Civil Action No.')} {case_caption.get('case_number', '________________')}"
        ).bold = True
        if case_caption.get("lead_case_number"):
            lead_case = document.add_paragraph()
            lead_case.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            lead_case.add_run(
                f"{case_caption.get('lead_case_number_label', 'Lead Case No.')} {case_caption['lead_case_number']}"
            ).bold = True
        if case_caption.get("related_case_number"):
            related_case = document.add_paragraph()
            related_case.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            related_case.add_run(
                f"{case_caption.get('related_case_number_label', 'Related Case No.')} {case_caption['related_case_number']}"
            ).bold = True
        if case_caption.get("assigned_judge"):
            judge = document.add_paragraph()
            judge.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            judge.add_run(
                f"{case_caption.get('assigned_judge_label', 'Assigned Judge')}: {case_caption['assigned_judge']}"
            ).bold = True
        if case_caption.get("courtroom"):
            courtroom = document.add_paragraph()
            courtroom.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            courtroom.add_run(
                f"{case_caption.get('courtroom_label', 'Courtroom')}: {case_caption['courtroom']}"
            ).bold = True

        title = document.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run(draft.get("case_caption", {}).get("document_title", "COMPLAINT"))
        title_run.bold = True
        title_run.font.size = Pt(14)
        if draft.get("case_caption", {}).get("jury_demand_notice"):
            jury_notice = document.add_paragraph()
            jury_notice.alignment = WD_ALIGN_PARAGRAPH.CENTER
            jury_notice_run = jury_notice.add_run(draft["case_caption"]["jury_demand_notice"])
            jury_notice_run.bold = True
            jury_notice_run.font.size = Pt(12)

        self._add_docx_section(document, "Nature of the Action", draft.get("nature_of_action", []))
        self._add_docx_section(
            document,
            "Parties",
            [
                f"Plaintiff: {', '.join(draft.get('parties', {}).get('plaintiffs', []))}.",
                f"Defendant: {', '.join(draft.get('parties', {}).get('defendants', []))}.",
            ],
        )
        self._add_docx_section(
            document,
            "Jurisdiction and Venue",
            [draft.get("jurisdiction_statement"), draft.get("venue_statement")],
        )
        self._add_docx_numbered_facts(document, "Summary of Facts", draft.get("summary_of_facts", []))
        self._add_docx_numbered_facts(
            document,
            "Factual Allegations",
            draft.get("factual_allegations") or draft.get("summary_of_facts", []),
            groups=draft.get("factual_allegation_groups") if isinstance(draft.get("factual_allegation_groups"), list) else None,
        )
        chronology_lines = draft.get("anchored_chronology_summary", [])
        if chronology_lines:
            self._add_docx_numbered_facts(document, "Anchored Chronology", chronology_lines)

        legal_standards = draft.get("legal_standards", [])
        if legal_standards:
            self._add_docx_section(document, "Applicable Legal Standards", legal_standards)

        document.add_heading("Claims for Relief", level=1)
        for index, claim in enumerate(draft.get("claims_for_relief", []), start=1):
            document.add_heading(f"Count {_roman(index)} - {claim.get('count_title', 'Claim')}", level=2)
            self._add_docx_subsection(document, "Legal Standard", claim.get("legal_standards", []))
            incorporated_clause = self._format_incorporated_reference_clause(
                claim.get("allegation_references", []),
                claim.get("supporting_exhibits", []),
            )
            if incorporated_clause:
                self._add_docx_subsection(document, "Incorporated Support", [incorporated_clause])
            self._add_docx_subsection(document, "Claim-Specific Support", claim.get("supporting_facts", []))
            missing = claim.get("missing_elements", [])
            if missing:
                self._add_docx_subsection(document, "Open Support Gaps", missing)
            exhibits = claim.get("supporting_exhibits", [])
            if exhibits:
                document.add_paragraph("Supporting Exhibits:")
                for exhibit in exhibits:
                    paragraph = document.add_paragraph(style="List Bullet")
                    paragraph.add_run(f"{exhibit.get('label')}. {exhibit.get('title')}")
                    if exhibit.get("link"):
                        paragraph.add_run(" ")
                        self._append_docx_hyperlink(
                            paragraph,
                            exhibit["link"],
                            "Open exhibit",
                            RELATIONSHIP_TYPE,
                            OxmlElement,
                            qn,
                            RGBColor,
                        )

        self._add_docx_subsection(document, "Requested Relief", draft.get("requested_relief", []), numbered=True)
        jury_demand = draft.get("jury_demand", {}) if isinstance(draft.get("jury_demand"), dict) else {}
        if jury_demand:
            self._add_docx_section(document, jury_demand.get("title") or "Jury Demand", [jury_demand.get("text")])

        document.add_heading("Supporting Exhibits", level=1)
        for exhibit in draft.get("exhibits", []):
            paragraph = document.add_paragraph(style="List Bullet")
            paragraph.add_run(f"{exhibit.get('label')}. {exhibit.get('title')}")
            if exhibit.get("summary"):
                paragraph.add_run(f" - {exhibit.get('summary')}")
            if exhibit.get("link"):
                paragraph.add_run(" ")
                self._append_docx_hyperlink(
                    paragraph,
                    exhibit["link"],
                    "Open exhibit",
                    RELATIONSHIP_TYPE,
                    OxmlElement,
                    qn,
                    RGBColor,
                )

        affidavit = draft.get("affidavit", {}) if isinstance(draft.get("affidavit"), dict) else {}
        if affidavit:
            self._add_docx_section(
                document,
                affidavit.get("title") or "Affidavit in Support of Complaint",
                list(_coerce_list(affidavit.get("venue_lines")))
                + [affidavit.get("intro"), affidavit.get("knowledge_graph_note")],
            )
            self._add_docx_numbered_facts(document, "Affiant States as Follows", affidavit.get("facts", []))
            supporting_exhibits = affidavit.get("supporting_exhibits") if isinstance(affidavit.get("supporting_exhibits"), list) else []
            if supporting_exhibits:
                document.add_heading("Affidavit Supporting Exhibits", level=2)
                for exhibit in supporting_exhibits:
                    if not isinstance(exhibit, dict):
                        continue
                    paragraph = document.add_paragraph(style="List Bullet")
                    paragraph.add_run(f"{exhibit.get('label')}. {exhibit.get('title')}")
                    if exhibit.get("link"):
                        paragraph.add_run(f" ({exhibit['link']})")
            self._add_docx_section(
                document,
                "Affidavit Execution",
                [affidavit.get("dated"), affidavit.get("signature_line"), affidavit.get("jurat"), *_coerce_list(affidavit.get("notary_block"))],
            )

        verification = draft.get("verification", {}) if isinstance(draft.get("verification"), dict) else {}
        if verification:
            self._add_docx_section(
                document,
                verification.get("title") or "Verification",
                [verification.get("text"), verification.get("dated"), verification.get("signature_line")],
            )
        certificate_of_service = draft.get("certificate_of_service", {}) if isinstance(draft.get("certificate_of_service"), dict) else {}
        if certificate_of_service:
            self._add_docx_section(
                document,
                certificate_of_service.get("title") or "Certificate of Service",
                [certificate_of_service.get("text")]
                + _coerce_list(certificate_of_service.get("detail_lines"))
                + [certificate_of_service.get("dated"), certificate_of_service.get("signature_line")],
            )
        signature_block = draft.get("signature_block", {}) if isinstance(draft.get("signature_block"), dict) else {}
        self._add_docx_section(
            document,
            "Signature Block",
            self._build_signature_section_lines(signature_block, self._resolve_draft_forum_type(draft)),
        )

        document.save(path)

    def _add_docx_section(self, document: Any, title: str, paragraphs: List[str]) -> None:
        document.add_heading(title, level=1)
        for paragraph in paragraphs:
            if paragraph:
                document.add_paragraph(str(paragraph))

    def _add_docx_numbered_facts(self, document: Any, title: str, facts: List[str], groups: Optional[List[Dict[str, Any]]] = None) -> None:
        document.add_heading(title, level=1)
        if groups:
            for group in groups:
                if not isinstance(group, dict):
                    continue
                heading = str(group.get("title") or "").strip()
                paragraphs = group.get("paragraphs") if isinstance(group.get("paragraphs"), list) else []
                if heading:
                    document.add_paragraph(heading)
                for paragraph in paragraphs:
                    if not isinstance(paragraph, dict):
                        continue
                    number = paragraph.get("number")
                    text = str(paragraph.get("text") or "").strip()
                    if text:
                        document.add_paragraph(f"{number}. {text}" if number else text)
            return
        for index, fact in enumerate(facts, start=1):
            document.add_paragraph(f"{index}. {fact}")

    def _add_docx_subsection(
        self,
        document: Any,
        title: str,
        lines: List[str],
        numbered: bool = False,
    ) -> None:
        document.add_paragraph(title)
        for index, line in enumerate(lines, start=1):
            prefix = f"{index}. " if numbered else ""
            document.add_paragraph(f"{prefix}{line}", style="List Bullet")

    def _append_docx_hyperlink(
        self,
        paragraph: Any,
        url: str,
        text: str,
        relationship_type: Any,
        oxml_element: Any,
        qn: Any,
        rgb_color: Any,
    ) -> None:
        part = paragraph.part
        rel_id = part.relate_to(url, relationship_type.HYPERLINK, is_external=True)
        hyperlink = oxml_element("w:hyperlink")
        hyperlink.set(qn("r:id"), rel_id)
        run = oxml_element("w:r")
        properties = oxml_element("w:rPr")
        color = oxml_element("w:color")
        color.set(qn("w:val"), "0563C1")
        underline = oxml_element("w:u")
        underline.set(qn("w:val"), "single")
        properties.append(color)
        properties.append(underline)
        run.append(properties)
        text_element = oxml_element("w:t")
        text_element.text = text
        run.append(text_element)
        hyperlink.append(run)
        paragraph._p.append(hyperlink)

    def _render_pdf(self, draft: Dict[str, Any], path: Path) -> None:
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
        from reportlab.lib.pagesizes import LETTER
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import inch
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

        styles = getSampleStyleSheet()
        styles.add(
            ParagraphStyle(
                name="CourtHeader",
                parent=styles["Normal"],
                fontName="Times-Bold",
                fontSize=12,
                leading=14,
                alignment=TA_CENTER,
                spaceAfter=12,
            )
        )
        styles.add(
            ParagraphStyle(
                name="Caption",
                parent=styles["Normal"],
                fontName="Times-Roman",
                fontSize=12,
                leading=14,
                alignment=TA_CENTER,
                spaceAfter=12,
            )
        )
        styles.add(
            ParagraphStyle(
                name="SectionHeading",
                parent=styles["Heading1"],
                fontName="Times-Bold",
                fontSize=13,
                leading=15,
                textColor=colors.black,
                alignment=TA_LEFT,
                spaceBefore=10,
                spaceAfter=6,
            )
        )
        styles.add(
            ParagraphStyle(
                name="RightAligned",
                parent=styles["Normal"],
                fontName="Times-Bold",
                fontSize=12,
                leading=14,
                alignment=TA_RIGHT,
                spaceAfter=8,
            )
        )

        doc = SimpleDocTemplate(
            str(path),
            pagesize=LETTER,
            topMargin=inch,
            bottomMargin=inch,
            leftMargin=inch,
            rightMargin=inch,
        )
        case_caption = draft.get("case_caption", {}) if isinstance(draft.get("case_caption"), dict) else {}
        caption_party_lines = case_caption.get("caption_party_lines") if isinstance(case_caption.get("caption_party_lines"), list) else self._build_caption_party_lines(case_caption)
        story = [
            Paragraph(escape(draft.get("court_header", "")), styles["CourtHeader"]),
            Paragraph(
                "<br/><br/>".join(escape(line).replace("\n", "<br/>") for line in caption_party_lines),
                styles["Caption"],
            ),
            Paragraph(
                escape(
                    f"{case_caption.get('case_number_label', 'Civil Action No.')} {case_caption.get('case_number', '________________')}"
                    + (
                        f"\n{case_caption.get('lead_case_number_label', 'Lead Case No.')} {case_caption.get('lead_case_number')}"
                        if case_caption.get('lead_case_number')
                        else ""
                    )
                    + (
                        f"\n{case_caption.get('related_case_number_label', 'Related Case No.')} {case_caption.get('related_case_number')}"
                        if case_caption.get('related_case_number')
                        else ""
                    )
                    + (
                        f"\n{case_caption.get('assigned_judge_label', 'Assigned Judge')}: {case_caption.get('assigned_judge')}"
                        if case_caption.get('assigned_judge')
                        else ""
                    )
                    + (
                        f"\n{case_caption.get('courtroom_label', 'Courtroom')}: {case_caption.get('courtroom')}"
                        if case_caption.get('courtroom')
                        else ""
                    )
                ),
                styles["RightAligned"],
            ),
            Paragraph(
                escape(draft.get("case_caption", {}).get("document_title", "COMPLAINT")),
                styles["CourtHeader"],
            ),
            *(
                [
                    Paragraph(
                        escape(draft["case_caption"]["jury_demand_notice"]),
                        styles["CourtHeader"],
                    )
                ]
                if draft.get("case_caption", {}).get("jury_demand_notice")
                else []
            ),
            Spacer(1, 8),
        ]

        self._append_pdf_section(story, styles, "Nature of the Action", draft.get("nature_of_action", []))
        self._append_pdf_section(
            story,
            styles,
            "Parties",
            [
                f"Plaintiff: {', '.join(draft.get('parties', {}).get('plaintiffs', []))}.",
                f"Defendant: {', '.join(draft.get('parties', {}).get('defendants', []))}.",
            ],
        )
        self._append_pdf_section(
            story,
            styles,
            "Jurisdiction and Venue",
            [draft.get("jurisdiction_statement"), draft.get("venue_statement")],
        )
        self._append_pdf_numbered_section(story, styles, "Summary of Facts", draft.get("summary_of_facts", []))
        self._append_pdf_numbered_section(
            story,
            styles,
            "Factual Allegations",
            draft.get("factual_allegations") or draft.get("summary_of_facts", []),
            groups=draft.get("factual_allegation_groups") if isinstance(draft.get("factual_allegation_groups"), list) else None,
        )
        self._append_pdf_numbered_section(story, styles, "Anchored Chronology", draft.get("anchored_chronology_summary", []))
        self._append_pdf_section(
            story,
            styles,
            "Applicable Legal Standards",
            draft.get("legal_standards", []),
        )

        story.append(Paragraph("Claims for Relief", styles["SectionHeading"]))
        for index, claim in enumerate(draft.get("claims_for_relief", []), start=1):
            story.append(
                Paragraph(
                    escape(f"Count {_roman(index)} - {claim.get('count_title', 'Claim')}"),
                    styles["Heading2"],
                )
            )
            self._append_pdf_section(story, styles, "Legal Standard", claim.get("legal_standards", []), heading_style="Heading3")
            incorporated_clause = self._format_incorporated_reference_clause(
                claim.get("allegation_references", []),
                claim.get("supporting_exhibits", []),
            )
            if incorporated_clause:
                self._append_pdf_section(story, styles, "Incorporated Support", [incorporated_clause], heading_style="Heading3")
            self._append_pdf_section(story, styles, "Claim-Specific Support", claim.get("supporting_facts", []), heading_style="Heading3")
            if claim.get("missing_elements"):
                self._append_pdf_section(story, styles, "Open Support Gaps", claim.get("missing_elements", []), heading_style="Heading3")
            if claim.get("supporting_exhibits"):
                story.append(Paragraph("Supporting Exhibits", styles["Heading3"]))
                for exhibit in claim.get("supporting_exhibits", []):
                    story.append(
                        Paragraph(
                            self._pdf_exhibit_markup(exhibit),
                            styles["Normal"],
                        )
                    )

        self._append_pdf_numbered_section(story, styles, "Requested Relief", draft.get("requested_relief", []))
        jury_demand = draft.get("jury_demand", {}) if isinstance(draft.get("jury_demand"), dict) else {}
        if jury_demand:
            self._append_pdf_section(story, styles, jury_demand.get("title") or "Jury Demand", [jury_demand.get("text")])
        story.append(Paragraph("Supporting Exhibits", styles["SectionHeading"]))
        for exhibit in draft.get("exhibits", []):
            story.append(Paragraph(self._pdf_exhibit_markup(exhibit), styles["Normal"]))

        affidavit = draft.get("affidavit", {}) if isinstance(draft.get("affidavit"), dict) else {}
        if affidavit:
            self._append_pdf_section(
                story,
                styles,
                affidavit.get("title") or "Affidavit in Support of Complaint",
                list(_coerce_list(affidavit.get("venue_lines"))) + [affidavit.get("intro"), affidavit.get("knowledge_graph_note")],
            )
            self._append_pdf_numbered_section(story, styles, "Affiant States as Follows", affidavit.get("facts", []))
            supporting_exhibits = affidavit.get("supporting_exhibits") if isinstance(affidavit.get("supporting_exhibits"), list) else []
            if supporting_exhibits:
                story.append(Paragraph("Affidavit Supporting Exhibits", styles["Heading3"]))
                for exhibit in supporting_exhibits:
                    if not isinstance(exhibit, dict):
                        continue
                    story.append(Paragraph(self._pdf_exhibit_markup(exhibit), styles["Normal"]))
            self._append_pdf_section(
                story,
                styles,
                "Affidavit Execution",
                [affidavit.get("dated"), affidavit.get("signature_line"), affidavit.get("jurat"), *_coerce_list(affidavit.get("notary_block"))],
            )

        verification = draft.get("verification", {}) if isinstance(draft.get("verification"), dict) else {}
        if verification:
            self._append_pdf_section(
                story,
                styles,
                verification.get("title") or "Verification",
                [verification.get("text"), verification.get("dated"), verification.get("signature_line")],
            )
        certificate_of_service = draft.get("certificate_of_service", {}) if isinstance(draft.get("certificate_of_service"), dict) else {}
        if certificate_of_service:
            self._append_pdf_section(
                story,
                styles,
                certificate_of_service.get("title") or "Certificate of Service",
                [certificate_of_service.get("text")]
                + _coerce_list(certificate_of_service.get("detail_lines"))
                + [certificate_of_service.get("dated"), certificate_of_service.get("signature_line")],
            )
        signature_block = draft.get("signature_block", {}) if isinstance(draft.get("signature_block"), dict) else {}
        self._append_pdf_section(
            story,
            styles,
            "Signature Block",
            self._build_signature_section_lines(signature_block, self._resolve_draft_forum_type(draft)),
        )

        doc.build(story)

    def _append_pdf_section(
        self,
        story: List[Any],
        styles: Any,
        title: str,
        paragraphs: List[str],
        heading_style: str = "SectionHeading",
    ) -> None:
        from reportlab.platypus import Paragraph

        if not paragraphs:
            return
        story.append(Paragraph(escape(title), styles[heading_style]))
        for paragraph in paragraphs:
            story.append(Paragraph(escape(str(paragraph)), styles["Normal"]))

    def _append_pdf_numbered_section(
        self,
        story: List[Any],
        styles: Any,
        title: str,
        paragraphs: List[str],
        heading_style: str = "SectionHeading",
        groups: Optional[List[Dict[str, Any]]] = None,
    ) -> None:
        from reportlab.platypus import Paragraph

        if not paragraphs and not groups:
            return
        story.append(Paragraph(escape(title), styles[heading_style]))
        if groups:
            for group in groups:
                if not isinstance(group, dict):
                    continue
                group_title = str(group.get("title") or "").strip()
                entries = group.get("paragraphs") if isinstance(group.get("paragraphs"), list) else []
                if group_title:
                    story.append(Paragraph(escape(group_title), styles["Heading3"]))
                for entry in entries:
                    if not isinstance(entry, dict):
                        continue
                    number = entry.get("number")
                    text = str(entry.get("text") or "").strip()
                    if text:
                        prefix = f"{number}. " if number else ""
                        story.append(Paragraph(escape(f"{prefix}{text}"), styles["Normal"]))
            return
        for index, paragraph in enumerate(paragraphs, start=1):
            story.append(Paragraph(escape(f"{index}. {paragraph}"), styles["Normal"]))

    def _pdf_exhibit_markup(self, exhibit: Dict[str, Any]) -> str:
        title = escape(f"{exhibit.get('label')}. {exhibit.get('title')}")
        summary = escape(str(exhibit.get("summary") or ""))
        link = str(exhibit.get("link") or "").strip()
        if link:
            link_markup = f'<link href="{escape(link)}">Open exhibit</link>'
            if summary:
                return f"{title} - {summary} ({link_markup})"
            return f"{title} ({link_markup})"
        if summary:
            return f"{title} - {summary}"
        return title
