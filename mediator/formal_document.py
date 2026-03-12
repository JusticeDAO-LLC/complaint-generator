"""Formal complaint document assembly and export helpers."""

from __future__ import annotations

from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Sequence, Tuple

from complaint_phases import ComplaintPhase, NodeType

try:
    from docx import Document as DocxDocument
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt

    HAS_DOCX = True
except ImportError:  # pragma: no cover - optional dependency
    DocxDocument = None
    WD_ALIGN_PARAGRAPH = None
    Inches = None
    Pt = None
    HAS_DOCX = False

try:
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

    HAS_REPORTLAB = True
except ImportError:  # pragma: no cover - optional dependency
    colors = None
    TA_CENTER = None
    letter = None
    ParagraphStyle = None
    getSampleStyleSheet = None
    inch = None
    Paragraph = None
    SimpleDocTemplate = None
    Spacer = None
    Table = None
    TableStyle = None
    HAS_REPORTLAB = False


def _utc_now_isoformat() -> str:
    return datetime.now(timezone.utc).isoformat()


def _clean_text(value: Any) -> str:
    text = str(value or "").replace("\r", "\n")
    lines = [" ".join(line.split()) for line in text.split("\n")]
    return "\n".join(line for line in lines if line).strip()


def _clean_sentence(value: Any) -> str:
    text = _clean_text(value)
    if not text:
        return ""
    return text if text.endswith((".", "?", "!", ":")) else f"{text}."


def _listify(value: Any) -> List[Any]:
    if value is None:
        return []
    if isinstance(value, list):
        return list(value)
    if isinstance(value, tuple):
        return list(value)
    return [value]


def _roman_numeral(number: int) -> str:
    values: Sequence[Tuple[int, str]] = (
        (1000, "M"),
        (900, "CM"),
        (500, "D"),
        (400, "CD"),
        (100, "C"),
        (90, "XC"),
        (50, "L"),
        (40, "XL"),
        (10, "X"),
        (9, "IX"),
        (5, "V"),
        (4, "IV"),
        (1, "I"),
    )
    remaining = max(int(number or 0), 1)
    output: List[str] = []
    for integer, numeral in values:
        while remaining >= integer:
            output.append(numeral)
            remaining -= integer
    return "".join(output)


def _exhibit_label(index: int) -> str:
    ordinal = max(index, 1)
    letters: List[str] = []
    while ordinal > 0:
        ordinal -= 1
        letters.append(chr(ord("A") + (ordinal % 26)))
        ordinal //= 26
    return f"Exhibit {''.join(reversed(letters))}"


class ComplaintDocumentBuilder:
    """Build and export a court-style complaint draft from mediator state."""

    def __init__(self, mediator):
        self.mediator = mediator

    def build(
        self,
        *,
        court_name: Optional[str] = None,
        district: Optional[str] = None,
        division: Optional[str] = None,
        court_header_override: Optional[str] = None,
        case_number: Optional[str] = None,
        title_override: Optional[str] = None,
        plaintiff_names: Optional[List[str]] = None,
        defendant_names: Optional[List[str]] = None,
        requested_relief: Optional[List[str]] = None,
        user_id: Optional[str] = None,
        base_formal_complaint: Optional[Dict[str, Any]] = None,
    ) -> Dict[str, Any]:
        phase_manager = getattr(self.mediator, "phase_manager", None)
        kg = phase_manager.get_phase_data(ComplaintPhase.INTAKE, "knowledge_graph") if phase_manager else None
        dg = phase_manager.get_phase_data(ComplaintPhase.INTAKE, "dependency_graph") if phase_manager else None
        legal_graph = phase_manager.get_phase_data(ComplaintPhase.FORMALIZATION, "legal_graph") if phase_manager else None
        matching_results = phase_manager.get_phase_data(ComplaintPhase.FORMALIZATION, "matching_results") if phase_manager else None

        resolved_user_id = user_id or getattr(self.mediator.state, "username", None) or getattr(self.mediator.state, "hashed_username", None) or "anonymous"
        base = dict(base_formal_complaint or {})
        evidence_records = self._safe_call("get_user_evidence", resolved_user_id, default=[])
        authority_records = self._safe_call("get_legal_authorities", resolved_user_id, default=[])
        support_links = self._safe_call("get_claim_support", resolved_user_id, default=[])
        intake_summary = self._collect_intake_summary()
        parties = self._build_parties(
            kg,
            base.get("parties"),
            plaintiff_names=plaintiff_names,
            defendant_names=defendant_names,
        )
        jurisdiction = str(base.get("jurisdiction") or self._infer_jurisdiction(legal_graph) or "federal")
        caption = self._build_caption(
            parties=parties,
            jurisdiction=jurisdiction,
            court_name=court_name,
            district=district,
            division=division,
            court_header_override=court_header_override,
            case_number=case_number,
            title=title_override or str(base.get("title") or "").strip(),
        )
        factual_allegations = self._build_factual_allegations(
            base.get("factual_allegations"),
            kg,
            intake_summary,
        )
        exhibits = self._build_exhibits(evidence_records, support_links)
        claims = self._build_claims(
            matching_results=matching_results,
            dependency_graph=dg,
            legal_graph=legal_graph,
            authority_records=authority_records,
            exhibits=exhibits,
            factual_allegations=factual_allegations,
            user_id=resolved_user_id,
        )
        requested_relief = self._build_requested_relief(requested_relief or base.get("prayer_for_relief"))
        nature_of_action = self._build_nature_of_action(base.get("statement_of_claim"), claims)
        jurisdiction_statement = self._build_jurisdiction_statement(jurisdiction, authority_records)
        venue_statement = self._build_venue_statement(district, division)
        signature_block = self._build_signature_block(parties)
        legal_standards = [
            {
                "claim_name": claim.get("claim_name", ""),
                "claim_type": claim.get("claim_type", ""),
                "standard": claim.get("legal_standard", ""),
                "citations": [
                    item.get("citation", "")
                    for item in claim.get("legal_standard_elements", [])
                    if item.get("citation")
                ],
            }
            for claim in claims
        ]

        draft = {
            **base,
            "generated_at": _utc_now_isoformat(),
            "court_header": caption["court_header"],
            "caption": caption,
            "case_number": caption["case_number"],
            "title": caption["case_title"],
            "parties": parties,
            "nature_of_action": nature_of_action,
            "jurisdiction": jurisdiction,
            "jurisdiction_statement": jurisdiction_statement,
            "venue_statement": venue_statement,
            "statement_of_claim": str(base.get("statement_of_claim") or nature_of_action),
            "summary_of_facts": factual_allegations[: min(len(factual_allegations), 6)],
            "factual_allegations": factual_allegations,
            "legal_standards": legal_standards,
            "legal_claims": claims,
            "claims_for_relief": claims,
            "prayer_for_relief": requested_relief,
            "requested_relief": requested_relief,
            "supporting_documents": exhibits,
            "exhibits": exhibits,
            "supporting_exhibits": exhibits,
            "intake_summary": intake_summary,
            "signature_block": signature_block,
        }
        draft["draft_text"] = self.render_text(draft)
        return draft

    def export(self, draft: Dict[str, Any], output_path: str, *, format: Optional[str] = None) -> Dict[str, Any]:
        destination = Path(output_path)
        resolved_format = (format or destination.suffix.lstrip(".") or "txt").lower()
        if resolved_format == "txt":
            destination.write_text(draft.get("draft_text") or self.render_text(draft), encoding="utf-8")
        elif resolved_format == "docx":
            self._write_docx(draft, destination)
        elif resolved_format == "pdf":
            self._write_pdf(draft, destination)
        else:
            raise ValueError(f"Unsupported complaint export format: {resolved_format}")

        return {
            "path": str(destination),
            "format": resolved_format,
            "bytes_written": destination.stat().st_size,
        }

    def render_text(self, draft: Dict[str, Any]) -> str:
        caption = draft.get("caption", {}) if isinstance(draft.get("caption"), dict) else {}
        lines: List[str] = []
        lines.append(str(draft.get("court_header") or "IN THE COURT OF COMPETENT JURISDICTION"))
        if caption.get("division_line"):
            lines.append(str(caption["division_line"]))
        lines.append("")
        plaintiffs = ", ".join(draft.get("parties", {}).get("plaintiffs", []) or ["Plaintiff"])
        defendants = ", ".join(draft.get("parties", {}).get("defendants", []) or ["Defendant"])
        case_number = draft.get("case_number") or "________________"
        lines.append(f"{plaintiffs}, Plaintiff,")
        lines.append("v.")
        lines.append(f"{defendants}, Defendant.")
        lines.append(f"Case No.: {case_number}")
        lines.append("")
        lines.append("COMPLAINT")
        lines.append("")
        lines.append("NATURE OF THE ACTION")
        lines.append(_clean_sentence(draft.get("nature_of_action")))
        lines.append("")
        lines.append("PARTIES")
        for index, party in enumerate(draft.get("parties", {}).get("plaintiffs", []), 1):
            lines.append(f"{index}. Plaintiff {party} is an aggrieved party bringing this action.")
        start_index = len(draft.get("parties", {}).get("plaintiffs", []))
        for offset, party in enumerate(draft.get("parties", {}).get("defendants", []), 1):
            lines.append(f"{start_index + offset}. Defendant {party} is alleged to be responsible for the acts described below.")
        lines.append("")
        lines.append("JURISDICTION AND VENUE")
        lines.append(_clean_sentence(draft.get("jurisdiction_statement")))
        lines.append(_clean_sentence(draft.get("venue_statement")))
        lines.append("")
        lines.append("FACTUAL ALLEGATIONS")
        for index, allegation in enumerate(_listify(draft.get("factual_allegations")), 1):
            lines.append(f"{index}. {_clean_sentence(allegation)}")

        claims = _listify(draft.get("legal_claims"))
        if claims:
            lines.append("")
            lines.append("CLAIMS FOR RELIEF")
        for claim in claims:
            lines.append("")
            lines.append(str(claim.get("title") or claim.get("claim_name") or "Claim"))
            if claim.get("description"):
                lines.append(_clean_sentence(claim.get("description")))
            if claim.get("legal_standard"):
                lines.append("Legal Standard:")
                lines.append(_clean_sentence(claim.get("legal_standard")))
            for item in _listify(claim.get("legal_standard_elements")):
                element_text = _clean_sentence(item.get("element") or item.get("description"))
                citation = str(item.get("citation") or "").strip()
                if citation:
                    lines.append(f"- {element_text} ({citation})")
                elif element_text:
                    lines.append(f"- {element_text}")
            supporting_facts = _listify(claim.get("supporting_facts"))
            if supporting_facts:
                lines.append("Supporting Facts:")
                for fact in supporting_facts:
                    lines.append(f"- {_clean_sentence(fact)}")
            authorities = _listify(claim.get("supporting_authorities"))
            if authorities:
                lines.append("Supporting Authorities:")
                for authority in authorities:
                    authority_line = authority.get("citation") or authority.get("title") or "Authority"
                    if authority.get("title") and authority.get("citation"):
                        authority_line = f"{authority['citation']} - {authority['title']}"
                    lines.append(f"- {_clean_text(authority_line)}")
            supporting_exhibits = _listify(claim.get("supporting_exhibits"))
            if supporting_exhibits:
                exhibit_labels = ", ".join(exhibit.get("label", "") for exhibit in supporting_exhibits if exhibit.get("label"))
                if exhibit_labels:
                    lines.append(f"Exhibits Incorporated by Reference: {exhibit_labels}.")

        lines.append("")
        lines.append("PRAYER FOR RELIEF")
        for item in _listify(draft.get("requested_relief")):
            lines.append(f"- {_clean_sentence(item)}")

        exhibits = _listify(draft.get("exhibits"))
        if exhibits:
            lines.append("")
            lines.append("EXHIBITS")
            for exhibit in exhibits:
                reference = exhibit.get("reference") or exhibit.get("source_url") or exhibit.get("cid") or ""
                summary = exhibit.get("summary") or exhibit.get("description") or ""
                exhibit_line = f"{exhibit.get('label', 'Exhibit')} - {_clean_text(exhibit.get('title') or exhibit.get('description') or 'Supporting exhibit')}"
                if reference:
                    exhibit_line = f"{exhibit_line} ({reference})"
                lines.append(exhibit_line)
                if summary:
                    lines.append(f"  {_clean_sentence(summary)}")

        signature_block = draft.get("signature_block", {}) if isinstance(draft.get("signature_block"), dict) else {}
        lines.append("")
        lines.append("Respectfully submitted,")
        lines.append(signature_block.get("name") or "Plaintiff")
        if signature_block.get("contact"):
            lines.append(signature_block["contact"])
        return "\n".join(line for line in lines if line is not None)

    def _safe_call(self, method_name: str, *args, default=None, **kwargs):
        method = getattr(self.mediator, method_name, None)
        if not callable(method):
            return default
        try:
            return method(*args, **kwargs)
        except Exception:
            return default

    def _collect_intake_summary(self) -> List[Dict[str, str]]:
        summary: List[Dict[str, str]] = []
        seen = set()
        for inquiry in _listify(getattr(self.mediator.state, "inquiries", [])):
            if not isinstance(inquiry, dict):
                continue
            question = _clean_text(inquiry.get("question"))
            answer = _clean_text(inquiry.get("answer"))
            if not question or not answer:
                continue
            marker = (question, answer)
            if marker in seen:
                continue
            seen.add(marker)
            summary.append({"question": question, "answer": answer})

        answered_questions = getattr(self.mediator.state, "answered_questions", {})
        if isinstance(answered_questions, dict):
            for question, answer in answered_questions.items():
                if str(question) == "last_question":
                    continue
                question_text = _clean_text(question)
                answer_text = _clean_text(answer)
                if not question_text or not answer_text:
                    continue
                marker = (question_text, answer_text)
                if marker in seen:
                    continue
                seen.add(marker)
                summary.append({"question": question_text, "answer": answer_text})
        return summary

    def _build_parties(
        self,
        knowledge_graph,
        base_parties: Any,
        *,
        plaintiff_names: Optional[List[str]] = None,
        defendant_names: Optional[List[str]] = None,
    ) -> Dict[str, List[str]]:
        plaintiffs: List[str] = []
        defendants: List[str] = []
        others: List[str] = []

        plaintiffs.extend(_clean_text(name) for name in _listify(plaintiff_names))
        defendants.extend(_clean_text(name) for name in _listify(defendant_names))

        if isinstance(base_parties, dict):
            plaintiffs.extend(_clean_text(name) for name in _listify(base_parties.get("plaintiffs")))
            defendants.extend(_clean_text(name) for name in _listify(base_parties.get("defendants")))

        if knowledge_graph is not None:
            for entity in knowledge_graph.get_entities_by_type("person"):
                role = str(entity.attributes.get("role") or "").lower()
                name = _clean_text(entity.name)
                if not name:
                    continue
                if any(token in role for token in ("plaintiff", "complainant", "claimant", "petitioner")):
                    plaintiffs.append(name)
                elif any(token in role for token in ("defendant", "respondent", "manager", "supervisor", "owner", "employer")):
                    defendants.append(name)
                else:
                    others.append(name)

            organizations = [_clean_text(entity.name) for entity in knowledge_graph.get_entities_by_type("organization") if _clean_text(entity.name)]
            defendants.extend(organizations)

        plaintiffs = self._dedupe(plaintiffs) or ["Plaintiff"]
        defendants = self._dedupe(defendants) or ["Defendant"]
        return {
            "plaintiffs": plaintiffs,
            "defendants": defendants,
            "other_parties": self._dedupe(others),
        }

    def _build_caption(
        self,
        *,
        parties: Dict[str, List[str]],
        jurisdiction: str,
        court_name: Optional[str],
        district: Optional[str],
        division: Optional[str],
        court_header_override: Optional[str],
        case_number: Optional[str],
        title: str,
    ) -> Dict[str, str]:
        district_text = _clean_text(district).upper()
        division_text = _clean_text(division).upper()
        if court_header_override:
            court_header = _clean_text(court_header_override).upper()
        elif court_name:
            court_header = _clean_text(court_name).upper()
        elif district_text:
            court_header = f"IN THE UNITED STATES DISTRICT COURT FOR THE DISTRICT OF {district_text}"
        elif str(jurisdiction).lower() in {"federal", "us", "united states"}:
            court_header = "IN THE UNITED STATES DISTRICT COURT"
        else:
            court_header = "IN THE COURT OF COMPETENT JURISDICTION"
        resolved_title = title or f"{parties['plaintiffs'][0]} v. {parties['defendants'][0]}"
        return {
            "court_header": court_header,
            "division_line": division_text,
            "case_title": resolved_title,
            "case_number": _clean_text(case_number) or "________________",
        }

    def _build_factual_allegations(self, base_allegations: Any, knowledge_graph, intake_summary: List[Dict[str, str]]) -> List[str]:
        allegations = [_clean_sentence(item) for item in _listify(base_allegations) if _clean_text(item)]
        if allegations:
            return allegations
        if knowledge_graph is not None:
            fact_entities = []
            for entity in knowledge_graph.entities.values():
                if getattr(entity, "type", "") == "fact":
                    fact_entities.append(_clean_sentence(entity.name))
            if fact_entities:
                return self._dedupe(fact_entities)
        fallback = []
        for item in intake_summary:
            question = item.get("question", "")
            answer = item.get("answer", "")
            if answer:
                fallback.append(f"{question}: {answer}" if question else answer)
        return self._dedupe([_clean_sentence(item) for item in fallback]) or ["Plaintiff will supplement the factual record with additional detail."]

    def _build_exhibits(self, evidence_records: Any, support_links: Any) -> List[Dict[str, Any]]:
        support_by_ref: Dict[str, Dict[str, Any]] = {}
        for link in _listify(support_links):
            if not isinstance(link, dict):
                continue
            support_ref = str(link.get("support_ref") or "").strip()
            if support_ref and support_ref not in support_by_ref:
                support_by_ref[support_ref] = link

        exhibits: List[Dict[str, Any]] = []
        for index, record in enumerate(_listify(evidence_records), 1):
            if not isinstance(record, dict):
                continue
            cid = str(record.get("cid") or "").strip()
            metadata = record.get("metadata", {}) if isinstance(record.get("metadata"), dict) else {}
            parse_metadata = record.get("parse_metadata", {}) if isinstance(record.get("parse_metadata"), dict) else {}
            filename = _clean_text(metadata.get("filename") or parse_metadata.get("filename") or "")
            source_url = _clean_text(record.get("source_url") or metadata.get("source_url") or "")
            reference = source_url or (f"ipfs://{cid}" if cid else "")
            support_link = support_by_ref.get(cid) or support_by_ref.get(source_url) or {}
            exhibits.append(
                {
                    "label": _exhibit_label(index),
                    "title": _clean_text(record.get("description") or filename or record.get("type") or f"Exhibit {index}"),
                    "description": _clean_text(record.get("description") or ""),
                    "claim_type": _clean_text(record.get("claim_type") or support_link.get("claim_type") or "").lower(),
                    "claim_element": _clean_text(record.get("claim_element") or support_link.get("claim_element_text") or ""),
                    "reference": reference,
                    "source_url": source_url,
                    "cid": cid,
                    "summary": _clean_text(record.get("parsed_text_preview") or ""),
                    "fact_count": int(record.get("fact_count") or 0),
                }
            )
        return exhibits

    def _build_claims(
        self,
        *,
        matching_results: Any,
        dependency_graph,
        legal_graph,
        authority_records: Any,
        exhibits: List[Dict[str, Any]],
        factual_allegations: List[str],
        user_id: str,
    ) -> List[Dict[str, Any]]:
        claim_entries = []
        raw_claims = []
        if isinstance(matching_results, dict):
            raw_claims.extend(_listify(matching_results.get("claims")))
        if not raw_claims and dependency_graph is not None:
            for node in dependency_graph.get_nodes_by_type(NodeType.CLAIM):
                raw_claims.append(
                    {
                        "claim_name": node.name,
                        "claim_type": node.attributes.get("claim_type", node.name),
                        "satisfied_requirements": 0,
                        "legal_requirements": 0,
                        "requirements": [],
                        "missing_requirements": [],
                    }
                )

        for index, claim in enumerate(raw_claims, 1):
            if not isinstance(claim, dict):
                continue
            claim_name = _clean_text(claim.get("claim_name") or claim.get("title") or f"Claim {index}")
            claim_type = _clean_text(claim.get("claim_type") or claim_name).lower().replace(" ", "_")
            legal_requirements = legal_graph.get_requirements_for_claim_type(claim_type) if legal_graph is not None else []
            claim_authorities = self._filter_authorities(authority_records, claim_type)
            supporting_facts = self._build_supporting_facts(claim_type, claim_name, factual_allegations, user_id)
            supporting_exhibits = [exhibit for exhibit in exhibits if not exhibit.get("claim_type") or exhibit.get("claim_type") == claim_type]
            legal_standard_elements = []
            for requirement in legal_requirements:
                legal_standard_elements.append(
                    {
                        "element": _clean_text(requirement.description or requirement.name),
                        "citation": _clean_text(requirement.citation),
                    }
                )
            if not legal_standard_elements:
                for requirement in _listify(claim.get("requirements")):
                    if not isinstance(requirement, dict):
                        continue
                    legal_standard_elements.append(
                        {
                            "element": _clean_text(requirement.get("requirement_description") or requirement.get("requirement_name") or ""),
                            "citation": _clean_text(requirement.get("citation") or ""),
                        }
                    )

            legal_standard = self._compose_legal_standard(claim_name, legal_standard_elements, claim_authorities)
            missing_requirements = [
                {
                    "name": _clean_text(item.get("requirement_name") or ""),
                    "citation": _clean_text(item.get("citation") or ""),
                    "suggested_action": _clean_text(item.get("suggested_action") or ""),
                }
                for item in _listify(claim.get("missing_requirements"))
                if isinstance(item, dict)
            ]
            claim_entries.append(
                {
                    "count": index,
                    "title": f"COUNT {_roman_numeral(index)} - {claim_name.upper()}",
                    "claim_name": claim_name,
                    "claim_type": claim_type,
                    "description": _clean_sentence(claim.get("description") or f"Plaintiff realleges the foregoing paragraphs and asserts {claim_name}."),
                    "elements_satisfied": f"{int(claim.get('satisfied_requirements') or 0)}/{int(claim.get('legal_requirements') or 0)}",
                    "legal_standard": legal_standard,
                    "legal_standard_elements": legal_standard_elements,
                    "supporting_facts": supporting_facts,
                    "supporting_authorities": claim_authorities,
                    "supporting_exhibits": supporting_exhibits,
                    "missing_requirements": missing_requirements,
                }
            )
        return claim_entries

    def _build_supporting_facts(self, claim_type: str, claim_name: str, factual_allegations: List[str], user_id: str) -> List[str]:
        facts = self._safe_call("get_claim_support_facts", claim_type, user_id, default=[])
        supporting_facts = []
        for fact in _listify(facts):
            if not isinstance(fact, dict):
                continue
            text = _clean_sentence(fact.get("text") or fact.get("fact_text") or "")
            if text:
                supporting_facts.append(text)
        if supporting_facts:
            return self._dedupe(supporting_facts)
        lowered_claim_name = claim_name.lower()
        filtered = [
            allegation for allegation in factual_allegations
            if claim_type.replace("_", " ") in allegation.lower() or lowered_claim_name in allegation.lower()
        ]
        return self._dedupe(filtered or factual_allegations[:3])

    def _filter_authorities(self, authority_records: Any, claim_type: str) -> List[Dict[str, Any]]:
        matched = []
        for authority in _listify(authority_records):
            if not isinstance(authority, dict):
                continue
            authority_claim_type = _clean_text(authority.get("claim_type") or "").lower()
            if authority_claim_type and authority_claim_type != claim_type:
                continue
            matched.append(
                {
                    "citation": _clean_text(authority.get("citation") or ""),
                    "title": _clean_text(authority.get("title") or authority.get("authority_type") or "Authority"),
                    "url": _clean_text(authority.get("url") or ""),
                    "relevance_score": authority.get("relevance_score"),
                }
            )
        return matched[:5]

    def _build_requested_relief(self, base_relief: Any) -> List[str]:
        relief_items = [_clean_sentence(item) for item in _listify(base_relief) if _clean_text(item)]
        if relief_items:
            return relief_items
        return [
            "Compensatory damages in an amount to be proven at trial",
            "Injunctive and declaratory relief sufficient to stop the challenged conduct",
            "Costs, fees, and any statutory fee-shifting relief authorized by law",
            "Such other and further relief as the Court deems just and proper",
        ]

    def _build_nature_of_action(self, statement_of_claim: Any, claims: List[Dict[str, Any]]) -> str:
        statement = _clean_sentence(statement_of_claim)
        if statement:
            return statement
        claim_names = [claim.get("claim_name", "") for claim in claims if claim.get("claim_name")]
        if claim_names:
            if len(claim_names) == 1:
                return f"This civil action arises from {claim_names[0]} and seeks relief for the injuries caused by that conduct."
            joined = ", ".join(claim_names[:-1]) + f", and {claim_names[-1]}"
            return f"This civil action arises from {joined} and seeks relief for the injuries caused by that conduct."
        return "This civil action seeks relief for unlawful conduct described in the factual allegations below."

    def _build_jurisdiction_statement(self, jurisdiction: str, authority_records: Any) -> str:
        citations = []
        for authority in _listify(authority_records):
            if not isinstance(authority, dict):
                continue
            citation = _clean_text(authority.get("citation") or "")
            if citation:
                citations.append(citation)
        jurisdiction_lower = str(jurisdiction or "").lower()
        if jurisdiction_lower in {"federal", "us", "united states"}:
            if citations:
                return f"This Court has subject-matter jurisdiction under federal law, including {citations[0]}, because Plaintiff alleges violations arising under the Constitution, laws, or treaties of the United States."
            return "This Court has subject-matter jurisdiction under 28 U.S.C. § 1331 because Plaintiff alleges claims arising under federal law."
        return "This Court has subject-matter jurisdiction because the claims arise under the governing law identified in this pleading."

    def _build_venue_statement(self, district: Optional[str], division: Optional[str]) -> str:
        district_text = _clean_text(district)
        division_text = _clean_text(division)
        if district_text and division_text:
            return f"Venue is proper in the {division_text} Division of the District of {district_text} because a substantial part of the events or omissions giving rise to these claims occurred there."
        if district_text:
            return f"Venue is proper in the District of {district_text} because a substantial part of the events or omissions giving rise to these claims occurred there."
        return "Venue is proper in this Court because a substantial part of the events or omissions giving rise to these claims occurred in this judicial district."

    def _build_signature_block(self, parties: Dict[str, List[str]]) -> Dict[str, str]:
        plaintiff_name = parties.get("plaintiffs", ["Plaintiff"])[0]
        return {
            "name": plaintiff_name,
            "contact": "Address and contact information to be completed.",
        }

    def _compose_legal_standard(self, claim_name: str, legal_standard_elements: List[Dict[str, str]], authorities: List[Dict[str, Any]]) -> str:
        if legal_standard_elements:
            citations = [item.get("citation", "") for item in legal_standard_elements if item.get("citation")]
            if citations:
                return f"To state {claim_name}, Plaintiff must allege facts satisfying the governing elements recognized by {citations[0]}."
            return f"To state {claim_name}, Plaintiff must allege facts satisfying each essential element recognized by governing law."
        if authorities:
            authority = authorities[0]
            reference = authority.get("citation") or authority.get("title") or "governing authority"
            return f"The legal standard for {claim_name} is supplied by {reference}."
        return f"Plaintiff alleges facts sufficient to satisfy the legal standard for {claim_name}."

    def _infer_jurisdiction(self, legal_graph) -> str:
        if legal_graph is None:
            return "federal"
        jurisdictions = []
        for element in getattr(legal_graph, "elements", {}).values():
            jurisdiction = _clean_text(getattr(element, "jurisdiction", ""))
            if jurisdiction:
                jurisdictions.append(jurisdiction)
        if not jurisdictions:
            return "federal"
        first = jurisdictions[0].lower()
        if first in {"us", "united states"}:
            return "federal"
        return first

    def _dedupe(self, values: Iterable[str]) -> List[str]:
        seen = set()
        deduped = []
        for value in values:
            cleaned = _clean_text(value)
            if not cleaned:
                continue
            marker = cleaned.lower()
            if marker in seen:
                continue
            seen.add(marker)
            deduped.append(cleaned)
        return deduped

    def _write_docx(self, draft: Dict[str, Any], destination: Path) -> None:
        if not HAS_DOCX:
            raise RuntimeError("DOCX export requires python-docx to be installed")
        document = DocxDocument()
        normal_style = document.styles["Normal"]
        normal_style.font.name = "Times New Roman"
        normal_style.font.size = Pt(12)

        for line in [draft.get("court_header"), draft.get("caption", {}).get("division_line")]:
            if not line:
                continue
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(str(line))
            run.bold = True

        caption = draft.get("caption", {}) if isinstance(draft.get("caption"), dict) else {}
        table = document.add_table(rows=1, cols=2)
        table.columns[0].width = Inches(4.75)
        table.columns[1].width = Inches(2.0)
        table.cell(0, 0).text = (
            f"{', '.join(draft.get('parties', {}).get('plaintiffs', []) or ['Plaintiff'])}, Plaintiff,\n"
            f"v.\n"
            f"{', '.join(draft.get('parties', {}).get('defendants', []) or ['Defendant'])}, Defendant."
        )
        table.cell(0, 1).text = f"Case No.: {caption.get('case_number', '________________')}"

        title = document.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.add_run("COMPLAINT").bold = True

        self._docx_section(document, "Nature of the Action", [draft.get("nature_of_action")])
        party_lines = []
        for plaintiff in draft.get("parties", {}).get("plaintiffs", []):
            party_lines.append(f"Plaintiff {plaintiff} is an aggrieved party bringing this action.")
        for defendant in draft.get("parties", {}).get("defendants", []):
            party_lines.append(f"Defendant {defendant} is alleged to be responsible for the conduct described below.")
        self._docx_section(document, "Parties", party_lines)
        self._docx_section(document, "Jurisdiction and Venue", [draft.get("jurisdiction_statement"), draft.get("venue_statement")])
        self._docx_section(document, "Factual Allegations", draft.get("factual_allegations", []), numbered=True)

        claims_heading = document.add_paragraph()
        claims_heading.add_run("Claims for Relief").bold = True
        for claim in _listify(draft.get("legal_claims")):
            claim_heading = document.add_paragraph()
            claim_heading.add_run(str(claim.get("title") or "Claim")).bold = True
            if claim.get("description"):
                document.add_paragraph(_clean_sentence(claim.get("description")))
            if claim.get("legal_standard"):
                paragraph = document.add_paragraph()
                paragraph.add_run("Legal Standard: ").bold = True
                paragraph.add_run(_clean_sentence(claim.get("legal_standard")))
            for item in _listify(claim.get("legal_standard_elements")):
                paragraph = document.add_paragraph(style="List Bullet")
                text = _clean_sentence(item.get("element") or "")
                citation = _clean_text(item.get("citation") or "")
                paragraph.add_run(f"{text} ({citation})" if citation else text)
            for fact in _listify(claim.get("supporting_facts")):
                paragraph = document.add_paragraph(style="List Bullet")
                paragraph.add_run(_clean_sentence(fact))
            exhibit_labels = ", ".join(exhibit.get("label", "") for exhibit in _listify(claim.get("supporting_exhibits")) if exhibit.get("label"))
            if exhibit_labels:
                document.add_paragraph(f"Supported by {exhibit_labels}.")

        self._docx_section(document, "Prayer for Relief", draft.get("requested_relief", []), bulleted=True)
        exhibit_lines = []
        for exhibit in _listify(draft.get("exhibits")):
            line = f"{exhibit.get('label', 'Exhibit')} - {exhibit.get('title', 'Supporting exhibit')}"
            if exhibit.get("reference"):
                line = f"{line} ({exhibit['reference']})"
            exhibit_lines.append(line)
        self._docx_section(document, "Exhibits", exhibit_lines)
        document.add_paragraph("Respectfully submitted,")
        document.add_paragraph(draft.get("signature_block", {}).get("name") or "Plaintiff")
        document.save(str(destination))

    def _docx_section(self, document, heading: str, paragraphs: Sequence[Any], *, numbered: bool = False, bulleted: bool = False) -> None:
        heading_paragraph = document.add_paragraph()
        heading_paragraph.add_run(heading).bold = True
        for index, value in enumerate(_listify(paragraphs), 1):
            text = _clean_sentence(value)
            if not text:
                continue
            if numbered:
                document.add_paragraph(f"{index}. {text}")
            elif bulleted:
                paragraph = document.add_paragraph(style="List Bullet")
                paragraph.add_run(text)
            else:
                document.add_paragraph(text)

    def _write_pdf(self, draft: Dict[str, Any], destination: Path) -> None:
        if not HAS_REPORTLAB:
            raise RuntimeError("PDF export requires reportlab to be installed")
        styles = getSampleStyleSheet()
        centered = ParagraphStyle("ComplaintCentered", parent=styles["Heading2"], alignment=TA_CENTER)
        section = ParagraphStyle("ComplaintSection", parent=styles["Heading3"], spaceBefore=10, spaceAfter=6)
        body = ParagraphStyle("ComplaintBody", parent=styles["BodyText"], leading=15)
        story = []
        story.append(Paragraph(str(draft.get("court_header") or "IN THE COURT OF COMPETENT JURISDICTION"), centered))
        if draft.get("caption", {}).get("division_line"):
            story.append(Paragraph(str(draft["caption"]["division_line"]), centered))
        story.append(Spacer(1, 0.2 * inch))

        left_caption = (
            f"{', '.join(draft.get('parties', {}).get('plaintiffs', []) or ['Plaintiff'])}, Plaintiff,<br/>"
            f"v.<br/>{', '.join(draft.get('parties', {}).get('defendants', []) or ['Defendant'])}, Defendant."
        )
        right_caption = f"Case No.: {draft.get('case_number') or '________________'}"
        caption_table = Table([[Paragraph(left_caption, body), Paragraph(right_caption, body)]], colWidths=[4.5 * inch, 2.0 * inch])
        caption_table.setStyle(TableStyle([
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("BOX", (0, 0), (-1, -1), 0.5, colors.black),
            ("INNERGRID", (0, 0), (-1, -1), 0.5, colors.black),
            ("LEFTPADDING", (0, 0), (-1, -1), 8),
            ("RIGHTPADDING", (0, 0), (-1, -1), 8),
        ]))
        story.append(caption_table)
        story.append(Spacer(1, 0.2 * inch))
        story.append(Paragraph("COMPLAINT", centered))

        self._pdf_section(story, section, body, "Nature of the Action", [draft.get("nature_of_action")])
        party_lines = []
        for plaintiff in draft.get("parties", {}).get("plaintiffs", []):
            party_lines.append(f"Plaintiff {plaintiff} is an aggrieved party bringing this action.")
        for defendant in draft.get("parties", {}).get("defendants", []):
            party_lines.append(f"Defendant {defendant} is alleged to be responsible for the conduct described below.")
        self._pdf_section(story, section, body, "Parties", party_lines)
        self._pdf_section(story, section, body, "Jurisdiction and Venue", [draft.get("jurisdiction_statement"), draft.get("venue_statement")])
        self._pdf_section(story, section, body, "Factual Allegations", draft.get("factual_allegations", []), numbered=True)

        story.append(Paragraph("Claims for Relief", section))
        for claim in _listify(draft.get("legal_claims")):
            story.append(Paragraph(str(claim.get("title") or "Claim"), section))
            if claim.get("description"):
                story.append(Paragraph(_clean_sentence(claim.get("description")), body))
            if claim.get("legal_standard"):
                story.append(Paragraph(f"<b>Legal Standard:</b> {_clean_sentence(claim.get('legal_standard'))}", body))
            for item in _listify(claim.get("legal_standard_elements")):
                text = _clean_sentence(item.get("element") or "")
                citation = _clean_text(item.get("citation") or "")
                story.append(Paragraph(f"• {text} ({citation})" if citation else f"• {text}", body))
            for fact in _listify(claim.get("supporting_facts")):
                story.append(Paragraph(f"• {_clean_sentence(fact)}", body))
            exhibit_labels = ", ".join(exhibit.get("label", "") for exhibit in _listify(claim.get("supporting_exhibits")) if exhibit.get("label"))
            if exhibit_labels:
                story.append(Paragraph(f"Supported by {exhibit_labels}.", body))

        self._pdf_section(story, section, body, "Prayer for Relief", draft.get("requested_relief", []), bulleted=True)
        exhibit_lines = []
        for exhibit in _listify(draft.get("exhibits")):
            line = f"{exhibit.get('label', 'Exhibit')} - {exhibit.get('title', 'Supporting exhibit')}"
            if exhibit.get("reference"):
                line = f"{line} ({exhibit['reference']})"
            exhibit_lines.append(line)
        self._pdf_section(story, section, body, "Exhibits", exhibit_lines)
        story.append(Spacer(1, 0.2 * inch))
        story.append(Paragraph("Respectfully submitted,", body))
        story.append(Paragraph(draft.get("signature_block", {}).get("name") or "Plaintiff", body))
        document = SimpleDocTemplate(str(destination), pagesize=letter)
        document.build(story)

    def _pdf_section(self, story, section_style, body_style, heading: str, paragraphs: Sequence[Any], *, numbered: bool = False, bulleted: bool = False) -> None:
        story.append(Paragraph(heading, section_style))
        for index, value in enumerate(_listify(paragraphs), 1):
            text = _clean_sentence(value)
            if not text:
                continue
            if numbered:
                story.append(Paragraph(f"{index}. {text}", body_style))
            elif bulleted:
                story.append(Paragraph(f"• {text}", body_style))
            else:
                story.append(Paragraph(text, body_style))